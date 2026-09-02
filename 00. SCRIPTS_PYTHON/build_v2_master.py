import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH

doc = docx.Document()

# Márgenes
for s in doc.sections:
    s.top_margin = Inches(1)
    s.bottom_margin = Inches(1)
    s.left_margin = Inches(1)
    s.right_margin = Inches(1)

styles = doc.styles

# Título Principal (Title)
title_style = styles['Title']
title_style.font.name = 'Calibri'
title_style.font.size = Pt(18)
title_style.font.bold = True
title_style.font.color.rgb = RGBColor(0x1A, 0x0A, 0x2E)

# Título 1 (Heading 1)
h1_style = styles['Heading 1']
h1_style.font.name = 'Calibri'
h1_style.font.size = Pt(15)
h1_style.font.bold = True
h1_style.font.color.rgb = RGBColor(0x0D, 0x1B, 0x4A)

# Título 2 (Heading 2)
h2_style = styles['Heading 2']
h2_style.font.name = 'Calibri'
h2_style.font.size = Pt(13.5)
h2_style.font.bold = True
h2_style.font.color.rgb = RGBColor(0x2E, 0x5B, 0x88)

# Título 3 (Heading 3) - Para Estilo y Prompts numéricos (para el Índice de Word)
h3_style = styles['Heading 3']
h3_style.font.name = 'Calibri'
h3_style.font.size = Pt(12)
h3_style.font.bold = True
h3_style.font.color.rgb = RGBColor(0x1A, 0x0A, 0x2E)

# Título del Documento
p_title = doc.add_paragraph('REPOSITORIO MAESTRO: ESTILOS DE IMÁGENES CON IA V2', style='Title')
p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER

p_sub = doc.add_paragraph()
p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
r_sub = p_sub.add_run('Manual y Repositorio Integral para el Aula (Mayores 60+): Crear en cualquier IA (< 200 car.) y Transformar en Gemini\n4 Bloques Maestros | 48+ Estilos Catalogados | 120+ Ejemplos Reales y 26 Prácticas de Edición (100% Numerados en Título 3 para el Índice de Word)')
r_sub.italic = True
r_sub.font.size = Pt(11)
r_sub.font.color.rgb = RGBColor(0x55, 0x55, 0x55)

doc.add_paragraph('----------------------------------------------------------------------------------------------------')

p_metodo = doc.add_paragraph()
r_m1 = p_metodo.add_run('METODOLOGÍA DOCENTE DEL CURSO: EL PUENTE IA → GEMINI\n')
r_m1.bold = True
r_m1.font.size = Pt(12.5)
r_m1.font.color.rgb = RGBColor(0x0D, 0x1B, 0x4A)
p_metodo.add_run(
    '1. CREAR EN TU IA FAVORITA: Copia el ejemplo o diapositiva (<200 caracteres) en tu generador de imágenes (Bing Image Creator, Gemini, Grok, DALL-E...) y crea la foto.\n'
    '2. DESCARGAR TU CREACIÓN: Guarda o descarga la imagen generada en tu ordenador (pulsa la flecha ↓ o botón derecho \"Guardar imagen\").\n'
    '3. TRANSFORMAR EN GEMINI: Abre gemini.google.com, sube tu foto descargada con el botón (+) y pégale la orden de edición que te proponemos en el reto práctico de cada estilo.\n'
    '4. TU CREACIÓN PROPIA: ¡Invéntate tú una modificación final y sorprende a tus compañeros!'
)

doc.add_paragraph('----------------------------------------------------------------------------------------------------')

def add_header_1(text):
    return doc.add_heading(text, level=1)

def add_header_2(text):
    return doc.add_heading(text, level=2)

def add_style_multi(num, title, desc, examples_list):
    p_style = doc.add_heading(f'{num} {title}', level=3)
    for r in p_style.runs:
        r.font.size = Pt(12.5)
        r.font.color.rgb = RGBColor(0x1A, 0x0A, 0x2E)
    
    p_desc = doc.add_paragraph()
    p_desc.add_run('• Descripción original: ').bold = True
    p_desc.add_run(desc)
    
    clean_num = num.strip('.')
    
    for idx, item in enumerate(examples_list, 1):
        if len(item) == 4:
            short_title, label, ex, bridge_edit = item
            assert len(ex) < 200, f'Prompt over 200 chars [{len(ex)}]: {ex}'
            
            p_prompt_h = doc.add_heading(f'{clean_num}.{idx}. {short_title}', level=3)
            for r in p_prompt_h.runs:
                r.font.size = Pt(11.5)
                r.font.color.rgb = RGBColor(0x00, 0x5A, 0x9E)
            
            p_ex = doc.add_paragraph()
            p_ex.add_run(f'• {label} listo para copiar en IA: ').bold = True
            r_ex_text = p_ex.add_run(f'\"{ex}\"')
            r_ex_text.font.color.rgb = RGBColor(0x00, 0x5A, 0x9E)
            r_ex_text.bold = True
            p_ex.add_run(f'  [{len(ex)} car.]').italic = True
            
            p_bridge = doc.add_paragraph()
            r_b_h = p_bridge.add_run(f'• 🧪 Reto Práctico de Edición ({label}) con Gemini:\n')
            r_b_h.bold = True
            r_b_h.font.color.rgb = RGBColor(0x8B, 0x00, 0x8B)
            p_bridge.add_run('Descarga la imagen generada en tu IA, súbela a Gemini (gemini.google.com) y escríbele lo siguiente:\n')
            
            r_edit = p_bridge.add_run(f'\"{bridge_edit}\"\n')
            r_edit.bold = True
            r_edit.font.color.rgb = RGBColor(0x2E, 0x5B, 0x88)
            
            r_call = p_bridge.add_run('👉 Ahora te toca a ti: ¡Haz tú una modificación que se te ocurra y sorpréndenos!')
            r_call.bold = True
            r_call.font.color.rgb = RGBColor(0x0D, 0x1B, 0x4A)
        else:
            short_title, label_series, slides_list, bridge_edit = item
            p_prompt_h = doc.add_heading(f'{clean_num}.{idx}. {short_title}', level=3)
            for r in p_prompt_h.runs:
                r.font.size = Pt(11.5)
                r.font.color.rgb = RGBColor(0x00, 0x5A, 0x9E)
            
            p_intro = doc.add_paragraph()
            p_intro.add_run(f'• {label_series}:\n').bold = True
            
            for s_idx, (s_title, s_ex) in enumerate(slides_list, 1):
                assert len(s_ex) < 200, f'Prompt over 200 chars [{len(s_ex)}]: {s_ex}'
                p_slide = doc.add_paragraph()
                p_slide.add_run(f'   - Diapositiva {s_idx} ({s_title}): ').bold = True
                r_ex = p_slide.add_run(f'\"{s_ex}\"')
                r_ex.font.color.rgb = RGBColor(0x00, 0x5A, 0x9E)
                r_ex.bold = True
                p_slide.add_run(f'  [{len(s_ex)} car.]').italic = True
            
            p_bridge = doc.add_paragraph()
            r_b_h = p_bridge.add_run('• 🧪 Reto Práctico de Edición de Serie con Gemini:\n')
            r_b_h.bold = True
            r_b_h.font.color.rgb = RGBColor(0x8B, 0x00, 0x8B)
            p_bridge.add_run('Descarga cualquiera de las fotos de esta serie o sube el cuadro original a Gemini (gemini.google.com) y escríbele lo siguiente:\n')
            
            r_edit = p_bridge.add_run(f'\"{bridge_edit}\"\n')
            r_edit.bold = True
            r_edit.font.color.rgb = RGBColor(0x2E, 0x5B, 0x88)
            
            r_call = p_bridge.add_run('👉 Ahora te toca a ti: ¡Haz tú una modificación que se te ocurra y sorpréndenos!')
            r_call.bold = True
            r_call.font.color.rgb = RGBColor(0x0D, 0x1B, 0x4A)
            
    doc.add_paragraph()

# ==============================================================================
# BLOQUE 1: BÁSICOS
# ==============================================================================
add_header_1('I. ESTILOS BÁSICOS')

add_header_2('1. Estilos Artísticos')

add_style_multi('1.1.', 'REALISTA (FOTORREALISMO)',
    'Imágenes detalladas y fotorrealistas que imitan la realidad con nitidez fotográfica.',
    [
        ('Faro al Amanecer', 'Ejemplo A (Paisaje Costero)',
         'Fotografía profesional 8K de un viejo faro costero resistiendo el oleaje al amanecer, luz dorada, enfoque nítido en la piedra húmeda, texturas hiperrealistas, National Geographic.',
         'Conserva exactamente el faro y las rocas del mar, pero cambia el amanecer por una noche de tormenta dramática con rayos iluminando el cielo oscuro y quita la espuma blanca más alta de las olas.'),
        ('Bodegón de Frutas', 'Ejemplo B (Naturaleza Muerta Clásica)',
         'Fotografía fotorrealista de un bodegón clásico con frutas variadas, jarra de vino de latón grabada, copa medio llena, cuchillo de madera y pan rústico sobre paño de lino arrugado.',
         'Mantén el bodegón de frutas y la jarra de latón, pero cambia el pan rústico por un tazón de cerámica lleno de fresas con nata y haz que la luz del fondo sea más dorada y cálida.'),
        ('Cervantes y Shakespeare en el Siglo de Oro', 'Ejemplo C (Encuentro Histórico Fotorrealista)',
         'Retrato fotorrealista 8K de Miguel de Cervantes y William Shakespeare conversando en una biblioteca del siglo XVI con pergaminos, iluminación de plumas y velas, trajes de época nítidos.',
         'Conserva a Cervantes y Shakespeare conversando en la biblioteca del siglo XVI, pero haz que sobre la mesa de madera grande aparezca un moderno ordenador portátil abierto luminoso.')
    ]
)

add_style_multi('1.2.', 'IMPRESIONISTA Y POST-IMPRESIONISTA',
    'Pinceladas sueltas, colores vibrantes y énfasis en la luz, similar a Monet, Van Gogh o Klimt.',
    [
        ('Campo de Amapolas', 'Ejemplo A (Paisaje al Óleo)',
         'Pintura al óleo de un campo de amapolas silvestres bajo un cielo nublado, estilo impresionista similar a Monet, pinceladas sueltas y gruesas, textura de lienzo visible, colores vivos.',
         'Mantén el estilo impresionista de óleo y el hermoso campo de amapolas, pero añade a una mujer con un vestido blanco de época y una sombrilla paseando tranquilamente por el centro del prado.'),
        ('Serie Carrusel: Monumentos de Madrid con Cielos de los 7 Grandes Maestros',
         'Repertorio maestro de 7 variaciones artísticas de monumentos emblemáticos de Madrid bajo los cielos inconfundibles de los grandes genios de la pintura mundial',
         [
             ('1. Vincent van Gogh (Estanque del Retiro)', 'Fotografía del Monumento a Alfonso XII en el Estanque del Retiro de Madrid, pero con un cielo de óleo impresionista con espirales amarillas y azules al estilo de Vincent van Gogh, reflejos.'),
             ('2. Claude Monet (Palacio de Cristal)', 'Fotografía del Palacio de Cristal del Retiro en Madrid junto al lago, bajo un cielo sereno impresionista con niebla suave y luz de amanecer pastel estilo Claude Monet, reflejos en el agua.'),
             ('3. J.M.W. Turner (Puerta de Alcalá)', 'Fotografía de la Puerta de Alcalá de Madrid al amanecer, envuelta en una niebla dorada romántica y luz etérea que difumina el sol entre nubes dramáticas al estilo del pintor J.M.W. Turner, 8K.'),
             ('4. Caspar David Friedrich (Templo de Debod)', 'Fotografía del Templo de Debod de Madrid al crepúsculo, bajo un cielo romántico sombrío y melancólico sobre un mar de nubes bajas y luz mística al estilo de Caspar David Friedrich, misterio.'),
             ('5. John Constable (Palacio Real y Jardines)', 'Fotografía del Palacio Real de Madrid desde los Jardines del Campo del Moro, bajo un cielo nublado dinámico y cambiante con nubes blancas y grises al estilo naturalista de John Constable.'),
             ('6. Edvard Munch (Plaza Mayor al Atardecer)', 'Fotografía de la Plaza Mayor de Madrid al atardecer, bajo un cielo dramático expresionista de nubes onduladas en tonos rojo fuego, rosa y azul marino oscuro al estilo de Edvard Munch, 8K.'),
             ('7. Giorgio de Chirico (Puerta del Sol desierta)', 'Fotografía de la Puerta del Sol de Madrid desierta, bajo un cielo surrealista amplio y vacío de color azul verdoso con luz fría, sombras alargadas y arquitectura onírica estilo De Chirico.')
         ],
         'Conserva el monumento de Madrid elegido y las aguas o suelo tranquilos, pero cambia el estilo del cielo por un cielo nocturno moderno lleno de fuegos artificiales multicolores.'),
        ('Serie Carrusel Magistral: El Cuadro de Aguirre de la Puerta de San Vicente bajo 9 Genios del Arte',
         'Recreación y repintado por IA de la célebre obra del siglo XVIII de Giménez Aguirre de la Puerta de San Vicente en el río Manzanares (Museo de Historia de Madrid) al estilo de 9 maestros de la pintura universal',
         [
             ('1. Obra Original (Giménez Aguirre - S. XVIII)', 'Pintura al óleo costumbrista del siglo XVIII del pintor Aguirre de la Puerta de San Vicente en el río Manzanares de Madrid, carruajes de caballos, lavanderas y vista clásica clasicista, 8K.'),
             ('2. Estilo Vincent van Gogh (Post-Impresionismo)', 'Pintura de la Puerta de San Vicente de Madrid del siglo XVIII, pero al óleo estilo Vincent van Gogh con pinceladas gruesas en espiral, cielos vibrantes en azul y amarillo y texturas dinámicas.'),
             ('3. Estilo Pablo Picasso (Cubismo y Vanguardia)', 'Recreación de la Puerta de San Vicente y carruajes del siglo XVIII en Madrid en estilo cubista de Pablo Picasso, perspectivas geométricas facetadas simultáneas y paleta de tonos tierra y ocre.'),
             ('4. Estilo Leonardo da Vinci (Renacimiento y Sfumato)', 'Pintura de la Puerta de San Vicente de Madrid en estilo renacentista de Leonardo da Vinci, técnica de sfumato suave atmosférico, paisaje difuminado al fondo y paleta de tonos ámbar sepia.'),
             ('5. Estilo Salvador Dalí (Surrealismo Onírico)', 'Pintura surrealista de la Puerta de San Vicente de Madrid al estilo de Salvador Dalí, con arcos clásicos alargados en un desierto onírico, sombras infinitas y luz crepuscular misteriosa, 8K.'),
             ('6. Estilo Diego Velázquez (Barroco Realista S. XVII)', 'Pintura al óleo de la Puerta de San Vicente de Madrid en estilo de Diego Velázquez, realismo costumbrista, claroscuro natural, atmósfera aérea capturada en el Manzanares y pincelada maestra.'),
             ('7. Estilo El Greco (Manierismo Místico)', 'Recreación de la Puerta de San Vicente de Madrid en el estilo místico de El Greco, figuras humanas alargadas y espirituales, cielo tormentoso dramático y paleta vibrante en cian y gris plata.'),
             ('8. Estilo Joaquín Sorolla (Luminismo Español)', 'Pintura al óleo de la Puerta de San Vicente de Madrid y el río Manzanares en estilo luminista de Joaquín Sorolla, luz solar blanca deslumbrante, reflejos puros en el agua y colores vivos.'),
             ('9. Estilo Sandro Botticelli (Renacimiento Italiano)', 'Recreación de la Puerta de San Vicente de Madrid en el elegante estilo renacentista de Botticelli, líneas curvas delicadas, figuras alegóricas de perfil y paleta de colores primaverales.'),
             ('10. Estilo Tiziano (Renacimiento Veneciano)', 'Pintura al óleo de la Puerta de San Vicente de Madrid en estilo renacentista veneciano de Tiziano, colores intensos rojo rubí y oro cálido, pincelada suelta en ropajes y luz de atardecer real.')
         ],
         'Sube la foto original del cuadro de la Puerta de San Vicente de Aguirre a Gemini y escríbele: \"Repinta este cuadro histórico del siglo XVIII transformando el paisaje del río Manzanares en un puerto marítimo veneciano con góndolas y reflejos dorados en el agua al estilo de Sorolla y Tiziano.\"'),
        ('Klimt: Dama en Pan de Oro', 'Ejemplo D (Simbolismo Dorado)',
         'Pintura estilo Gustav Klimt de una dama elegante con manto labrado en pan de oro brillante y motivos geométricos mosaico, fondo dorado ornamental, mosaicos y óleo rico en textura.',
         'Conserva el estilo de Gustav Klimt con pan de oro y los motivos de mosaico, pero cambia a la dama clásica por un hermoso gato persa blanco sentado sobre un cojín imperial de terciopelo verde.')
    ]
)

add_style_multi('1.3.', 'SURREALISTA',
    'Elementos oníricos y abstractos, inspirados en Dalí o atmósferas imposibles.',
    [
        ('Relojes en Desierto', 'Ejemplo A (Sueño Clásico)',
         'Pintura surrealista de relojes de bolsillo flotando y derritiéndose en un desierto iluminado por dos lunas, estilo Salvador Dalí, atmósfera onírica, detalles nítidos y misteriosos.',
         'Conserva el desierto onírico y el cielo con dos lunas, pero elimina los relojes derretidos y en su lugar pon grandes libros abiertos flotando cuyas páginas doradas se convierten en pájaros.'),
        ('Cuadros en Ruinas en Desierto', 'Ejemplo B (Galería Surrealista al Aire Libre)',
         'Pintura surrealista de marcos clásicos dorados de cuadros gigantes plantados en medio de las dunas de un desierto al atardecer, ruinas de galería de arte al aire libre, luz dramática.',
         'Conserva las dunas del desierto al atardecer y los marcos de cuadros clásicos en ruinas, pero haz que dentro del marco más grande del centro se vea un océano azul cristalino con olas saltando.')
    ]
)

add_style_multi('1.4.', 'CÓMIC / MANGA Y HALFTONE RETRO',
    'Dibujos con líneas definidas, colores planos, o tramado de puntos al estilo del cómic retro y manga anime.',
    [
        ('Tren Nevado Manga', 'Ejemplo A (Anime de Línea Nítida)',
         'Ilustración de cómic manga de un tren de pasajeros cruzando un bosque nevado por un puente de hierro, línea de tinta negra nítida, colores planos vivos, sombreado cel-shading, anime.',
         'Conserva el puente de hierro, la nieve del bosque y el estilo manga de tinta, pero cambia el tren moderno por una locomotora de vapor clásica echando una nube de humo rosa brillante por la chimenea.'),
        ('Serie Carrusel: Retratos en Grabado Halftone Retro Pop Art',
         'Variaciones monocromáticas con técnica de tramado de puntos estilo periódico antiguo o cómic clásico',
         [
             ('Capitán de Barco', 'Retrato en estilo cómic Halftone retro de tramado de puntos de un anciano capitán de barco con gorra naval y barba poblada, alto contraste de tinta negra sobre fondo claro, grabado pop art.'),
             ('Mujer Glamour Años 50', 'Retrato estilo Halftone pop art de tramado de puntos de una mujer elegante de los años 50 con ondas en el pelo, maquillaje nítido y mirada expresiva, tinta negra en contraste puro, cómic retro.'),
             ('Caballo Blanco al Galope', 'Ilustración dinámica en estilo Halftone de tramado de puntos de un majestuoso caballo blanco galopando por un prado verde luminoso bajo un cielo cian con trama de puntos, cómic pop art retro.')
         ],
         'Conserva el estilo de grabado Halftone retro de puntos de tinta negra, pero cambia el personaje por el retrato de un astronauta sonriente con su casco espacial de visor reflejante.')
    ]
)

add_style_multi('1.5.', 'PIXEL ART',
    'Gráficos de baja resolución similares a los videojuegos retro.',
    [
        ('Cafetería Retro 16-bit', 'Ejemplo real',
         'Escena en Pixel Art 16-bit de una acogedora cafetería de barrio de noche con luces de neón en la ventana, estilo videojuego retro arcade, paleta de colores cálidos y píxeles nítidos.',
         'Conserva la estética de videojuego de 16-bit y la fachada de la cafetería, pero cambia la noche por una mañana soleada y añade a un gato naranja durmiendo plácidamente en el alféizar de la ventana.')
    ]
)

add_header_2('2. Estilos Fotográficos')

add_style_multi('2.1.', 'HDR (HIGH DYNAMIC RANGE)',
    'Mayor contraste y detalles con gran impacto en luces y sombras.',
    [
        ('Calles de Toledo', 'Ejemplo A (Calle Adoquinada tras Lluvia)',
         'Fotografía HDR de alto rango dinámico de las calles adoquinadas de Toledo tras la lluvia, máximo contraste en reflejos de farolas, microdetalles en la piedra, fotorrealismo 8K nítido.',
         'Conserva exactamente los reflejos en los adoquines mojados de Toledo y el alto contraste HDR, pero quita la lluvia y añade a un farolero antiguo encendiendo una farola con una larga vara de madera.'),
        ('Bodegón del Mar', 'Ejemplo B (Marisco y Pescado Fresco)',
         'Bodegón HDR fotorrealista de pescados frescos, langostas y mejillones sobre mesa rústica, botella de vino blanco enfriada, copas de cristal brillante, limones y microdetalles en hielo.',
         'Conserva el marisco fresco y las copas brillantes con efecto HDR, pero cambia la botella de vino blanco por una elegante sopera de cerámica azul humeante llena de sopa de marisco caliente.')
    ]
)

add_style_multi('2.2.', 'BLANCO Y NEGRO (MONOCROMÁTICO)',
    'Imágenes monocromáticas con énfasis en sombras y luces.',
    [
        ('Anciano Violinista', 'Ejemplo A (Retrato de Estudio)',
         'Fotografía clásica en blanco y negro de un anciano violinista tocando en la ventana, alto contraste de claroscuro, iluminación lateral sobre sus manos, grano fino de película de 35mm.',
         'Conserva el blanco y negro, la luz dramática y la figura del anciano violinista, pero haz que su violín tenga un sutil color ámbar dorado en contraste con el resto del cuadro monocromático.'),
        ('La Joven de la Perla y Monalisa en Claroscuro', 'Ejemplo B (Fusión de Retratos Célebres)',
         'Fotografía artística en blanco y negro de La Joven de la Perla de Vermeer junto a la Monalisa de Da Vinci posando juntas en un estudio fotográfico del siglo XIX, claroscuro y elegancia.',
         'Conserva el retrato monocromático en blanco y negro de las dos mujeres clásicas, pero haz que el famoso pendiente de perla brille intensamente con un color azul zafiro luminoso.')
    ]
)

add_style_multi('2.3.', 'ANÁLOGO / VINTAGE Y DOCUMENTAL HISTÓRICO',
    'Efecto de cámaras antiguas, con grano, tonos sepia y nostalgia de épocas pasadas o grandes eventos documentados.',
    [
        ('Picnic Años 70', 'Ejemplo A (Familia con Cámara Kodak)',
         'Fotografía analógica vintage años 70 de una familia merendando en el campo junto a un coche clásico, cámara Kodak de carrete, tonos cálidos nostálgicos, grano suave y destello de luz.',
         'Conserva el estilo vintage años 70 de carrete Kodak, la familia y el campo verde, pero cambia el coche clásico por una preciosa furgoneta Volkswagen camperizada en color azul pastel marino.'),
        ('La Dama del Té', 'Ejemplo B (Bistró en París 1900)',
         'Fotografía analógica vintage de una dama elegante de 1900 tomando té en una mesa de bistró clásico con la Torre Eiffel de fondo, tonos sepia cálidos, luz de tarde, cámara de época.',
         'Conserva la dama clásica, el té y el estilo vintage con la Torre Eiffel al fondo, pero cambia la taza de té por una copa de champán con burbujas y añade un ramo de rosas en la mesa.'),
        ('Serie Carrusel: El Desastre del Dirigible Hindenburg (1937)',
         'Recreación documental histórica del famoso dirigible sobrevolando Nueva Jersey',
         [
             ('Aproximación Solemne', 'Fotografía documental analógica de 1937 en blanco y negro del gigantesco dirigible Hindenburg surcando el cielo nublado cerca de su mástil de amarre en Lakehurst, grano histórico retro.'),
             ('Momento del Estallido', 'Fotografía periodística histórica en sepia del instante en que la cola del dirigible Hindenburg arde en llamas sobre el campo de aterrizaje, multitudes y camiones de bomberos al fondo.'),
             ('Descenso en Llamas', 'Fotografía de prensa dramática de los años 30 de la estructura de aluminio del Hindenburg colapsando contra el suelo en una bola de fuego, alto contraste monocromático y humo real.')
         ],
         'Conserva la fotografía histórica en blanco y negro del dirigible volando majestuosamente, pero quita cualquier rastro de humo o llamas y haz que el cielo esté totalmente despejado y pacífico.'),
        ('Serie Carrusel: Un Día en las Carreras del Hipódromo de la Castellana',
         'Crónica costumbrista de la alta sociedad madrileña en el antiguo Hipódromo a principios del siglo XX',
         [
             ('Damas con Sombreros en la Grada', 'Fotografía analógica sepia 1910 de damas elegantes de la alta sociedad con grandes sombreros y vestidos largos charlando en la tribuna del Hipódromo de la Castellana en Madrid, luz suave.'),
             ('Caballos al Galope en Pista', 'Fotografía documental de época de jinetes montados en caballos purasangre galopando en la recta final del Hipódromo de la Castellana, polvo levantándose en la pista, público al fondo.'),
             ('Carruajes en el Exterior', 'Fotografía vintage 1912 de lujosos carruajes de caballos y primeros automóviles aparcados a las puertas del Hipódromo en el Paseo de la Castellana, caballeros con chistera y bastón.')
         ],
         'Conserva el estilo fotográfico sepia de 1910 y la tribuna del hipódromo con las damas elegantes, pero añade en la pista central una moderna bicicleta de carreras de color rojo brillante.')
    ]
)

add_style_multi('2.4.', 'MACRO Y ORFEBRERÍA HIPERREALISTA',
    'Capturas de primerísimo plano con gran nivel de detalle milimétrico, texturas nobles, esmaltes y desenfoque del fondo.',
    [
        ('Gota de Rocío', 'Ejemplo A (Naturaleza en Pétalo de Rosa)',
         'Fotografía macro de primerísimo plano de una gota de rocío equilibrada sobre un pétalo de rosa roja, lente macro de 100mm, reflejo del sol dentro de la gota, fondo desenfocado bokeh.',
         'Mantén el primerísimo plano macro y la gota de rocío sobre la rosa roja, pero haz que dentro del reflejo cristalino de la gota de agua se vea con claridad la silueta diminuta de una mariquita.'),
        ('Joyas Mujer Hindú', 'Ejemplo B (Primerísimo Plano de Joyas)',
         'Retrato macro en primerísimo plano de una mujer hindú con maquillaje vibrante, joyas tradicionales, mang tikka en la frente, pendientes decorados y mirada intensa con bindi, 8K.',
         'Conserva el primerísimo plano de las joyas doradas y la mirada vibrante de la mujer hindú, pero cambia el color de las piedras preciosas del collar y pendientes a un verde esmeralda brillante.'),
        ('Juego Café Porcelana', 'Ejemplo C (Detalle Porcelana Fina)',
         'Primerísimo plano macro de una tetera inglesa de porcelana fina con motivos florales delicados y filos dorados, tazas a juego, azucarero y bandeja de dulces con luz suave interior.',
         'Mantén la tetera de porcelana fina y los motivos florales en primerísimo plano, pero añade saliendo del pico de la tetera una delicada nube de vapor que forme un pequeño corazón en el aire.'),
        ('Caras con Pan de Oro', 'Ejemplo D (Textura y Relieve Artístico)',
         'Primerísimo plano macro de un rostro humano sereno cuya piel está cubierta con delicadas láminas de pan de oro fracturadas, reflejos metálicos dorados sobre fondo negro mate, 8K.',
         'Conserva el rostro macro cubierto con láminas de pan de oro fracturadas, pero haz que en las grietas entre el oro se vea una suave luz interna de color cian turquesa resplandeciente.'),
        ('Joyas Imposibles Corona y Huevo', 'Ejemplo E (Alta Orfebrería)',
         'Fotografía macro 8K de una corona imperial de oro blanco con incrustaciones milimétricas de diamantes y rubíes flotando sobre un huevo de cristal tallado, microdetalles y reflejos puros.',
         'Conserva la corona imperial y el huevo de cristal tallado con todos sus reflejos macro, pero cambia el color de los rubíes rojos por espléndidos zafiros de color azul marino profundo.'),
        ('Serie Carrusel: Máscaras Cloisonné de Turquesa y Oro (Estilo Veneciano)',
         'Alta orfebrería con técnica de esmaltado Cloisonné, relieves de miniaturas en turquesa y oro sobre rostros expresivos',
         [
             ('Dama Veneciana Cloisonné', 'Retrato fotorrealista 8K de una mujer con máscara veneciana de alta orfebrería en estilo Cloisonné turquesa y oro brillante, microesculturas en relieve y reflejos de luz sobre esmalte y piedras.'),
             ('Carnaval de Venecia en Relieve', 'Primer plano macro de una dama con máscara de porcelana turquesa labrada en filigrana de oro Cloisonné, relieves en miniatura inspirados en canales venecianos, ojos expresivos y luz de atardecer.')
         ],
         'Descarga la foto de la mujer con máscara turquesa y oro, súbela a Gemini y escríbele: \"Colorea las partes de oro amarillo en un brillante platino plateado, y haz que en los ojos tras la máscara brille un misterioso fulgor de color amatista púrpura.\"')
    ]
)

# 3. Estilos 3D y Generativos
add_header_2('3. Estilos 3D y Generativos')
add_style_multi('3.1.', 'RENDER 3D',
    'Modelos con iluminación y texturas realistas.',
    [
        ('Biblioteca Flotante Espacial', 'Ejemplo real',
         'Render 3D de una biblioteca futurista flotando en el espacio, motor Unreal Engine 5, iluminación global con reflejos en el suelo de cristal, texturas metálicas hiperdetalladas y limpias.',
         'Conserva el render 3D hiperdetallado de la biblioteca flotante en el espacio y sus reflejos en el suelo, pero cambia el vacío estelar del fondo por una nebulosa brillante llena de estrellas de colores.')
    ]
)

add_style_multi('3.2.', 'LOW-POLY',
    'Estilo simplificado con pocas caras poligonales.',
    [
        ('Zorro en Montaña', 'Ejemplo real',
         'Ilustración en estilo Low-Poly 3D de un zorro naranja sentado en la cumbre de una montaña nevada, caras poligonales planas geométricas con aristas nítidas, paleta de colores suaves.',
         'Mantén el estilo Low-Poly geométrico del zorro naranja en la cumbre, pero quita la nieve de la montaña y conviértela en una colina de hierba primaveral salpicada de pequeñas margaritas poligonales.')
    ]
)

add_style_multi('3.3.', 'CYBERPUNK',
    'Estética futurista con luces neón y paisajes urbanos oscuros.',
    [
        ('Callejón Lluvioso Tokio', 'Ejemplo real',
         'Escena futurista estilo Cyberpunk de un callejón nocturno en Tokio bajo la lluvia, reflejos de letreros de neón rosa y cian sobre el suelo mojado, humo de vapor y alta tecnología.',
         'Conserva el callejón de Tokio con lluvia y los reflejos de neón rosa en el asfalto, pero añade a un robot humanoides amable con un paraguas transparente paseando por el centro del callejón.')
    ]
)

add_style_multi('3.4.', 'STEAMPUNK E INVENTOS DE ÉPOCA',
    'Mezcla de elementos victorianos con tecnología retro-futurista, electricidad bobinas y cobre.',
    [
        ('Locomotora Voladora Londres', 'Ejemplo A (Aventuras Victorianas)',
         'Escena estilo Steampunk victoriano de una locomotora voladora con alas de latón y engranajes mecánicos sobrevolando Londres, tuberías de cobre, vapor de caldera y luz de tarde dorada.',
         'Conserva la locomotora voladora con engranajes de latón sobre Londres, pero cambia la luz de tarde dorada por un cielo crepuscular lleno de globos aerostáticos de estilo victoriano en el horizonte.'),
        ('Serie Carrusel: El Laboratorio del Genio Nikola Tesla',
         'Recreación histórica y retro-futurista del laboratorio del gran inventor de la corriente alterna',
         [
             ('Bobinas de Rayos en el Laboratorio', 'Fotografía sepia de época del laboratorio de Nikola Tesla en Colorado Springs en 1899, gigantescas bobinas eléctricas emitiendo arcos y rayos de electricidad azul entre maquinaria de cobre.'),
             ('Tesla Junto a la Bobina Gigante', 'Retrato histórico en blanco y negro de Nikola Tesla sentado serenamente leyendo un libro en su laboratorio mientras enormes arcos eléctricos cruzan el aire sobre su cabeza, iluminación clara.'),
             ('Maquinaria y Motores de Latón', 'Fotografía fotorrealista 8K de los generadores y prototipos de motores magnéticos de Tesla sobre mesas de madera rústica, engranajes de latón pulido, chispa eléctrica en penumbra.')
         ],
         'Conserva el laboratorio histórico de Tesla con las bobinas eléctricas de latón y cobre, pero haz que los rayos de electricidad cambien del color azul o blanco a un color verde esmeralda neón.')
    ]
)

# 4. Estilos Conceptuales
add_header_2('4. Estilos Conceptuales y Abstracción')
add_style_multi('4.1.', 'MINIMALISTA',
    'Imágenes con pocos elementos y colores planos.',
    [
        ('Árbol Solitario Rojo', 'Ejemplo real',
         'Diseño gráfico minimalista de un árbol solitario de hojas rojas sobre una colina blanca, abundante espacio en blanco negativo, paleta de dos colores puros, composición serena y limpia.',
         'Mantén el estilo minimalista, el espacio en blanco y el árbol de hojas rojas en la colina, pero añade una pequeña silueta limpia de un pájaro negro volando justo encima de la copa del árbol.')
    ]
)

add_style_multi('4.2.', 'ABSTRACTO',
    'Formas y colores sin una estructura reconocible.',
    [
        ('Pinceladas Azules y Oro', 'Ejemplo real',
         'Pintura de arte abstracto puro en tonos azules profundos y dorados, formas geométricas y corrientes fluidas entrelazadas, texturas de óleo empastado con relieves dramáticos y armonía.',
         'Conserva la textura de óleo empastado y la composición abstracta fluida, pero cambia la paleta de colores azules profundos por tonos cálidos otoñales como rojo rubí, naranja mandarina y oro metálico.')
    ]
)

add_style_multi('4.3.', 'GLITCH ART Y DISTORSIÓN DIGITAL',
    'Efectos de errores digitales y distorsiones visuales.',
    [
        ('Escultura Mármol Glitch', 'Ejemplo A (Rostro Fragmentado)',
         'Ilustración en estilo Glitch Art de un rostro humano de mármol que se fragmenta en error de señal digital, líneas de escaneo CRT, aberración cromática RGB cian y magenta, bloques rotos.',
         'Conserva la escultura de mármol y las líneas de escaneo CRT retro, pero cambia las distorsiones de aberración cromática cian y magenta por destellos digitales en color verde esmeralda luminoso.'),
        ('Mujeres de Mármol Glitch Digital', 'Ejemplo B (Estatuas Clásicas Fragmentadas)',
         'Fotografía de esculturas clásicas de mujeres griegas en mármol blanco puro distorsionadas por ondas de error informático, bloques de píxeles flotantes y cortes geométricos de luz cian.',
         'Conserva las estatuas griegas de mármol blanco y el corte de error informático, pero añade una corona de laurel de oro intacta y brillante sobre la cabeza de la estatua central.')
    ]
)

# ==============================================================================
# BLOQUE 2: SOFISTICADOS Y ESPECÍFICOS
# ==============================================================================
add_header_1('II. ESTILOS DE IMÁGENES GENERADAS POR IA MÁS SOFISTICADOS Y ESPECÍFICOS')

# 1. Estilos Artísticos Avanzados
add_header_2('1. Estilos Artísticos Avanzados')

add_style_multi('1.1.', 'BARROCO DIGITAL Y MITOLOGÍA',
    'Inspirado en el arte barroco, con un uso dramático de la luz y la sombra (claroscuro), detalles ornamentales y composiciones complejas.',
    [
        ('Biblioteca Renacentista', 'Ejemplo A (Interior Teatral)',
         'Pintura estilo Barroco Digital de una biblioteca renacentista con globos terráqueos, iluminación claroscuro dramática tipo Caravaggio, oro labrado, drapeados oscuros de terciopelo y luz.',
         'Conserva la biblioteca de estilo Barroco Digital y el claroscuro dramático, pero añade sobre una de las mesas de madera tallada un gran candelabro de plata encendido con tres velas.'),
        ('Huevo Joya Fabergé', 'Ejemplo B (Orfebrería Imperial)',
         'Fotografía fotorrealista de un huevo joya estilo Fabergé exquisitamente colocado sobre un tocador de lujo barroco, ricamente decorado con oro y piedras preciosas, brillo de joyas y luz tenue.',
         'Conserva el huevo joya de oro y pedrería sobre el tocador de lujo, pero haz que la parte superior del huevo esté abierta y salga del interior un pequeño pájaro mecánico de oro cantando.'),
        ('Serie Carrusel: Dioses del Olimpo en el Trono',
         'Recreaciones majestuosas de las deidades de la mitología clásica en entornos de palacio celestial',
         [
             ('Zeus con el Rayo', 'Pintura estilo Barroco Digital hiperrealista del dios Zeus sentado en su trono de oro en el Monte Olimpo, barba blanca solemne, empuñando un rayo brillante, nubes tormentosas y águila imperial.'),
             ('Atenea con Armadura de Oro', 'Retrato barroco hiperrealista de la diosa Atenea con casco y armadura de oro labrado, sosteniendo una lanza y un escudo con el rostro de Medusa, lechuza posada en su hombro, luz sagrada.'),
             ('Poseidón entre las Olas', 'Pintura dramática del dios Poseidón emergiendo del mar agitado con su tridente de oro, corona de conchas, corceles de espuma blanca saltando entre olas y claroscuro celestial.')
         ],
         'Conserva el estilo majestuoso barroco del dios en su trono celestial, pero añade junto al estrado un pequeño perro golden retriever tranquilo mirando pacíficamente hacia la luz del Olimpo.')
    ]
)

add_style_multi('1.2.', 'ROCOCÓ MODERNO',
    'Colores pastel, formas curvas y un aire de lujo y romanticismo aristocrático aplicado a entornos y personajes.',
    [
        ('Salón de Té Rococó', 'Ejemplo A (Ambiente Aristocrático)',
         'Pintura estilo Rococó Moderno de un salón de té con lámparas de araña y espejos curvos, paleta pastel en rosa y celadón, ornamentos florales de oro, lujo romántico y sedas delicadas.',
         'Mantén el salón de té estilo Rococó y la paleta de colores pastel rosa y oro, pero añade en una de las sillas de seda a un elegante caniche blanco con un lazo rosa al cuello.'),
        ('María Antonieta Vistiéndose', 'Ejemplo B (Retrato Histórico de Palacio)',
         'Pintura Rococó de María Antonieta vistiéndose ayudada por su doncella en un salón de palacio con terciopelo y espejos dorados, corsé ajustado, vestido de seda pastel y luz suave por la ventana.',
         'Conserva a María Antonieta y el lujoso salón de palacio rococó, pero cambia el vestido de seda pastel rosa por un espectacular vestido en color azul cielo real con bordados de perlas.'),
        ('Hora Té Harrods', 'Ejemplo C (Costumbrismo Inglés S. XIX)',
         'Escena fotorrealista en un salón de té estilo Harrods siglo XIX con dos damas elegantes tomando té, vajilla inglesa de porcelana floral, bandeja de cuatro pisos con macarons y pastas de té.',
         'Conserva las dos damas elegantes y la bandeja de dulces del té de Harrods, pero añade en la mesa central una pequeña tarta de cumpleaños con velas encendidas rodeada de flores de azúcar.')
    ]
)

add_style_multi('1.3.', 'NEO-EXPRESIONISMO',
    'Representaciones emocionales y distorsionadas con pinceladas fuertes y colores intensos.',
    [
        ('Violinista Apasionado', 'Ejemplo real',
         'Pintura de Neo-Expresionismo de un violinista tocando con pasión en un escenario, trazos agresivos de espátula, paleta de colores puros y salvajes contrastados, óleo empastado rugoso.',
         'Conserva los trazos de espátula y al violinista, pero cambia el color del fondo oscuro por una explosión de pinceladas amarillas, naranjas y rojas que transmitan pura energía y alegría.')
    ]
)

add_style_multi('1.4.', 'ARTE GENERATIVO PARAMÉTRICO Y PROPORCIÓN ÁUREA',
    'Imágenes creadas a partir de fórmulas matemáticas, con patrones geométricos y estructuras algorítmicas.',
    [
        ('Esfera Curvas Bézier', 'Ejemplo A (Geometría Paramétrica)',
         'Arte Generativo Paramétrico de una esfera hueca formada por infinitas curvas de Bézier entrelazadas, teselaciones geométricas de líneas finas de luz dorada sobre fondo azul noche zen.',
         'Mantén la esfera paramétrica geométrica de líneas finas, pero cambia la luz dorada de las curvas de Bézier por un gradiente luminoso arcoíris que cambie suavemente de color por toda la esfera.'),
        ('Proporción Áurea en Espiral', 'Ejemplo B (Fibonacci en la Naturaleza y Arquitectura)',
         'Ilustración geométrica hiperrealista de la espiral de la Proporción Áurea de Fibonacci sobrepuesta perfectamente en la concha del caracol Nautilus de nácar y una escalera de caracol dorada.',
         'Conserva la geometría sagrada de la espiral de proporción áurea y la concha de Nautilus, pero añade en el centro de la espiral una brillante perla blanca iluminada con luz suave.')
    ]
)

# 2. Estilos Fotográficos y Cinemáticos
add_header_2('2. Estilos Fotográficos y Cinemáticos')

add_style_multi('2.1.', 'CINEMATOGRÁFICO NOIR / ÉPICO Y GRANDES BATALLAS',
    'Alto contraste, sombras profundas, encuadres dramáticos y atmósfera de gran película de Hollywood o superproducción.',
    [
        ('Detective en Despacho', 'Ejemplo A (Cine Negro / Noir Clásico)',
         'Fotograma de película estilo Cinematográfico Noir de un detective mirando por la ventana de su despacho por la noche, sombras en persianas, humo de cigarrillo y luz lateral de farola.',
         'Conserva al detective en blanco y negro y las sombras de las persianas, pero elimina el humo del cigarrillo y pon en su mano una elegante taza de café humeante.'),
        ('Barco Pirata Épico', 'Ejemplo B (Superproducción de Aventuras)',
         'Fotograma de película épica de un majestuoso barco pirata navegando en mar agitado, capitán en proa con sombrero de tres picos y espada, marineros en cubierta bajo cielo dramático con rayos.',
         'Conserva el majestuoso barco pirata y el mar agitado, pero haz que del agua salga el enorme tentáculo de un calamar gigante de color rojo esmeralda salpicando espuma junto al barco.'),
        ('Moby Dick en Tormenta', 'Ejemplo C (Clímax Cinematográfico en el Mar)',
         'Fotograma cinematográfico épico del Capitán Ahab gritando en un mar de tormenta mientras lanza un arpón a la boca gigante de la ballena blanca Moby Dick, olas salpicando, 8K.',
         'Conserva la ballena gigante Moby Dick y el mar en tormenta, pero haz que el cielo en lo alto se abra en un círculo claro por donde entre un rayo de sol dorado iluminando al barco.'),
        ('Serie Carrusel: La Armada Invencible en Batalla (1588)',
         'Superproducción histórica naval sobre la gran flota española del siglo XVI en combate en el Canal de la Mancha',
         [
             ('Galeones Españoles a Toda Vela', 'Fotograma cinematográfico épico 8K de una flota de majestuosos galeones españoles de la Armada Invencible en 1588 navegando a toda vela con banderas cruz de Borgoña al viento, mar agitado.'),
             ('Cañoneo entre Niebla y Humo', 'Pintura cinematográfica dramática del combate naval de la Armada Invencible, galeones disparando sus cañones entre niebla, humo de pólvora, fuego en la cubierta y cielo de tormenta rugiendo.'),
             ('El Almirante en la Toldilla', 'Retrato cinematográfico fotorrealista del almirante español con armadura grabada y capa contemplando la batalla desde la cubierta de popa de su galeón bajo un cielo rojo crepuscular.')
         ],
         'Conserva los galeones navales de 1588 navegando en el mar agitado, pero elimina todo rastro de batalla o humo de cañón y haz que naveguen pacíficamente bajo una mañana soleada con gaviotas.'),
        ('Carreras de Cuadrigas en el Circo Romano', 'Ejemplo E (Acción Épica Antigüedad)',
         'Fotograma cinematográfico de acción épica 8K de una carrera de cuadrigas de caballos en el Circo Máximo de Roma, aurigas gritando con riendas en mano, arena levantándose y público en gradas.',
         'Conserva la carrera de cuadrigas romanas y la arena levantándose, pero cambia el color de la capa del auriga ganador de color blanco a un púrpura imperial romano resplandeciente.'),
        ('Serie Carrusel: Retratos de Perfil Épicos en Oro y Sombra (Cleopatra y Napoleón)',
         'Imágenes de perfil con hermoso diseño realista en oro, claroscuro dramático, textura hiperdetallada y fondo oscuro insondable',
         [
             ('Perfil Épico de Cleopatra en Oro', 'Retrato de perfil de Cleopatra en estilo cinematográfico 8K, fondo oscuro dramático, hermoso diseño dorado ornamental realista en corona y joyas, textura hiperdetallada y claroscuro épico.'),
             ('Perfil Épico de Napoleón Bonaparte', 'Retrato de perfil de Napoleón Bonaparte en estilo cinematográfico épico 8K, fondo oscuro en penumbra, uniforme con bordados de oro realista, microdetalles de textura facial y claroscuro de estudio.')
         ],
         'Conserva el perfil épico y todo el diseño dorado ornamental en alta textura, pero cambia el fondo oscuro y vacío por el interior de una monumental sala real con lámparas de araña desenfocadas.')
    ]
)

add_style_multi('2.2.', 'HIPERREALISMO AI',
    'Imágenes que imitan la textura de la piel, reflejos en hielo o cristal y microdetalles al nivel de fotografía de alta gama.',
    [
        ('Marinero Barba Blanca', 'Ejemplo A (Retrato Fotorrealista de Textura)',
         'Retrato de Hiperrealismo AI de un anciano marinero con barba blanca y gorro de lana, microtextura en poros y arrugas, reflejo de un barco en sus ojos, iluminación suave de estudio 8K.',
         'Conserva el retrato hiperrealista y el rostro del marinero con su barba blanca, pero cambia el gorro de lana de color oscuro por un clásico gorro de capitán de barco color blanco y azul navy.'),
        ('Marisco y Vino', 'Ejemplo B (Bodegón Gastronómico Hiperrealista)',
         'Bodegón hiperrealista 8K de pescados frescos, langostas, mejillones y gambas sobre mesa rústica, botella de vino blanco enfriada y dos copas de cristal brillante, reflejos y limones cortados.',
         'Conserva el marisco hiperrealista y las copas brillantes, pero añade alrededor de la langosta varias rodajas de limón amarillo espolvoreadas con perejil fresco muy verde.')
    ]
)

add_style_multi('2.3.', 'ESTÉTICA ANALOG HORROR',
    'Efectos de distorsión VHS, ruido visual y paletas de colores inquietantes inspiradas en el terror retro.',
    [
        ('Parque Infantil Nocturno', 'Ejemplo real',
         'Escena de estética Analog Horror años 80 de un parque infantil vacío por la noche visto por monitor de seguridad, cinta VHS con ruido estático y líneas de distorsión de tracking, oscuro.',
         'Mantén el efecto retro VHS y el parque infantil nocturno, pero añade en la esquina superior derecha las típicas letras digitales en verde de una grabadora analógica que digan \"REC 23:45 PM\".')
    ]
)

add_style_multi('2.4.', 'FOTOGRAFÍA DE ILUMINACIÓN REMBRANDT Y RETRATO HISTÓRICO',
    'Uso de luz suave con sombras triangulares o claroscuro costumbrista, imitando técnicas de pintores clásicos y personajes ilustres.',
    [
        ('Mujer Mayor Leyendo', 'Ejemplo A (Retrato Solemnidad Clásica)',
         'Retrato fotográfico con Iluminación Rembrandt de una mujer mayor leyendo un libro antiguo, triángulo claro de luz sobre su mejilla en sombra, fondo en penumbra y tonos cálidos solemnes.',
         'Conserva el retrato con luz Rembrandt de la mujer mayor y el fondo en penumbra, pero haz que el libro antiguo que lee esté abierto y emita una suave luz dorada mágica hacia su rostro.'),
        ('Bodegón de Matanza', 'Ejemplo B (Bodegón Costumbrista Rústico)',
         'Bodegón fotorrealista de productos de matanza sobre mesa rústica de madera, jarra de vino de barro, vasos, ristra de ajos y pan rústico con iluminación Rembrandt tenue de claroscuro lateral.',
         'Mantén el bodegón rústico de matanza con la jarra de barro y los ajos, pero añade en el lado izquierdo una vela encendida que ilumine suavemente el pan rústico y el cuchillo.'),
        ('Jorge Juan en su Despacho Naval (Siglo XVIII)', 'Ejemplo C (Retrato Ilustrado Marítimo)',
         'Pintura hiperrealista con iluminación Rembrandt del célebre marino y científico español Jorge Juan en su despacho naval del siglo XVIII con compás, cartas náuticas y catalejo, claroscuro.',
         'Conserva el retrato de Jorge Juan en su despacho con luz Rembrandt y cartas náuticas, pero haz que a través del ventanal del fondo se vea un velero navegando bajo la luz de la luna clara.')
    ]
)

# 3. Estilos 3D y Futuristas
add_header_2('3. Estilos 3D y Futuristas')
add_style_multi('3.1.', 'BIOPUNK',
    'Un cruce entre cyberpunk y biotecnología, con paisajes orgánicos y sintéticos combinados.',
    [
        ('Laboratorio Botánico Futurista', 'Ejemplo real',
         'Escena estilo Biopunk de un laboratorio botánico futurista donde plantas gigantes crecen integradas con cables bioluminiscentes verde esmeralda y membranas de cristal translúcido.',
         'Mantén el laboratorio botánico y las plantas biotecnológicas, pero cambia el color de los cables bioluminiscentes de verde esmeralda a un color cian turquesa resplandeciente.')
    ]
)

add_style_multi('3.2.', 'FUTURISMO RETRO (CASSETTE FUTURISM)',
    'Visión del futuro desde la perspectiva de los años 70 y 80, con dispositivos análogos y estéticas de ciencia ficción clásica.',
    [
        ('Cabina Espacial Años 80', 'Ejemplo real',
         'Escena de Cassette Futurism de la cabina de mando de una nave espacial años 80, pantallas CRT de fósforo ámbar con gráficos vectoriales, botones físicos de colores y chasis color marfil.',
         'Conserva la cabina retro-futurista y los botones físicos de colores, pero cambia lo que se ve a través de la ventana frontal: en vez del espacio, pon la vista del planeta Tierra desde la órbita.')
    ]
)

add_style_multi('3.3.', 'ARTE GIGERIANO (BIOMECÁNICO)',
    'Inspirado en H.R. Giger, con criaturas y estructuras biomecánicas inquietantes.',
    [
        ('Templo Alienígena Biomecánico', 'Ejemplo real',
         'Escena de arte Biomecánico estilo H.R. Giger del interior de un templo catedralicio alienígena, columnas de conductos orgánicos fundidos con acero pulido oscuro, paleta en gris plomo.',
         'Conserva la monumental arquitectura biomecánica del templo alienígena en gris plomo, pero añade un rayo de luz blanca pura entrando desde un lucernario en el techo hasta el centro del suelo.')
    ]
)

add_style_multi('3.4.', 'ARTE PROCEDURAL ORGÁNICO',
    'Estructuras y formas que evolucionan como organismos naturales pero creadas con algoritmos generativos.',
    [
        ('Escultura Coral Nácar', 'Ejemplo real',
         'Arte Procedural Orgánico de una escultura flotante con ramificaciones algorítmicas similares a corales marinos infinitos y micelio fúngico en crecimiento, relieves en nácar y oro suave.',
         'Conserva la escultura procedural flotante similar a un coral en nácar, pero añade en su interior pequeñas perlas esféricas de luz dorada que floten en el centro de las ramificaciones.')
    ]
)

# 4. Estilos Conceptuales
add_header_2('4. Estilos Conceptuales y Experimentales')
add_style_multi('4.1.', 'ARTE NEURAL DREAMSCAPE',
    'Paisajes y escenarios generados por IA con mezclas surrealistas de formas y colores en constante cambio.',
    [
        ('Océano Olas Cristal', 'Ejemplo real',
         'Arte Neural Dreamscape de un océano pacífico cuyas olas de cristal se disuelven suavemente en nubes y nebulosas de colores tornasolados fluyentes en el cielo, surrealismo etéreo y zen.',
         'Conserva el océano de cristal que se disuelve en las nubes tornasoladas, pero añade en el centro del mar una pequeña barca de madera con un farol encendido reflejándose en el agua zen.')
    ]
)

add_style_multi('4.2.', 'FRAGMENTACIÓN CUBISTA DIGITAL',
    'Interpretaciones cubistas con variaciones digitales y efectos de glitch.',
    [
        ('Violonchelo Facetado Cubista', 'Ejemplo real',
         'Pintura de Fragmentación Cubista Digital de un violonchelo y un rostro facetados en múltiples puntos de vista geométricos simultáneos con reflejos de glitch prismático en fondo blanco.',
         'Conserva el estilo cubista digital y las facetas geométricas del violonchelo y el rostro, pero cambia el fondo blanco por un fondo azul oscuro profundo para hacer resaltar los reflejos prismáticos.')
    ]
)

add_style_multi('4.3.', 'PSICODELIA FRACTAL',
    'Imágenes caleidoscópicas con patrones fractales infinitos y colores vibrantes.',
    [
        ('Búho Cósmico Mandalas', 'Ejemplo real',
         'Ilustración de Psicodelia Fractal de un búho cósmico cuyas alas desplegadas forman simetrías caleidoscópicas de mandalas con espirales logarítmicas, colores neón ultravioleta hipnóticos.',
         'Conserva la geometría fractal hipnótica y al búho cósmico, pero cambia la paleta de colores neón ultravioleta por una paleta elegante de tonos oro brillante, blanco puro y azul turquesa.')
    ]
)

add_style_multi('4.4.', 'HYPERPOP AESTHETIC',
    'Combinaciones de colores neón, distorsión visual y elementos futuristas de la cultura pop.',
    [
        ('Zapatillas Cromadas Neón', 'Ejemplo real',
         'Estética Hyperpop Y2K de unas zapatillas deportivas futuristas sobre una nube de plata cromada líquida, paleta saturada en rosa chicle y cian neón con esferas 3D flotantes y destellos.',
         'Conserva las zapatillas futuristas y las nubes de plata cromada líquida, pero añade detrás unas grandes gafas de sol futuristas flotantes reflejando los destellos en color rosa chicle.')
    ]
)

# ==============================================================================
# BLOQUE 3: AÚN MÁS AVANZADOS
# ==============================================================================
add_header_1('III. ESTILOS AÚN MÁS AVANZADOS, POST-HUMANOS Y RECREACIÓN DE MUNDOS PERDIDOS')

add_header_2('1. Estilos de IA Post-Humanos')
add_style_multi('1.1.', 'METAHUMANO SURREALISTA',
    'Fusión de elementos hiperrealistas con estructuras imposibles, como rostros que se desintegran en códigos fractales.',
    [
        ('Rostro Desintegrado Fractales', 'Ejemplo real',
         'Retrato Metahumano Surrealista de una mujer clásica 8K cuya piel se desintegra en partículas y códigos fractales de luz dorada flotando, mirada serena, fondo oscuro insondable.',
         'Conserva el rostro de la mujer y la desintegración fractal en luz dorada, pero haz que los códigos numéricos flotantes alrededor se conviertan en pequeñas mariposas de luz de color azul turquesa.')
    ]
)

add_style_multi('1.2.', 'POST-SINGULARIDAD ESTÉTICA',
    'Imágenes generadas con estructuras algorítmicas que simulan una visión de la IA después de la singularidad tecnológica.',
    [
        ('Ciudadela Cuántica Cristal', 'Ejemplo real',
         'Estética Post-Singularidad de una ciudadela cuántica de cristal inteligente flotando sobre anillos concéntricos de luz pura y campos de energía dorados, trascendencia digital y paz.',
         'Conserva la ciudadela cuántica flotando sobre los anillos de luz pura, pero añade varias naves espaciales con forma de anillo plateado entrando suavemente hacia el centro de la ciudadela.')
    ]
)

add_style_multi('1.3.', 'ARTESANÍA CUÁNTICA GENERATIVA',
    'Composiciones basadas en cálculos cuánticos, con patrones emergentes imposibles y miniaturas cristalinas en 3D.',
    [
        ('Rosa Cósmica Oro', 'Ejemplo A (Escultura Cuántica)',
         'Artesanía Cuántica Generativa de una rosa cósmica esculpida por ondas de superposición cuántica cristalizadas en láminas de oro puro y cobalto iridiscente, patrones de luz infinitos.',
         'Conserva la escultura de láminas de oro y cobalto en forma de rosa cósmica, pero añade una fina capa de rocío brillante en los pétalos, como si fueran diamantes reflejando las estrellas.'),
        ('Caja Música Nieve', 'Ejemplo B (Miniatura Cristalina)',
         'Artesanía fotorrealista de una delicada caja de música de oro y madera tallada que sostiene una bola de nieve de cristal con un paisaje invernal y copos flotando en su interior, luz dorada.',
         'Conserva la caja de música y la bola de nieve cristalina, pero haz que dentro de la bola de nieve el paisaje invernal tenga una casita de madera con luces encendidas y humo en la chimenea.')
    ]
)

add_style_multi('1.4.', 'EVOLUCIÓN NEURAL EXPANDIDA',
    'Figuras en constante transformación generadas por redes neuronales profundas, imitando la evolución orgánica.',
    [
        ('Metamorfosis Criatura Cristal', 'Ejemplo real',
         'Arte de Evolución Neural Expandida de una criatura mitológica en metamorfosis fluida entre alas de mariposa de cristal, pétalos de orquídea blanca y filamentos de sinapsis de luz azul.',
         'Conserva la metamorfosis fluida entre alas de cristal y pétalos de orquídea blanca, pero cambia los filamentos de sinapsis azul por filamentos en color oro cálido para darle un tono solar.')
    ]
)

add_header_2('2. Estilos de Computación Visual Avanzada')
add_style_multi('2.1.', 'ARTE TOPOLÓGICO DINÁMICO',
    'Uso de modelos matemáticos para generar formas que fluyen entre dimensions espaciales y fractales.',
    [
        ('Cinta Möbius Dicroica', 'Ejemplo real',
         'Arte Topológico Dinámico de una cinta de Möbius continua esculpida en cristal dicroico translúcido que fluye en dimensions infinitas, refracción de luz de arcoíris sobre fondo negro puro.',
         'Conserva la cinta de Möbius de cristal dicroico y sus reflejos arcoíris sobre el fondo negro, pero haz que en el centro del bucle flote pacíficamente una pequeña esfera dorada perfectamente pulida.')
    ]
)

add_style_multi('2.2.', 'VISUALIZACIÓN DE DATOS ARTÍSTICA',
    'Transformación de big data en patrones visuales complejos y artísticos, con colores y estructuras inspiradas en visualización de redes neuronales.',
    [
        ('Red Neuronal Cerebro', 'Ejemplo real',
         'Visualización de Datos Artística de una red neuronal del cerebro humano en 3D, millones de nodos y sinapsis de luz dorada y cian conectando datos sobre fondo azul zafiro oscuro, 8K.',
         'Mantén la red neuronal en 3D del cerebro y sus conexiones brillantes de luz dorada y cian, pero añade en el centro exacto de la red una brillante luz blanca que pulse con fuerza.')
    ]
)

add_style_multi('2.3.', 'HIPERTEXTURAS CUÁNTICAS',
    'Representaciones visuales que combinan hipertexturas fluidas y materiales imposibles, como líquidos que flotan en el vacío.',
    [
        ('Gotas Mercurio Flotantes', 'Ejemplo real',
         'Escultura de Hipertexturas Cuánticas con gotas de mercurio cromado líquido flotando en gravedad cero en el vacío, microrelieves fractales iridiscentes en su superficie y luz de sol de tarde.',
         'Conserva las gotas de mercurio cromado líquido flotando en gravedad cero con microrelieves, pero cambia el fondo vacío por el interior de una elegante sala de museo con luz cenital de galería.')
    ]
)

add_style_multi('2.4.', 'ESTÉTICA DEL CAMPO MORFOGÉNICO',
    'Inspirado en la teoría de los campos mórficos, generando estructuras visuales que evolucionan en patrones biológicos autoorganizados.',
    [
        ('Mandala Ondas Ámbar', 'Ejemplo real',
         'Estética del Campo Morfogénico de ondas de energía vital resonante formando un mandala orgánico autoorganizado en tonos de ámbar translúcido y verde clorofila brillante, armonía zen.',
         'Conserva las ondas biológicas y el mandala en verde clorofila y ámbar, pero haz que desde el centro del mandala emerja un pequeño brote verde con dos hojas perfectas creciendo hacia la luz.')
    ]
)

add_header_2('3. Estilos de Ficción Especulativa y Recreación Histórica de Mundos')

add_style_multi('3.1.', 'ARQUEOLOGÍA DIGITAL Y RECREACIÓN DE MUNDOS PERDIDOS OR MADRID ANTIGUO',
    'Imágenes generadas para simular descubrimientos o reconstruir milimétricamente civilizaciones, edificios históricos perdidos y monumentos del pasado.',
    [
        ('Templo Cibernético Sumergido', 'Ejemplo A (Ficción Especulativa)',
         'Arqueología Digital de un templo sumergido en el núcleo de un servidor olvidado, ruinas de piedra cibernética cubiertas de musgo bioluminiscente e ideogramas láser de luz turquesa.',
         'Conserva las ruinas del templo cibernético sumergido y el musgo bioluminiscente, pero añade en primer plano a un buzo científico observando los ideogramas láser con una linterna.'),
        ('Serie Carrusel: Construcción del Palacio de Cibeles (1908-1910)',
         'Recreación histórica progresiva en 5 etapas de la edificación monumental del Palacio de Comunicaciones de Madrid',
         [
             ('Fase 1: Cimientos y Estructura', 'Fotografía documental vintage 1908 en blanco y negro de los cimientos y primera estructura de acero y andamios de madera del Palacio de Cibeles en Madrid, obreros y carruajes, grano retro.'),
             ('Fase 2: Fachada Principal en Obras', 'Fotografía analógica de época de la construcción avanzada del Palacio de Cibeles de Madrid en 1909, fachada de piedra medio levantada con torres entre andamios de madera, caballos en plaza.'),
             ('Fase 3: Torre Central Levantándose', 'Fotografía histórica en blanco y negro 1909 de la torre central del Palacio de Cibeles cubierta de andamios de madera, arco principal terminado, primeros automóviles y carruajes en la calle.'),
             ('Fase 4: Esculturas y Reloj en Detalle', 'Fotografía vintage de 1910 del Palacio de Cibeles con el reloj y esculturas terminadas, alas laterales emergiendo de andamios de madera, damas y caballeros paseando junto a carruajes de caballos.'),
             ('Fase 5: Palacio y Fuente Terminados', 'Fotografía histórica 1910 en tono sepia del Palacio de Cibeles terminado frente a la Fuente de la diosa Cibeles en su carro, automóviles clásicos de la época y personas paseando por el adoquín.')
         ],
         'Conserva la fotografía histórica de 1910 del Palacio y la Fuente de Cibeles, pero cambia el cielo gris por una noche estrellada con la luna llena iluminando la piedra blanca y el agua.'),
        ('Serie Carrusel: Construcción del Edificio Metrópolis (1907-1911)',
         'Crónica visual en 3 etapas del levantamiento de uno de los edificios más emblemáticos de la Gran Vía y Alcalá',
         [
             ('Estructura y Andamios en la Esquina', 'Fotografía histórica 1908 de la esquina de Alcalá y Gran Vía de Madrid con la estructura del Edificio Metrópolis rodeada de tupidos andamios de madera, tranvías de caballos y peatones.'),
             ('Colocación de la Cúpula de Pizarra', 'Fotografía documental sepia de los obreros instalando la característica cúpula redonda y ornamentos de piedra en lo alto del Edificio Metrópolis de Madrid en 1910, luz de tarde clara.'),
             ('Inauguración Brillante', 'Fotografía fotorrealista de época del Edificio Metrópolis recién terminado en 1911, cúpula coronada por la estatua alada, elegantes transeúntes con chistera y carruajes en la calle Alcalá.')
         ],
         'Conserva la fotografía histórica de 1911 del Edificio Metrópolis recién terminado, pero haz que la cúpula superior esté iluminada con una cálida luz dorada que realce sus detalles arquitectónicos.'),
        ('Serie Carrusel: El Alcázar Real de Madrid Antes del Incendio de 1734',
         'Recreación histórica rigurosa de la antigua residencia real de los Austrias en el emplazamiento del actual Palacio Real',
         [
             ('Vista Monumental desde el Manzanares', 'Pintura histórica fotorrealista de 1700 del imponente Alcázar Real de Madrid con sus torres de pizarra y fachada clasicista visto desde el río Manzanares, carruajes y cielo despejado.'),
             ('La Plaza de Palacio en el Siglo XVII', 'Recreación fotorrealista de la plaza frente al Alcázar Real de Madrid en tiempos de Felipe IV, caballeros con capa y espada, damas con guardainfante y guardias en la puerta principal.')
         ],
         'Conserva la vista monumental del antiguo Alcázar Real de Madrid y sus torres de pizarra, pero añade sobre el cielo una bandada de palomas blancas sobrevolando el edificio.'),
        ('Proyecto de la Exposición Universal de 1888', 'Ejemplo E (Arquitectura Neoclásica Monumental)',
         'Pintura monumental neoclásica 8K de un proyecto para la Exposición Universal de 1888, bulevares con jardines y cúpulas, tranvías de caballos, multitudes elegantes y gran palacio al fondo.',
         'Conserva la gran avenida neoclásica y los pabellones de la Exposición de 1888, pero cambia la luz de tarde por una espectacular iluminación nocturna de época con farolas de gas encendidas.'),
        ('Puerto Fluvial del Manzanares y Esfera Armilar', 'Ejemplo F (Proyectos Ilustrados de Madrid)',
         'Ilustración histórica 8K de un gran puerto fluvial navegable en el río Manzanares de Madrid en el siglo XVIII con muelles de piedra y barcos, y una gran Esfera Armilar astronómica de bronce.',
         'Conserva los muelles de piedra del puerto en el Manzanares y la Esfera Armilar de bronce, pero añade un elegante velero de tres mástiles amarrado en el muelle principal descargando mercaderías.')
    ]
)

add_style_multi('3.2.', 'MISTICISMO ALGORÍTMICO',
    'Arte inspirado en traditions esotéricas reinterpretadas por redes neuronales, creando símbolos y patrones arcanos que parecen significativos pero son generados aleatoriamente.',
    [
        ('Altar Obsidiana Rúnico', 'Ejemplo real',
         'Arte de Misticismo Algorítmico de una esfera cósmica flotando en un altar de obsidiana pulida, rodeada por anillos rúnicos y símbolos arcanos de luz dorada y geometría sagrada.',
         'Conserva el altar de obsidiana pulida y los símbolos rúnicos de luz dorada, pero cambia la esfera cósmica flotante del centro por una gran pluma cristalina de color azul zafiro brillante.')
    ]
)

add_style_multi('3.3.', 'ARTESANÍA CYBERDRÓNICA',
    'Imágenes que combinan estructuras de ciberpunk extremo con formas biológicas emergentes, fusionando lo orgánico y lo mecánico en composiciones complejas.',
    [
        ('Microdrones Ala Libélula', 'Ejemplo real',
         'Artesanía Cyberdrónica de enjambres de microdrones de titanio pulido ensamblando un ala de libélula biónica ultraligera con membranas solares de cristal en un taller con luz cenital.',
         'Conserva los microdrones y el taller, pero cambia el ala de libélula biónica que están construyendo por el delicado engranaje de un reloj de bolsillo mecánico con piezas transparentes de cristal.')
    ]
)

add_style_multi('3.4.', 'LENGUAJE VISUAL DE IAS AUTÓNOMAS',
    'Creación de imágenes a partir de lenguajes visuales generados por inteligencias artificiales sin intervención humana.',
    [
        ('Monolito Negro Criptográfico', 'Ejemplo real',
         'Lenguaje Visual de IAs Autónomas en un monolito de metal mate negro infinito, grabado de arriba a abajo con un alfabeto criptográfico de luz blanca cegadora en geometría pura.',
         'Conserva el monolito de metal mate negro y los símbolos geométricos del alfabeto criptográfico, pero cambia el color de la luz cegadora blanca por una suave y elegante luz violeta neón.')
    ]
)

add_header_2('4. Estilos de Simulación y Meta-Realidad')
add_style_multi('4.1.', 'ESTÉTICA DE SIMULACIÓN DE REALIDADES ALTERNAS (RA/VR/AR)',
    'Creación de imágenes que imitan realidades paralelas, con filtros de percepción alterados y estructuras visuales extrañas.',
    [
        ('Ballenas Luz Madrid', 'Ejemplo real',
         'Escena de Realidad Aumentada AR en la Gran Vía de Madrid, hologramas gigantes de ballenas de luz azul flotando sobre el tráfico al atardecer, interfaz HUD sutil y moderna en el aire.',
         'Conserva los edificios de la Gran Vía de Madrid al atardecer y la interfaz HUD digital, pero cambia las ballenas de luz por gigantescos peces koi de luz dorada y naranja nadando en el aire.')
    ]
)

add_style_multi('4.2.', 'ARTE TRANSDIMENSIONAL',
    'Imágenes que parecen superpuestas en múltiples dimensiones, con perspectivas imposibles y efectos de profundidad hiperespacial.',
    [
        ('Escalinata Imposible Escher', 'Ejemplo real',
         'Arte Transdimensional de una escalinata de palacio que se despliega en perspectivas imposibles tipo Escher hacia el infinito, con cristales conectando cielos estrellados y amaneceres.',
         'Conserva las escaleras en perspectivas imposibles tipo Escher y los cristales flotantes, pero añade en los peldaños a una figura con una larga capa roja caminando boca abajo por el techo.')
    ]
)

add_style_multi('4.3.', 'PAISAJES DE LA MENTE IA',
    'Generación de escenarios abstractos que representan la \"psique\" de una IA, con formas que reflejan el procesamiento de información y conexiones neuronales.',
    [
        ('Santuario Cañones Cuarzo', 'Ejemplo real',
         'Paisaje de la Mente de una IA como un santuario interior sereno, cascadas de luz líquida dorada cayendo por cañones monumentales de cristal de cuarzo e infinitas sinapsis flotantes.',
         'Conserva el santuario interior con los cañones monumentales de cuarzo blanco, pero haz que la cascada en vez de luz dorada sea una cascada cristalina de color esmeralda reflejando arcoíris.')
    ]
)

add_style_multi('4.4.', 'ARTE DE LOS SUEÑOS LÚCIDOS GENERADOS',
    'Uso de técnicas de redes neuronales para capturar la esencia de sueños y estados alterados de conciencia en formas visuales.',
    [
        ('Escalera Mármol Luna', 'Ejemplo real',
         'Arte de Sueño Lúcido de una escalera de caracol de mármol blanco emergiendo de un mar en calma hacia una luna gigante y dorada en el cielo de crepúsculo, atmósfera mágica y pacífica.',
         'Conserva el mar en calma, la luna gigante dorada y la escalera de caracol, pero haz que en la cumbre de la escalera de mármol esté abierto un arco de piedra que conduzca directamente a un jardín de estrellas.')
    ]
)

# ==============================================================================
# BLOQUE 4: TALLER PRÁCTICO DE EDICIÓN Y TRANSFORMACIÓN CON GEMINI GRATUITO
# ==============================================================================
add_header_1('IV. TALLER PRÁCTICO DE EDICIÓN Y TRANSFORMACIÓN DE IMÁGENES CON GEMINI GRATUITO')

p_intro_b4 = doc.add_paragraph()
r_i4 = p_intro_b4.add_run(
    'ADVERTENCIA PEDAGÓGICA PARA EL AULA (MAYORES 60+):\n'
    'En los entornos técnicos y vídeos en inglés se conoce a la tecnología de visión avanzada y edición por instrucciones como \"Nano-banana\" (modelo Gemini 2.5 Flash Image Preview en Google AI Studio).\n'
    'Sin embargo, para que nuestros alumnos mayores puedan utilizarlo de forma gratuita, sencilla, en español y SIN necesidad de suscripciones de pago ni cuentas de desarrollador, todas estas 26 Prácticas de Taller se realizan o practican subiendo la fotografía directamente con el botón (+) en gemini.google.com o en la herramienta gratuita de creación de imágenes de Google.\n\n'
    'A continuación se presentan los 26 Retos Prácticos catalogados por módulos y listos para ejecutar en clase (100% excluidas las funciones de vídeo y animación marcadas con NO):'
)
r_i4.italic = True
r_i4.font.color.rgb = RGBColor(0x33, 0x33, 0x33)

doc.add_paragraph('----------------------------------------------------------------------------------------------------')

def add_taller_item(num_str, title, desc_taller, prompt_gemini):
    p_h = doc.add_heading(f'{num_str} {title}', level=3)
    for r in p_h.runs:
        r.font.size = Pt(12)
        r.font.color.rgb = RGBColor(0x1A, 0x0A, 0x2E)
    
    p_d = doc.add_paragraph()
    p_d.add_run('• Objetivo del Taller: ').bold = True
    p_d.add_run(desc_taller)
    
    p_g = doc.add_paragraph()
    p_g.add_run('• 🧪 Orden / Prompt listo para pegar en Gemini (gemini.google.com): ').bold = True
    r_gt = p_g.add_run(f'\"{prompt_gemini}\"')
    r_gt.bold = True
    r_gt.font.color.rgb = RGBColor(0x00, 0x5A, 0x9E)
    p_g.add_run(f'  [{len(prompt_gemini)} car.]').italic = True
    doc.add_paragraph()

add_header_2('1. Módulo A: Retoque Fotográfico y Mejora Cotidiana')
add_taller_item('4.1.1.', 'Mejora automática de foto oscura o quemada',
                'Corregir automáticamente el balance de blancos, la exposición y el contraste de una fotografía tomada en malas condiciones de luz.',
                'Mejora esta fotografía automáticamente: equilibra el balance de blancos, aclara las zonas oscuras en sombra y ajusta el contraste para que los colores se vean nítidos y naturales.')

add_taller_item('4.1.2.', 'Corrección e iluminación facial de retratos',
                'Eliminar imperfecciones temporales como acné, brillos o sombras duras en la piel, manteniendo los rasgos naturales del rostro.',
                'Retoca suavemente el rostro de esta foto: elimina pequeñas marcas o imperfecciones de la piel, suaviza la textura y aporta una iluminación facial clara y favorecedora de estudio.')

add_taller_item('4.1.3.', 'Restauración y coloreado de fotos antiguas',
                'Recuperar fotos familiares en blanco y negro, sepia o dañadas por el paso del tiempo, devolviéndoles la nitidez y un color realista 8K.',
                'Restaura esta fotografía antigua en blanco y negro: repara los arañazos y zonas borrosas, mejora la nitidez facial y colóreala con tonos realistas y naturales de alta definición.')

add_taller_item('4.1.4.', 'Foto de Perfil Profesional para LinkedIn / Currículum',
                'Transformar una foto informal (de vacaciones o callejera) en un retrato profesional de estudio con vestimenta elegante y fondo sobrio.',
                'Convierte esta foto informal en un retrato profesional para currículum: viste a la persona con una chaqueta americana elegante y cambia el fondo por una oficina moderna desenfocada.')

add_header_2('2. Módulo B: Modificación Mágica y Recomposición')
add_taller_item('4.2.1.', 'Eliminación de objetos o personas molestos',
                'Quitar elementos no deseados de una foto (turistas de fondo, farolas, cables, gafas de sol) rellenando el fondo de manera inteligente.',
                'Elimina de esta fotografía las gafas de sol y las personas que aparecen de fondo, rellenando el paisaje de forma coherente e imperceptible para que quede limpia.')

add_taller_item('4.2.2.', 'Añadir ropa y complementos al retrato',
                'Incorporar prendas o accesorios realistas a una persona en la foto, ajustándose a la iluminación y postura.',
                'Añade a la persona de la fotografía unas elegantes gafas de sol de montura negra y una chaqueta de cuero estilo retro perfectamente ajustada a su postura e iluminación natural.')

add_taller_item('4.2.3.', 'Cambio de look y peinado retro (Estilo Años 80)',
                'Modificar el corte de pelo, el peinado y el color capilar de una persona en foto para adaptarlo a una estética o década concreta.',
                'Cambia el peinado de la persona en esta foto por un voluminoso peinado con ondas estilo años 80, y dale un toque de color castaño cobrizo brillante natural.')

add_taller_item('4.2.4.', 'Probador virtual de ropa (E-commerce)',
                'Vestir de forma realista a una modelo en una foto utilizando otra imagen de referencia donde se muestre la prenda por separado.',
                'Observa la prenda de ropa de la segunda imagen de referencia y viste con ella exactamente a la modelo de la primera foto, respetando los pliegues y la luz.')

add_taller_item('4.2.5.', 'Reconstrucción de objetos o fotos rotas',
                'Reconstruir una imagen cortada por la mitad o reparar objetos rotos que aparezcan en ella (como el reflejo en un espejo fracturado).',
                'Reconstruye y completa los fragmentos que faltan en esta fotografía del espejo roto, restaurando el reflejo nítido y perfecto como si el espejo estuviera intacto y limpio.')

add_header_2('3. Módulo C: Arquitectura, Bocetos y 3D')
add_taller_item('4.3.1.', 'Restauración digital de monumentos históricos',
                'Reconstruir digitalmente ruinas o monumentos antiguos (como el Coliseo de Roma o el Alcázar real) para mostrarlos en su esplendor original.',
                'Reconstruye digitalmente este monumento histórico en ruinas: muestra exactamente cómo se vería en su época de máximo esplendor con todas sus columnas, mármoles y estatuas intactas.')

add_taller_item('4.3.2.', 'Reinterpretación fantástica de lugares conocidos',
                'Transformar una fotografía urbana o paisaje célebre en un escenario de ciencia ficción o fantasía pura.',
                'Transforma esta fotografía de París en un escenario de ciencia ficción fantástica: coloca la ciudad bajo una gigantesca cúpula de cristal submarina con reflejos y luz bioluminiscente.')

add_taller_item('4.3.3.', 'De Boceto a Foto Real (Sketch to Real 8K)',
                'Subir una foto de un dibujo a lápiz, boceto esquemático o acuarela simple y convertirlo en una fotografía fotorrealista 8K.',
                'Interpreta este boceto y dibujo esquemático y conviértelo en una fotografía fotorrealista 8K de alta definición, añadiendo texturas reales de materiales, sombras y luz solar de tarde.')

add_taller_item('4.3.4.', 'De Captura de Mapa a Vista Fotorrealista 3D',
                'Convertir una captura de pantalla de un plano, mapa o callejero topográfico en una vista aérea fotorrealista en relieve de la zona.',
                'Convierte esta captura de plano y mapa en una vista aérea fotorrealista 3D de la ciudad: levanta los edificios con texturas reales, calles arboladas y luz de atardecer.')

add_taller_item('4.3.5.', 'De Foto a Diseño de Tatuaje Minimalista',
                'Extraer la silueta y rasgos de una fotografía real para convertirla en un diseño artístico lineal de tatuaje en tinta negra.',
                'Transforma esta fotografía en un diseño para tatuaje minimalista: crea una ilustración limpia de línea fina en tinta negra pura sobre fondo blanco, con elegancia y trazo nítido.')

add_taller_item('4.3.6.', 'Extracción y Aislamiento 3D de un objeto',
                'Aislar un objeto o producto del fondo de una foto de manera impecable y generarlo con iluminación neutra como modelo de catálogo 3D.',
                'Aísla y recorta perfectamente el objeto central de esta fotografía, eliminando todo el fondo y presentándolo como un modelo 3D nítido sobre un fondo blanco puro de estudio.')

add_header_2('4. Módulo D: Diseño Gráfico, Tipografía y Creación de Personajes')
add_taller_item('4.4.1.', 'Creación de hoja de personaje (Model Sheet)',
                'A partir de la foto de un solo personaje, generar una lámina técnica con sus 3 vistas: frontal, perfil y trasera sobre fondo blanco.',
                'Genera una hoja de personaje (model sheet) coherente con 3 vistas sobre fondo blanco: muestra al personaje de esta foto en vista frontal exacta, de perfil derecho y de espaldas.')

add_taller_item('4.4.2.', 'Escenas coherentes manteniendo el personaje',
                'Utilizar al personaje recién creado para colocarlo en diferentes situaciones, aventuras y encuadres sin que cambie su rostro ni ropa.',
                'Coloca a exactamente este mismo personaje de referencia sentado leyendo en un acogedor café de París en día lluvioso, manteniendo idénticos sus rasgos faciales, peinado y vestimenta.')

add_taller_item('4.4.3.', 'Escena de acción a partir de boceto y personajes',
                'Combinar las imágenes de los personajes de referencia con un boceto esquemático para crear una composición dinámica de acción.',
                'Combina a los personajes de las fotos de referencia y colócalos exactamente en la posición que indica el boceto esquemático para generar una escena cinematográfica dramática de acción.')

add_taller_item('4.4.4.', 'Collage Múltiple Inteligente (Hasta 13 referencias)',
                'Poner a trabajar juntas múltiples imágenes de referencia (personajes, objetos, fondo, iluminación) y fusionarlas en un solo cuadro magistral.',
                'Fusiona armónicamente todas las imágenes de referencia adjuntas en una sola escena fotográfica 8K coherente, unificando la luz, la perspectiva y los reflejos en un cuadro fotorrealista.')

add_taller_item('4.4.5.', 'Textos con Texturas 3D (Efecto CGI)',
                'Escribir palabras o frases en pantalla aplicando a las letras la textura física de un objeto de referencia (lana, hielo, pan de oro, peluche).',
                'Escribe la palabra \"MADRID 2026\" en letras 3D gigantes, aplicando a cada letra exactamente la textura suave de lana o peluche del abrigo de la foto de referencia, luz de estudio.')

add_taller_item('4.4.6.', 'Diseño de tipografía a partir de referencia',
                'Imitar el estilo de letra o tipografía artística que aparece en una foto de referencia para escribir un nuevo título en español.',
                'Imita el estilo tipográfico exacto de la letra que aparece en la imagen de referencia para escribir con esa misma fuente y color el título \"TALLER DE INTELIGENCIA ARTIFICIAL\".')

add_taller_item('4.4.7.', 'Diseño de Logotipos profesionales desde descripción',
                'Crear emblemas vectoriales, logotipos o insignias limpias para asociaciones, cursos o marcas desde una instrucción de texto E E.g.',
                'Diseña un logotipo vectorial minimalista y elegante sobre fondo blanco para un club de lectura y tecnología: combina un libro abierto que se transforma en un búho geométrico azul y oro.')

add_taller_item('4.4.8.', 'Variaciones automáticas de un Logotipo',
                'A partir de un logo existente, crear múltiples adaptaciones prácticas: en blanco y negro, en sello circular, en metálico dorado o bordado.',
                'A partir de este logotipo de referencia, crea 4 variaciones profesionales: una versión monocromática negra, una en oro metálico sobre fondo marino, un sello circular y un bordado.')

add_taller_item('4.4.9.', 'Diseño y Aplicación de Etiquetas de Producto',
                'Crear una etiqueta comercial (para una botella de vino, aceite de oliva o cerveza artesanal) y aplicarla en una foto de producto real.',
                'Diseña una etiqueta elegante estilo vintage en oro y verde para una botella de aceite de oliva virgen extra y pégala de forma hiperrealista sobre una botella de cristal oscuro en mesa rústica.')

add_taller_item('4.4.10.', 'Banners Publicitarios Completos',
                'Combinar en un solo cartel publicitario la foto de un producto, el logotipo de la marca y un fondo atractivo en perfecta armonía.',
                'Crea un banner publicitario elegante combinando la foto del producto, el logotipo en la esquina superior y un fondo cálido desenfocado con una iluminación comercial atractiva.')

add_taller_item('4.4.11.', 'Rediseño Cultural de Marca (Locales Internacionales)',
                'Adaptar la estética de un anuncio o cartel occidental al estilo visual de otra cultura (por ejemplo, al minimalismo zen japonés).',
                'Rediseña este cartel publicitario occidental adaptando toda su estética, colores y composición al estilo minimalista y zen japonés, logrando un ambiente sereno, limpio y sofisticado.')

output_path = '/Users/externo/Library/Mobile Documents/com~apple~CloudDocs/PERSONAL/CLASES DE TECNOLOGÍA/CURSO-IA/ESTILOS DE IMAGENES CON IA V2.docx'

    # ==========================================
    # ==========================================
    # ==========================================
# ==========================================
# ==========================================
# ==========================================
# ==========================================
# BLOQUE V: LÓGICA, CREATIVIDAD LITERARIA, EDICIÓN COMERCIAL Y CASOS PRÁCTICOS COTIDIANOS
# ==========================================
add_header_1("V. LÓGICA, CREATIVIDAD LITERARIA, EDICIÓN COMERCIAL Y CASOS PRÁCTICOS COTIDIANOS CON IA")

p_intro_v = doc.add_paragraph()
p_intro_v.add_run("Este quinto bloque amplía el repositorio maestro reuniendo técnicas esenciales de razonamiento lógico, creatividad textual, gestión documental ágil y casos prácticos cotidianos. Todas las instrucciones se han adaptado al método pedagógico del curso para adultos y mayores (60+), priorizando la sencillez, las aplicaciones reales y la optimización de instrucciones (< 200 caracteres).")

# 5.1
add_header_2("5.1. Razonamiento Lógico y Progresión Geométrica (El Ajedrez de Sissa)")
add_taller_item("5.1.1.", "El Cálculo de Sissa y la Analogía del Trigo en el Tablero", 
                "Demostrar cómo la IA calcula progresiones exponenciales (hasta la casilla 64 = 18.446.744.073.709.551.615 granos) y las traduce en analogías comprensibles: 922 billones de toneladas de trigo que superarían toda la superficie cultivable de la Tierra.", 
                "Actúa como profesor de matemáticas. Explícame el cálculo de los granos de trigo en el tablero de ajedrez y dame una analogía sencilla con el peso mundial.")

# 5.2
add_header_2("5.2. Creatividad Literaria y Juego de Roles (Cervantes vs. Shakespeare)")
add_taller_item("5.2.1.", "Duelo Poético en Verso y Brindis Histórico Conjunto", 
                "Adoptar personalidades literarias e históricas para que Cervantes y Shakespeare entablen un duelo en verso, recriminándose con humor sus estilos poéticos para terminar con un brindis de amistad: ¡Brindemos, amigo, que juntos brillamos!.", 
                "Escribe un diálogo corto en verso siglo XVII donde Cervantes y Shakespeare debatan con humor sobre sus obras y terminen brindando como amigos.")

# 5.3
add_header_2("5.3. Análisis Lingüístico y Figuras Retóricas en Literatura")
add_taller_item("5.3.1.", "Definición y Ejemplificación Poética de las 7 Figuras Clave", 
                "Sintetizar definiciones claras y generar ejemplos poéticos memorables para las 7 figuras retóricas fundamentales del curso: anáfora, antítesis, anadiplosis, tricolon, quiasmo, epífora y metáfora.", 
                "Defíneme con claridad y ponme un ejemplo poético breve de estas 7 figuras retóricas: anáfora, antítesis, anadiplosis, tricolon, quiasmo, epífora y metáfora.")

# 5.4
add_header_2("5.4. Flujo Rápido de Trabajo con Documentos PDF (y la Revolución NotebookLM)")
add_taller_item("5.4.1.", "El Flujo de 4 Pasos: Resumen, Artículo, Cambio de Tono y Diapositivas", 
                "Ejecutar las 4 fases de procesamiento de documentos: 1. Resumen -> 2. Artículo divulgativo para principiantes -> 3. Cambio de tono profesional -> 4. Estructura por diapositivas. NOTA PEDAGÓGICA CLAVE (60+): Hoy en día la herramienta líder y más sencilla para documentos largos, apuntes de clase y libros es NotebookLM de Google, permitiendo chatear con hasta 50 documentos y crear podcasts explicativos de audio con un solo clic.", 
                "Hazme un resumen en 3 puntos clave de este documento y luego escribe un artículo de 150 palabras con tono divulgativo y sencillo para principiantes.")

# 5.5
add_header_2("5.5. Visión IA en el Móvil: La Cámara en la Cocina (¿Qué cocino hoy?)")
add_taller_item("5.5.1.", "Reconocimiento Visual de Ingredientes en Nevera y Creación de Recetas", 
                "Abrir la app de Gemini o ChatGPT en el móvil, pulsar el icono de la Cámara, fotografiar ingredientes sueltos de la cocina o nevera (rúcula, zanahorias, patatas, lechuga y cebolla) y solicitar recetas inmediatas y saludables en tiempo real.", 
                "Mira la foto adjunta de los ingredientes que tengo en mi cocina. Sugíreme 2 recetas sencillas, saludables y rápidas que puedo preparar hoy mismo con ellos.")

# 5.6
add_header_2("5.6. Edición Creativa y Merchandising con Gemini (El Reto de la Chica y Coca-Cola)")
add_taller_item("5.6.1.", "Product Placement en Mano y Diseño de Logotipo en Camiseta", 
                "Realizar transformaciones comerciales y publicitarias sobre fotografías reales de personas usando Gemini Gratuito (Clase 8): 1) Colocar de forma natural una lata o botella de Coca-Cola en su mano igualando iluminación y sombras; 2) Modificar el diseño de su camiseta estampando el logotipo y tipografía oficial de Coca-Cola en el pecho.", 
                "Modifica la foto de la chica adjunta: ponle una lata de Coca-Cola en la mano con reflejos reales y añade el logotipo oficial de Coca-Cola en el centro de su camiseta.")

doc.save(output_path)
print("REPOSITORIO MAESTRO V2 COMPLETADO CON BLOQUES I, II, III, IV y V.")
doc = docx.Document(doc_path)
import re
doc = docx.Document(doc_path)
import re
root_xml = ET.Element('map', {'version': 'freeplane 1.11.1'})

id_counter = 1
def get_id():
    global id_counter
    res = f'ID_{id_counter}'
    id_counter += 1
    return res

root_node = ET.SubElement(root_xml, 'node', {
    'ID': get_id(),
    'TEXT': 'REPOSITORIO MAESTRO V2: ÍNDICE GENERAL (MAYORES 60+)'
})
ET.SubElement(root_node, 'font', {'NAME': 'Calibri', 'SIZE': '14', 'BOLD': 'true'})

def get_level(p):
    st = p.style.name.lower() if p.style else ''
    txt = p.text.strip()
    if not txt: return 0
    if 'heading 1' in st or 'título 1' in st: return 1
    if 'heading 2' in st or 'título 2' in st: return 2
    if 'heading 3' in st or 'título 3' in st: return 3
    if re.match(r'^(I|II|III|IV|V)\.\s+', txt) or re.match(r'^BLOQUE\s+[0-9IV]+', txt, re.IGNORECASE): return 1
    if re.match(r'^[1-9]\.\s+[A-ZÁÉÍÓÚÑ]', txt) or re.match(r'^[1-9]\.\s+Módulo\s+[A-Z]', txt) or re.match(r'^5\.[1-6]\.\s+[A-ZÁÉÍÓÚÑ]', txt):
        if not re.match(r'^[1-9]\.[0-9]+\.[0-9]+', txt) and not re.match(r'^[1-4]\.[1-9]\s+[A-ZÁÉÍÓÚÑ]', txt): return 2
    if re.match(r'^[1-4]\.[1-9]\.\s+[A-ZÁÉÍÓÚÑ]', txt) and not re.match(r'^[1-4]\.[1-9]\.[0-9]+', txt): return 2
    if re.match(r'^[0-9]+\.[0-9]+\.[0-9]+\.?\s+', txt): return 3
    return 0

current_h1 = root_node
current_h2 = root_node
current_h3 = root_node

for p in doc.paragraphs:
    txt = p.text.strip()
    if not txt: continue
    lvl = get_level(p)
    if lvl == 1:
        current_h1 = ET.SubElement(root_node, 'node', {'ID': get_id(), 'TEXT': txt, 'POSITION': 'right'})
        ET.SubElement(current_h1, 'font', {'NAME': 'Calibri', 'SIZE': '13', 'BOLD': 'true'})
        ET.SubElement(current_h1, 'edge', {'COLOR': '#003366', 'WIDTH': '2'})
        current_h2 = current_h1
        current_h3 = current_h1
    elif lvl == 2:
        current_h2 = ET.SubElement(current_h1, 'node', {'ID': get_id(), 'TEXT': txt})
        ET.SubElement(current_h2, 'font', {'NAME': 'Calibri', 'SIZE': '12', 'BOLD': 'true'})
        current_h3 = current_h2
    elif lvl == 3:
        current_h3 = ET.SubElement(current_h2, 'node', {'ID': get_id(), 'TEXT': txt})
        ET.SubElement(current_h3, 'font', {'NAME': 'Calibri', 'SIZE': '11', 'BOLD': 'true'})

def indent(elem, level=0):
    i = '
' + level * '  '
    if len(elem):
        if not elem.text or not elem.text.strip(): elem.text = i + '  '
        if not elem.tail or not elem.tail.strip(): elem.tail = i
        for s in elem: indent(s, level + 1)
        if not elem.tail or not elem.tail.strip(): elem.tail = i
    else:
        if level and (not elem.tail or not elem.tail.strip()): elem.tail = i

indent(root_xml)
xml_str = '<?xml version="1.0" encoding="UTF-8"?>
' + ET.tostring(root_xml, encoding='utf-8').decode('utf-8')

with open(mm_path, 'w', encoding='utf-8') as f:
    f.write(xml_str)
