# -*- coding: utf-8 -*-
"""
Crea la carpeta CLASES/2. INTRODUCCION_NLM
y genera dentro PROMPT_PRESENTACION_NLM.docx
con el prompt maestro para NotebookLM adaptado a mayores de 60 años.
"""

import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

ROOT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TARGET_DIR = os.path.join(ROOT_DIR, "CLASES", "2. INTRODUCCION_NLM")
TARGET_FILE = os.path.join(TARGET_DIR, "PROMPT_PRESENTACION_NLM.docx")

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=200, right=200):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def create_nlm_prompt_docx():
    os.makedirs(TARGET_DIR, exist_ok=True)
    if os.path.exists(TARGET_FILE):
        try: os.remove(TARGET_FILE)
        except Exception: pass

    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("CURSO DE INTELIGENCIA ARTIFICIAL PARA PERSONAS MAYORES (60+)")
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(11)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("PROMPT MAESTRO PARA NOTEBOOKLM (NLM)\nTALLER 3: GUÍA VISUAL PASO A PASO DE NOTEBOOKLM")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(15)
    r_sub.font.bold = True
    r_sub.font.color.rgb = RGBColor(0x1A, 0x0A, 0x2E)

    p_inst = doc.add_paragraph()
    p_inst.paragraph_format.space_before = Pt(8)
    p_inst.paragraph_format.space_after = Pt(12)
    r_inst = p_inst.add_run("📋 INSTRUCCIÓN: Copia todo el contenido del recuadro inferior y pégalo en el chat de tu cuaderno de NotebookLM para generar la presentación visual (diapositiva a diapositiva) que explicará NotebookLM a los alumnos de forma clara, amena y sin tecnicismos.")
    r_inst.font.name = 'Calibri'
    r_inst.font.size = Pt(10)
    r_inst.font.italic = True
    r_inst.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    prompt_text = """Actúa como un Diseñador Curricular y Pedagógico Senior especializado en la enseñanza de Inteligencia Artificial para personas mayores (65 a 75 años sin experiencia informática).

Tu misión es generar el guion y contenido EXACTO diapositiva por diapositiva para crear una PRESENTACIÓN VISUAL titulada:
«NOTEBOOKLM: TU ASISTENTE DE DOCUMENTOS Y CUADERNO INTELIGENTE»

El objetivo es que los alumnos entiendan qué es NotebookLM, en qué se diferencia de Gemini, cómo subir sus propios documentos familiares, médicos o recetas, y cómo escuchar el programa de radio/audio automático que genera.

REGLAS PEDAGÓGICAS ESTRICTAS:
1. Lenguaje cálido, cercano, sin jerga informática y con letras grandes.
2. Cada diapositiva debe contener una sola idea principal, fácil de asimilar.
3. Explicar la diferencia clave con Gemini: Gemini sabe de todo internet; NotebookLM solo lee los papeles que tú le subes a su mesa (cero invenciones).
4. Destacar la función de audio como algo mágico y muy fácil de usar.

ESTRUCTURA EXACTA DE LAS 12 DIAPOSITIVAS:

SLIDE 1: PORTADA
- Título: NotebookLM: Tu Cuaderno Inteligente y Archivo Personal.
- Subtítulo: Pon a trabajar a la IA con tus propios documentos, recuerdos y recetas.
- Mensaje de bienvenida: "Una biblioteca inteligente que solo lee lo que tú le pides."

SLIDE 2: ¿EN QUÉ SE DIFERENCIA DE GEMINI?
- Metáfora visual: Gemini es como una enciclopedia abierta a todo el mundo; NotebookLM es tu carpeta personal de casa.
- La gran ventaja: NotebookLM solo responde con la verdad de tus papeles. No se inventa cosas de internet.
- Seguridad y rigor: Cada dato que te da tiene una pequeña etiqueta numerada para que veas de qué página lo sacó.

SLIDE 3: CÓMO ENTRAR A NOTEBOOKLM
- Dirección web: notebooklm.google.com en el navegador (Chrome o Edge).
- Inicio de sesión: Con tu misma cuenta de Google (Gmail).
- Sin registros raros: Entras con un solo clic porque ya estás identificado en Google.

SLIDE 4: LA PANTALLA PRINCIPAL: «TUS CUADERNOS»
- Elemento visual: La pantalla limpia con el botón "+ Nuevo cuaderno" en el centro.
- Concepto: Un "cuaderno" es como una carpeta de cartón donde guardas papeles de un mismo tema (ej: "Mis recetas familiares", "Informes del médico" o "La historia de mi pueblo").

SLIDE 5: EL PANEL DE FUENTES: «AÑADIR TUS PAPELES»
- Elemento visual: El panel lateral izquierdo con el botón "+ Añadir fuentes".
- Qué puedes meterle: Archivos PDF, fotos de cartas antiguas, documentos de Word o enlaces web.
- El gesto sencillo: Arrastrar el documento con el ratón directamente a la pantalla o pulsar "Subir desde el ordenador".

SLIDE 6: EL CHAT INTELIGENTE: «PREGUNTA A TUS PAPELES»
- Elemento visual: La barra inferior de chat con el texto de consulta.
- Cómo funciona: Le preguntas en español normal, como si hablaras con un archivador que se ha leído todos tus papeles en 3 segundos.
- Ejemplo: "Dime qué ingredientes lleva la receta de rosquillas de mi abuela" o "¿En qué fecha nos mudamos de piso según las cartas?".

SLIDE 7: LAS CITAS NUMERADAS: «DEMOSTRACIÓN DE VERDAD»
- Elemento visual: Un párrafo de respuesta con pequeños números [1], [2] al final de las frases.
- Explicación: Al pulsar sobre el número [1], la pantalla salta directamente al párrafo exacto de tu documento donde está escrito.
- Tranquilidad: Sabes al 100% que no se lo está inventando.

SLIDE 8: LA GUÍA DEL CUADERNO: ESQUEMAS Y RESÚMENES
- Elemento visual: El panel superior con botones como "Guía del cuaderno", "Preguntas frecuentes" y "Línea de tiempo".
- Utilidad: Te resume automáticamente un documento largo de 20 páginas en 5 puntos clave para que no tengas que leerlo todo de golpe.

SLIDE 9: LA FUNCIÓN ESTRELLA: EL PROGRAMA DE RADIO (AUDIO OVERVIEW) 🎙️
- Elemento visual: El botón "Generar" bajo el apartado de audio.
- La magia: La IA crea una tertulia de radio en español entre dos locutores simpáticos que comentan y explican tus documentos como si fuera una conversación de sobremesa.
- Cómo disfrutarlo: Le das al "Play" y lo escuchas tranquilamente mientras tomas un café o descansas.

SLIDE 10: CASO PRÁCTICO 1: TUS MEMORIAS Y RECETAS FAMILIARES
- Qué subes: Fotos de cuadernos viejos de cocina, recuerdos de juventud o anécdotas escritas en un folio.
- Qué te da: Un recetario ordenado por platos, o un relato continuo de tu vida para regalar a tus nietos.

SLIDE 11: CASO PRÁCTICO 2: INFORMES MÉDICOS Y PAPELES DE CASA
- Qué subes: El PDF de un análisis de sangre, el contrato de la luz o la carta del banco.
- Qué le pides: "Explícame en palabras sencillas qué significa este informe para que lo entienda sin ser médico ni abogado".
- Privacidad total: Tus documentos son solo tuyos y privados; nadie más puede verlos.

SLIDE 12: LAS 3 REGLAS DE ORO DE NOTEBOOKLM
- 1ª Regla: Cuanto mejor sea el documento que subas, mejor será la ayuda que te preste.
- 2ª Regla: Comprueba siempre las citas [1] para ver el papel original.
- 3ª Regla: Disfruta escuchando tus audios como si tuvieras tu propia emisora de radio en casa."""

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    cell = table.cell(0, 0)
    cell.width = Inches(6.8)
    set_cell_background(cell, "F8FAFC")
    set_cell_margins(cell, top=200, bottom=200, left=240, right=240)

    p_box = cell.paragraphs[0]
    p_box.paragraph_format.line_spacing = 1.15
    p_box.paragraph_format.space_after = Pt(4)

    for line in prompt_text.split('\n'):
        if not line.strip():
            p_line = cell.add_paragraph()
            p_line.paragraph_format.space_before = Pt(2)
            p_line.paragraph_format.space_after = Pt(2)
            continue
        p_line = cell.add_paragraph()
        p_line.paragraph_format.line_spacing = 1.15
        p_line.paragraph_format.space_after = Pt(2)
        r = p_line.add_run(line)
        r.font.name = 'Consolas'
        r.font.size = Pt(9)
        if line.startswith(('SLIDE', 'REGLAS', 'ESTRUCTURA')):
            r.font.bold = True
            r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        elif line.strip().startswith(('-', '•')):
            r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        else:
            r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    if len(cell.paragraphs) > 1 and not cell.paragraphs[0].text.strip():
        p_elem = cell.paragraphs[0]._p
        p_elem.getparent().remove(p_elem)

    doc.save(TARGET_FILE)
    print(f"✅ Archivo creado con éxito: {TARGET_FILE}")

if __name__ == '__main__':
    create_nlm_prompt_docx()
