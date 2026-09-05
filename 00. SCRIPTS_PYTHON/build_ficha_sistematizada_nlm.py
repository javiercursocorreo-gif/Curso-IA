# -*- coding: utf-8 -*-
"""
Generador de la Ficha Práctica Sistematizada de NotebookLM en formato Word (.docx) y PDF.
Organiza de forma pedagógica, ordenada y exhaustiva todas las funciones de NotebookLM
a partir de las 3 fuentes de ejemplo (TXT, PDF, YouTube):
1. Carga y observación del resumen automático central.
2. Resúmenes individuales y preguntas concretas con citas [1] en cada fuente.
3. Fijar resultados como notas (Pin).
4. El panel Studio: seleccionar solo YouTube y probar los botones de Studio.
5. Seleccionar TXT + PDF: pulsar botón Presentación con estilo visual sobrio y botón Revisar sobre la marcha.
6. Seleccionar las 3 fuentes: pulsar directamente el botón «Resumen de vídeo» en Studio sin prompts adicionales.
"""

import os
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib.units import cm

ROOT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TARGET_DIR = os.path.join(ROOT_DIR, "CLASES", "2. INTRODUCCION_NLM")
TARGET_DOCX = os.path.join(TARGET_DIR, "FICHA_PRACTICA_SISTEMATIZADA_NOTEBOOKLM.docx")
TARGET_PDF = os.path.join(TARGET_DIR, "FICHA_PRACTICA_SISTEMATIZADA_NOTEBOOKLM.pdf")

os.makedirs(TARGET_DIR, exist_ok=True)

# ---------------------------------------------------------------------------
# UTILIDADES PARA ESTILO DOCX
# ---------------------------------------------------------------------------
def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_callout_box(doc, title_text, body_paragraphs, border_color="0B4F6C", bg_color="F0F7FA"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)

    cell = table.cell(0, 0)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=160, bottom=160, left=220, right=220)

    # Borde izquierdo grueso
    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="none"/>'
        f'  <w:left w:val="single" w:sz="30" w:space="0" w:color="{border_color}"/>'
        f'  <w:bottom w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

    p_title = cell.paragraphs[0]
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    run_t = p_title.add_run(title_text)
    run_t.bold = True
    run_t.font.name = "Calibri"
    run_t.font.size = Pt(11)
    run_t.font.color.rgb = RGBColor.from_string(border_color)

    for body_text in body_paragraphs:
        p_body = cell.add_paragraph()
        p_body.paragraph_format.space_before = Pt(0)
        p_body.paragraph_format.space_after = Pt(3)
        run_b = p_body.add_run(body_text)
        run_b.font.name = "Calibri"
        run_b.font.size = Pt(10)
        run_b.font.color.rgb = RGBColor(40, 50, 60)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ---------------------------------------------------------------------------
# GENERADOR DEL DOCUMENTO WORD (.DOCX)
# ---------------------------------------------------------------------------
def generate_docx():
    print("📝 Generando Ficha Práctica Sistematizada en DOCX...")
    doc = docx.Document()

    # Márgenes de página
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.85)
        s.right_margin = Inches(0.85)

    C_NAVY = RGBColor(11, 37, 69)      # #0B2545
    C_BLUE = RGBColor(0, 102, 161)     # #0066A1
    C_DARK = RGBColor(33, 37, 41)
    C_MUTED = RGBColor(100, 116, 139)

    # --- ENCABEZADO Y TÍTULO ---
    p_inst = doc.add_paragraph()
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_inst.paragraph_format.space_after = Pt(2)
    r_inst = p_inst.add_run("CURSO DE INTELIGENCIA ARTIFICIAL Y TECNOLOGÍA PARA ADULTOS MAYORES (60+)")
    r_inst.font.name = "Calibri"
    r_inst.font.size = Pt(9.5)
    r_inst.font.bold = True
    r_inst.font.color.rgb = C_BLUE

    p_main = doc.add_paragraph()
    p_main.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_main.paragraph_format.space_after = Pt(4)
    r_main = p_main.add_run("TALLER MONOGRÁFICO DE NOTEBOOKLM")
    r_main.font.name = "Calibri"
    r_main.font.size = Pt(20)
    r_main.font.bold = True
    r_main.font.color.rgb = C_NAVY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(14)
    r_sub = p_sub.add_run("Ficha de Práctica Sistemática: Todas las Funciones de la Pantalla Real Paso a Paso")
    r_sub.font.name = "Calibri"
    r_sub.font.size = Pt(12)
    r_sub.font.italic = True
    r_sub.font.color.rgb = C_MUTED

    # Cuadro Objetivo
    add_callout_box(
        doc,
        "🎯 Objetivo del Taller Práctico",
        [
            "Aprender a manejar Google NotebookLM de forma rigurosa, visual y ordenada.",
            "Partiremos de 3 fuentes complementarias sobre La Llegada a la Luna (20 de julio de 1969): un recuerdo familiar (TXT), un informe histórico de la NASA (PDF) y un vídeo conmemorativo (YouTube).",
            "Descubriremos el resumen automático central, cómo hacer preguntas con citas [1] en cada papel, cómo fijar notas permanentes, cómo exprimir los botones de Studio, cómo crear una presentación con estilo y revisarla en vivo, y cómo generar un resumen de vídeo combinando las 3 fuentes."
        ],
        border_color="0066A1",
        bg_color="F0F7FC"
    )

    # --- FASE 0: PREPARACIÓN Y FUENTES ---
    p_h1 = doc.add_paragraph()
    p_h1.paragraph_format.space_before = Pt(12)
    p_h1.paragraph_format.space_after = Pt(6)
    r = p_h1.add_run("FASE 0: Las 3 Fuentes de Partida (Panel Izquierdo)")
    r.font.name = "Calibri"
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = C_NAVY

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run("En la columna izquierda («Fuentes»), pulsa ").font.color.rgb = C_DARK
    p.add_run("+ Añadir fuentes").bold = True
    p.add_run(" y carga estos 3 documentos que tienes preparados en tu carpeta:")

    # Tabla de las 3 Fuentes
    t_fuentes = doc.add_table(rows=4, cols=3)
    t_fuentes.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_fuentes.autofit = False
    col_widths = [Inches(1.2), Inches(2.3), Inches(3.0)]

    headers = ["Tipo", "Nombre del Archivo / Enlace", "¿Qué contiene? (Contenido clave)"]
    for j, h in enumerate(headers):
        c = t_fuentes.cell(0, j)
        c.width = col_widths[j]
        set_cell_background(c, "0B2545")
        set_cell_margins(c, top=120, bottom=120, left=140, right=140)
        p_h = c.paragraphs[0]
        r = p_h.add_run(h)
        r.bold = True
        r.font.name = "Calibri"
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(255, 255, 255)

    f_data = [
        ("📄 FUENTE 1 (Texto TXT)", "EJEMPLO_A_FUENTE_TEXTO_RECUERDO_LUNA_1969.txt", "El recuerdo personal: La noche calurosa en Madrid, la tele Telefunken en blanco y negro, la voz de Jesús Hermida y las lágrimas del abuelo que nació viajando en carro de mulas."),
        ("📑 FUENTE 2 (Documento PDF)", "EJEMPLO_A_FUENTE_HISTORIA_APOLO_11.pdf", "Los datos técnicos y científicos: 400.000 ingenieros, el Saturno V, el ordenador AGC con solo 4 KB de RAM y el papel crucial de la estación española de Fresnedillas de la Oliva (Madrid)."),
        ("🎥 FUENTE 3 (Vídeo YouTube)", "https://www.youtube.com/watch?v=Wbh9VAuZcUM (Euronews en español)", "La conmemoración oficial: 50 años del lanzamiento del Apolo 11, declaraciones de historiadores y la mirada hacia las futuras misiones lunares.")
    ]

    for i, (t_tipo, t_nom, t_desc) in enumerate(f_data, start=1):
        bg = "FFFFFF" if i % 2 != 0 else "F8FAFC"
        row_vals = [t_tipo, t_nom, t_desc]
        for j, val in enumerate(row_vals):
            c = t_fuentes.cell(i, j)
            c.width = col_widths[j]
            set_cell_background(c, bg)
            set_cell_margins(c, top=100, bottom=100, left=120, right=120)
            p_cell = c.paragraphs[0]
            p_cell.paragraph_format.space_after = Pt(0)
            r = p_cell.add_run(val)
            r.font.name = "Calibri"
            r.font.size = Pt(9)
            if j == 0:
                r.bold = True
            r.font.color.rgb = C_DARK

    doc.add_paragraph().paragraph_format.space_after = Pt(6)

    add_callout_box(
        doc,
        "💡 LA REGLA DE ORO DE LAS CASILLAS DE VERIFICACIÓN (Checkboxes)",
        [
            "Cada fuente tiene una pequeña casilla cuadrada a su izquierda [✔].",
            "NotebookLM es un archivador inteligente que SOLO lee las fuentes que tengan la casilla marcada.",
            "Si dejas activada una sola fuente, la IA 'olvidará' temporalmente las otras dos. ¡Esta es la clave para hacer análisis individuales y limpios!"
        ],
        border_color="0B7285",
        bg_color="E6FCF5"
    )

    # --- FASE 1: LA ZONA CENTRAL ---
    p_h1 = doc.add_paragraph()
    p_h1.paragraph_format.space_before = Pt(12)
    p_h1.paragraph_format.space_after = Pt(6)
    r = p_h1.add_run("FASE 1: La Zona Central (Observación, Resúmenes y Preguntas con Citas)")
    r.font.name = "Calibri"
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = C_NAVY

    # Paso 1.1
    p_h2 = doc.add_paragraph()
    p_h2.paragraph_format.space_before = Pt(6)
    p_h2.paragraph_format.space_after = Pt(3)
    r = p_h2.add_run("1.1. Observación guiada: El Resumen Automático Central")
    r.font.name = "Calibri"
    r.font.size = Pt(11.5)
    r.font.bold = True
    r.font.color.rgb = C_BLUE

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run("[  ] ").bold = True
    p.add_run("Con las 3 fuentes marcadas [✔], dirige tu mirada al centro de la pantalla. Verás la tarjeta titulada ").font.color.rgb = C_DARK
    p.add_run("«Guía del cuaderno»").bold = True
    p.add_run(". Lee atentamente:")

    add_callout_box(
        doc,
        "👀 QUÉ DEBE OBSERVAR EL ALUMNO EN EL CENTRO:",
        [
            "1. NotebookLM ha elaborado por sí mismo un resumen general sin que hayamos tecleado nada.",
            "2. Ha identificado que los tres documentos tratan sobre el alunizaje de 1969.",
            "3. En la parte inferior sugiere varias preguntas recomendadas para explorar los temas."
        ],
        border_color="D97706",
        bg_color="FFFBEB"
    )

    # Paso 1.2
    p_h2 = doc.add_paragraph()
    p_h2.paragraph_format.space_before = Pt(8)
    p_h2.paragraph_format.space_after = Pt(3)
    r = p_h2.add_run("1.2. Práctica de aislamiento: Resumen individual de cada una de las 3 fuentes")
    r.font.name = "Calibri"
    r.font.size = Pt(11.5)
    r.font.bold = True
    r.font.color.rgb = C_BLUE

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run("Ahora vamos a pedirle a la IA que examine cada fuente por separado:")

    p_a = doc.add_paragraph()
    p_a.paragraph_format.left_indent = Inches(0.2)
    p_a.paragraph_format.space_after = Pt(3)
    p_a.add_run("[  ] Ejercicio A (Solo TXT): ").bold = True
    p_a.add_run("Desmarca el PDF y el vídeo. Deja marcado solo el TXT. En el recuadro inferior de chat escribe:\n").font.color.rgb = C_DARK
    r_p = p_a.add_run("👉 «Resume en 3 líneas los recuerdos familiares y la emoción de aquella noche.»")
    r_p.bold = True
    r_p.font.color.rgb = C_NAVY

    p_b = doc.add_paragraph()
    p_b.paragraph_format.left_indent = Inches(0.2)
    p_b.paragraph_format.space_after = Pt(3)
    p_b.add_run("[  ] Ejercicio B (Solo PDF): ").bold = True
    p_b.add_run("Desmarca el TXT. Deja marcado solo el PDF. Escribe en el chat:\n").font.color.rgb = C_DARK
    r_p = p_b.add_run("👉 «Resume en 3 puntos técnicos los datos más asombrosos del programa Apolo 11.»")
    r_p.bold = True
    r_p.font.color.rgb = C_NAVY

    p_c = doc.add_paragraph()
    p_c.paragraph_format.left_indent = Inches(0.2)
    p_c.paragraph_format.space_after = Pt(6)
    p_c.add_run("[  ] Ejercicio C (Solo YouTube): ").bold = True
    p_c.add_run("Desmarca el PDF. Deja marcado solo el vídeo de YouTube. Escribe en el chat:\n").font.color.rgb = C_DARK
    r_p = p_c.add_run("👉 «Resume en 3 frases qué conmemora el vídeo de Euronews y qué destaca de la misión.»")
    r_p.bold = True
    r_p.font.color.rgb = C_NAVY

    # Paso 1.3
    p_h2 = doc.add_paragraph()
    p_h2.paragraph_format.space_before = Pt(8)
    p_h2.paragraph_format.space_after = Pt(3)
    r = p_h2.add_run("1.3. Búsqueda quirúrgica y verificación de citas numeradas [1]")
    r.font.name = "Calibri"
    r.font.size = Pt(11.5)
    r.font.bold = True
    r.font.color.rgb = C_BLUE

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run("Marca ahora las 3 fuentes a la vez [✔]. Vamos a hacer preguntas concretas para comprobar ").font.color.rgb = C_DARK
    p.add_run("la cita numerada [1] que demuestra de dónde sale cada dato:").bold = True

    p_q1 = doc.add_paragraph()
    p_q1.paragraph_format.left_indent = Inches(0.2)
    p_q1.paragraph_format.space_after = Pt(3)
    p_q1.add_run("[  ] Pregunta 1 (Dato personal): ").bold = True
    p_q1.add_run("«¿Qué marca era la televisión familiar y qué comparación hizo el abuelo entre su infancia y la llegada a la Luna?»\n").font.color.rgb = C_DARK
    p_q1.add_run("   🔍 Comprobación: Pulsa sobre el número [1] al final de la respuesta. Verás cómo la pantalla se mueve y subraya la palabra 'Telefunken' y la frase sobre el 'carro de mulas' en tu texto original.").font.italic = True

    p_q2 = doc.add_paragraph()
    p_q2.paragraph_format.left_indent = Inches(0.2)
    p_q2.paragraph_format.space_after = Pt(3)
    p_q2.add_run("[  ] Pregunta 2 (Dato técnico): ").bold = True
    p_q2.add_run("«¿Cuánta memoria RAM tenía el ordenador del Apolo 11 y qué estación española en Madrid escuchó a Neil Armstrong antes que Houston?»\n").font.color.rgb = C_DARK
    p_q2.add_run("   🔍 Comprobación: Haz clic en la cita [2]. Comprobarás que subraya '4 kilobytes de memoria RAM' y 'Fresnedillas de la Oliva' en el PDF.").font.italic = True

    p_q3 = doc.add_paragraph()
    p_q3.paragraph_format.left_indent = Inches(0.2)
    p_q3.paragraph_format.space_after = Pt(6)
    p_q3.add_run("[  ] Pregunta 3 (Dato audiovisual): ").bold = True
    p_q3.add_run("«¿Cuántos años celebra el reportaje de televisión y qué mensaje transmite?»\n").font.color.rgb = C_DARK
    p_q3.add_run("   🔍 Comprobación: Haz clic en la cita. Te remitirá al minuto exacto del vídeo de Euronews transcrito.").font.italic = True

    # Paso 1.4
    p_h2 = doc.add_paragraph()
    p_h2.paragraph_format.space_before = Pt(8)
    p_h2.paragraph_format.space_after = Pt(3)
    r = p_h2.add_run("1.4. La función «Fijar como Nota» (Guardar en el Corcho de Studio)")
    r.font.name = "Calibri"
    r.font.size = Pt(11.5)
    r.font.bold = True
    r.font.color.rgb = C_BLUE

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run("[  ] ").bold = True
    p.add_run("Escribe en el chat: ").font.color.rgb = C_DARK
    r_p = p.add_run("«Compara en una tabla la vivencia humana del salón madrileño con la frialdad técnica de los 4 KB de memoria.»")
    r_p.bold = True
    r_p.font.color.rgb = C_NAVY

    add_callout_box(
        doc,
        "📌 PASO OBLIGATORIO: FIJAR COMO NOTA (Pin)",
        [
            "1. En cuanto NotebookLM te responda con la tabla comparativa, sitúa el ratón encima de la respuesta.",
            "2. En la esquina superior derecha del mensaje verás el icono de una chincheta o botón «Guardar en nota» (Pin). ¡Haz clic en él!",
            "3. Observa el panel lateral derecho (Studio): la respuesta se ha convertido en una Nota permanente. Nunca se borrará aunque cierres el chat o apagues el ordenador."
        ],
        border_color="C2410C",
        bg_color="FFF7ED"
    )

    # --- FASE 2: PANEL DE STUDIO (9 BOTONES) ---
    p_h1 = doc.add_paragraph()
    p_h1.paragraph_format.space_before = Pt(12)
    p_h1.paragraph_format.space_after = Pt(6)
    r = p_h1.add_run("FASE 2: El Panel de Studio — Explorando los 9 Botones Automáticos (Solo con YouTube)")
    r.font.name = "Calibri"
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = C_NAVY

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(6)
    p.add_run("Ahora vamos a explorar la fábrica de contenidos de Studio. Sigue estos pasos:").font.color.rgb = C_DARK

    p_step = doc.add_paragraph()
    p_step.paragraph_format.left_indent = Inches(0.2)
    p_step.paragraph_format.space_after = Pt(4)
    p_step.add_run("1º En la columna izquierda: ").bold = True
    p_step.add_run("Desmarca el TXT [ ] y desmarca el PDF [ ]. Deja marcado ").font.color.rgb = C_DARK
    p_step.add_run("ÚNICAMENTE el vídeo de YouTube [✔]").bold = True
    p_step.add_run(". Así todo lo generado se basará 100% en el vídeo.")

    p_step2 = doc.add_paragraph()
    p_step2.paragraph_format.left_indent = Inches(0.2)
    p_step2.paragraph_format.space_after = Pt(8)
    p_step2.add_run("2º En la columna derecha (Studio): ").bold = True
    p_step2.add_run("Ve pulsando UNO A UNO cada uno de los 9 botones automáticos y observa el resultado:").font.color.rgb = C_DARK

    # Tabla de los 9 Botones
    t_botones = doc.add_table(rows=10, cols=3)
    t_botones.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_botones.autofit = False
    b_col_widths = [Inches(1.5), Inches(2.2), Inches(2.8)]

    b_headers = ["Botón en Studio", "Acción del Alumno", "¿Qué crea NotebookLM?"]
    for j, h in enumerate(b_headers):
        c = t_botones.cell(0, j)
        c.width = b_col_widths[j]
        set_cell_background(c, "0B2545")
        set_cell_margins(c, top=100, bottom=100, left=120, right=120)
        p_h = c.paragraphs[0]
        r = p_h.add_run(h)
        r.bold = True
        r.font.name = "Calibri"
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(255, 255, 255)

    botones_data = [
        ("1. Preguntas frecuentes (FAQ)", "[  ] Pulsa 'Preguntas frecuentes'", "Un cuestionario con las dudas clave que resolvería alguien que ve el vídeo de la conmemoración."),
        ("2. Guía de estudio 📚", "[  ] Pulsa 'Guía de estudio'", "Un temario pedagógico estructurado con preguntas de repaso y términos clave explicados."),
        ("3. Cronología / Línea de tiempo ⏱️", "[  ] Pulsa 'Línea de tiempo'", "Una lista cronológica con las fechas y años clave mencionados en el reportaje."),
        ("4. Documento informativo 📄", "[  ] Pulsa 'Documento informativo'", "Un informe formal y ejecutivo (Briefing) listo para imprimir o enviar por correo."),
        ("5. Tabla de contenidos 📑", "[  ] Pulsa 'Tabla de contenidos'", "Un índice estructurado con los capítulos y temas tratados a lo largo del vídeo."),
        ("6. Puntos clave / Ideas destacadas 💡", "[  ] Pulsa 'Puntos clave'", "Un decálogo con las perlas de información más importantes sintetizadas en viñetas."),
        ("7. Ensayo / Borrador narrativo ✍️", "[  ] Pulsa 'Borrador / Ensayo'", "Una redacción continua y literaria que narra la trascendencia histórica del acontecimiento."),
        ("8. Resumen breve", "[  ] Pulsa 'Resumen'", "Una síntesis rápida en un solo párrafo para entender la esencia del vídeo."),
        ("9. + Añadir nota manual 📝", "[  ] Pulsa '+ Añadir nota'", "Crea una nota propia en blanco para apuntar tus reflexiones (ej. 'Ver este vídeo con mis nietos').")
    ]

    for i, (b_name, b_act, b_res) in enumerate(botones_data, start=1):
        bg = "FFFFFF" if i % 2 != 0 else "F8FAFC"
        row_vals = [b_name, b_act, b_res]
        for j, val in enumerate(row_vals):
            c = t_botones.cell(i, j)
            c.width = b_col_widths[j]
            set_cell_background(c, bg)
            set_cell_margins(c, top=80, bottom=80, left=100, right=100)
            p_cell = c.paragraphs[0]
            p_cell.paragraph_format.space_after = Pt(0)
            r = p_cell.add_run(val)
            r.font.name = "Calibri"
            r.font.size = Pt(8.5)
            if j == 0:
                r.bold = True
            r.font.color.rgb = C_DARK

    doc.add_paragraph().paragraph_format.space_after = Pt(8)

    # --- FASE 3: EL BOTÓN DE PRESENTACIÓN Y REVISIÓN EN VIVO ---
    p_h1 = doc.add_paragraph()
    p_h1.paragraph_format.space_before = Pt(12)
    p_h1.paragraph_format.space_after = Pt(6)
    r = p_h1.add_run("FASE 3: El Botón de Presentación y Revisión en Vivo (TXT + PDF)")
    r.font.name = "Calibri"
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = C_NAVY

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run("En esta fase no escribiremos en el chat. Utilizaremos directamente el botón de presentación de Studio para generar diapositivas con estilo propio y después revisarlas en vivo:").font.color.rgb = C_DARK

    p_p1 = doc.add_paragraph()
    p_p1.paragraph_format.left_indent = Inches(0.2)
    p_p1.paragraph_format.space_after = Pt(4)
    p_p1.add_run("[  ] Paso 3.1: Selección de Fuentes: ").bold = True
    p_p1.add_run("En la columna izquierda, marca ").font.color.rgb = C_DARK
    p_p1.add_run("el TXT [✔] y el PDF [✔]").bold = True
    p_p1.add_run(". Deja desmarcado el vídeo de YouTube [ ].")

    p_p2 = doc.add_paragraph()
    p_p2.paragraph_format.left_indent = Inches(0.2)
    p_p2.paragraph_format.space_after = Pt(4)
    p_p2.add_run("[  ] Paso 3.2: Pulsar el Botón «Presentación» y fijar el estilo visual: ").bold = True
    p_p2.add_run("En el panel lateral de Studio (a la derecha), pulsa directamente sobre el botón ").font.color.rgb = C_DARK
    p_p2.add_run("«Presentación»").bold = True
    p_p2.add_run(". En la ventana o casilla de personalización que te ofrece NotebookLM para guiar la presentación, escribe este prompt de estilo visual y editorial:\n")

    r_p = p_p2.add_run(
        "👉 «Crea una presentación con un estilo visual sobrio, elegante y de crónica documental histórica. "
        "Utiliza una paleta espacial limpia (azul marino profundo, gris grafito y blanco de alto contraste), "
        "con tipografía grande y muy legible para personas mayores. Estructura las diapositivas equilibrando el rigor técnico "
        "de la NASA con la emoción humana del recuerdo familiar en Madrid, incorporando frases textuales entrecomilladas "
        "y evitando totalmente ilustraciones infantiles, colores estridentes o estilos de fantasía.»"
    )
    r_p.bold = True
    r_p.font.color.rgb = C_NAVY

    p_p3 = doc.add_paragraph()
    p_p3.paragraph_format.left_indent = Inches(0.2)
    p_p3.paragraph_format.space_after = Pt(6)
    p_p3.add_run("[  ] Paso 3.3: Revisar la Presentación resultante con el Botón «Revisar»: ").bold = True
    p_p3.add_run("Una vez generada la presentación en pantalla, entra en ella y pulsa el botón ").font.color.rgb = C_DARK
    p_p3.add_run("«Revisar»").bold = True
    p_p3.add_run(" (o Editar/Ajustar). En este paso no usaremos un texto cerrado: realizaremos la revisión directamente sobre la marcha en clase con el profesor, probando ajustes en vivo según lo que queramos perfeccionar.")

    add_callout_box(
        doc,
        "💡 LA CLAVE DE LA REVISIÓN EN VIVO",
        [
            "La inteligencia artificial ofrece un primer borrador excelente, pero el control pedagógico es siempre tuyo.",
            "Al pulsar el botón «Revisar» dentro de la misma presentación, puedes modificar cualquier diapositiva, añadir más énfasis a una frase del abuelo o pulir los datos técnicos al instante."
        ],
        border_color="0066A1",
        bg_color="F0F7FC"
    )

    # --- FASE 4: EL GRAN RESUMEN DE VÍDEO ---
    p_h1 = doc.add_paragraph()
    p_h1.paragraph_format.space_before = Pt(12)
    p_h1.paragraph_format.space_after = Pt(6)
    r = p_h1.add_run("FASE 4: El Gran Resumen de Vídeo (Las 3 Fuentes Juntas)")
    r.font.name = "Calibri"
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = C_NAVY

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run("El momento cumbre del taller: poner a trabajar la emoción humana (TXT), el rigor histórico (PDF) y el impacto audiovisual (YouTube) de forma totalmente automatizada en Studio.").font.color.rgb = C_DARK

    p_f1 = doc.add_paragraph()
    p_f1.paragraph_format.left_indent = Inches(0.2)
    p_f1.paragraph_format.space_after = Pt(4)
    p_f1.add_run("[  ] Paso 4.1: Selección Total de Fuentes: ").bold = True
    p_f1.add_run("En la columna izquierda, marca las ").font.color.rgb = C_DARK
    p_f1.add_run("3 fuentes a la vez [✔] TXT + [✔] PDF + [✔] YouTube").bold = True
    p_f1.add_run(".")

    p_f2 = doc.add_paragraph()
    p_f2.paragraph_format.left_indent = Inches(0.2)
    p_f2.paragraph_format.space_after = Pt(6)
    p_f2.add_run("[  ] Paso 4.2: Pulsar el Botón «Resumen de vídeo» en Studio (Sin prompts en el chat): ").bold = True
    p_f2.add_run("Dirígete al panel de Studio (a la derecha) y pulsa directamente sobre el botón ").font.color.rgb = C_DARK
    p_f2.add_run("«Resumen de vídeo»").bold = True
    p_f2.add_run(". No necesitas escribir ningún prompt en el chat ni añadir más instrucciones: simplemente pulsa el botón y observa cómo NotebookLM procesa e integra automáticamente las tres fuentes en una síntesis audiovisual.")

    add_callout_box(
        doc,
        "🎬 QUÉ HACE NOTEBOOKLM AL PULSAR «RESUMEN DE VÍDEO»",
        [
            "1. Lee al mismo tiempo la vivencia íntima del salón madrileño de 1969, la hazaña científica del Saturno V y de Fresnedillas, y la conmemoración de Euronews.",
            "2. Estructura una narrativa audiovisual equilibrada que combina imagen, texto y datos clave en un formato dinámico y moderno listo para proyectar."
        ],
        border_color="0B7285",
        bg_color="E6FCF5"
    )

    # --- TABLA DE AUTOEVALUACIÓN DEL ALUMNO ---
    p_h1 = doc.add_paragraph()
    p_h1.paragraph_format.space_before = Pt(12)
    p_h1.paragraph_format.space_after = Pt(6)
    r = p_h1.add_run("📋 Tabla de Autoevaluación y Checklist de Competencias")
    r.font.name = "Calibri"
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = C_NAVY

    t_check = doc.add_table(rows=9, cols=3)
    t_check.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_check.autofit = False
    c_col_widths = [Inches(1.0), Inches(3.8), Inches(1.7)]

    c_headers = ["Estado", "Competencia Práctica Adquirida", "¿Conseguido?"]
    for j, h in enumerate(c_headers):
        c = t_check.cell(0, j)
        c.width = c_col_widths[j]
        set_cell_background(c, "0B2545")
        set_cell_margins(c, top=100, bottom=100, left=120, right=120)
        p_h = c.paragraphs[0]
        r = p_h.add_run(h)
        r.bold = True
        r.font.name = "Calibri"
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(255, 255, 255)

    checklist_data = [
        ("Paso 1", "Sé subir un archivo TXT, un PDF y un enlace de YouTube a Fuentes.", "[  ] SÍ  /  [  ] DUDAS"),
        ("Paso 2", "Sé localizar y observar el Resumen Automático Central en la Guía del Cuaderno.", "[  ] SÍ  /  [  ] DUDAS"),
        ("Paso 3", "Sé marcar y desmarcar casillas para analizar una sola fuente de forma aislada.", "[  ] SÍ  /  [  ] DUDAS"),
        ("Paso 4", "Sé hacer preguntas concretas y pulsar las citas [1] para comprobar el texto original.", "[  ] SÍ  /  [  ] DUDAS"),
        ("Paso 5", "Sé fijar cualquier respuesta útil como Nota permanente en Studio con la chincheta (Pin).", "[  ] SÍ  /  [  ] DUDAS"),
        ("Paso 6", "He probado los botones automáticos de Studio seleccionando únicamente la fuente de YouTube.", "[  ] SÍ  /  [  ] DUDAS"),
        ("Paso 7", "Sé pulsar el botón Presentación (TXT+PDF) con estilo sobrio y realizar la revisión en vivo.", "[  ] SÍ  /  [  ] DUDAS"),
        ("Paso 8", "Sé seleccionar las 3 fuentes a la vez y pulsar directamente el botón Resumen de vídeo.", "[  ] SÍ  /  [  ] DUDAS")
    ]

    for i, (p_num, p_desc, p_eval) in enumerate(checklist_data, start=1):
        bg = "FFFFFF" if i % 2 != 0 else "F8FAFC"
        row_vals = [p_num, p_desc, p_eval]
        for j, val in enumerate(row_vals):
            c = t_check.cell(i, j)
            c.width = c_col_widths[j]
            set_cell_background(c, bg)
            set_cell_margins(c, top=70, bottom=70, left=100, right=100)
            p_cell = c.paragraphs[0]
            p_cell.paragraph_format.space_after = Pt(0)
            r = p_cell.add_run(val)
            r.font.name = "Calibri"
            r.font.size = Pt(9)
            if j == 0:
                r.bold = True
            r.font.color.rgb = C_DARK

    doc.save(TARGET_DOCX)
    print(f"✅ Documento Word actualizado con éxito en:\n   {TARGET_DOCX}")

# ---------------------------------------------------------------------------
# GENERADOR DEL DOCUMENTO PDF
# ---------------------------------------------------------------------------
def generate_pdf():
    print("\n📄 Generando Ficha Práctica Sistematizada en PDF (ReportLab)...")
    if os.path.exists(TARGET_PDF):
        try: os.remove(TARGET_PDF)
        except Exception: pass

    doc = SimpleDocTemplate(
        TARGET_PDF,
        pagesize=A4,
        leftMargin=1.8*cm,
        rightMargin=1.8*cm,
        topMargin=1.8*cm,
        bottomMargin=1.8*cm
    )

    styles = getSampleStyleSheet()

    c_primary = colors.HexColor('#0B2545')
    c_blue = colors.HexColor('#0066A1')
    c_dark = colors.HexColor('#1E293B')
    c_bg_box = colors.HexColor('#F0F7FC')

    p_header = ParagraphStyle('Head', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=8.5, textColor=c_blue, alignment=1, spaceAfter=2)
    p_title = ParagraphStyle('Title', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=16, leading=19, textColor=c_primary, alignment=1, spaceAfter=4)
    p_sub = ParagraphStyle('Sub', parent=styles['Normal'], fontName='Helvetica-Oblique', fontSize=10, leading=13, textColor=colors.HexColor('#64748B'), alignment=1, spaceAfter=10)
    
    p_h1 = ParagraphStyle('H1', parent=styles['Heading2'], fontName='Helvetica-Bold', fontSize=12, leading=15, textColor=c_primary, spaceBefore=8, spaceAfter=4)
    p_body = ParagraphStyle('Body', parent=styles['Normal'], fontName='Helvetica', fontSize=8.5, leading=11.5, textColor=c_dark, spaceAfter=3)
    p_prompt = ParagraphStyle('Prompt', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=8.5, leading=12, textColor=c_primary)
    p_box = ParagraphStyle('Box', parent=styles['Normal'], fontName='Helvetica', fontSize=8.2, leading=11.5, textColor=c_dark)
    p_cell = ParagraphStyle('Cell', parent=styles['Normal'], fontName='Helvetica', fontSize=8, leading=10.5, textColor=c_dark)
    p_cell_b = ParagraphStyle('CellB', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=8, leading=10.5, textColor=c_primary)
    p_th = ParagraphStyle('TH', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=8.2, leading=11, textColor=colors.white)

    story = []

    story.append(Paragraph("CURSO DE INTELIGENCIA ARTIFICIAL Y TECNOLOGÍA PARA ADULTOS MAYORES (60+)", p_header))
    story.append(Paragraph("TALLER MONOGRÁFICO DE NOTEBOOKLM", p_title))
    story.append(Paragraph("Ficha de Práctica Sistemática: Todas las Funciones de la Pantalla Real Paso a Paso", p_sub))
    story.append(HRFlowable(width="100%", thickness=1.5, color=c_blue, spaceAfter=8))

    # Caja Objetivo
    obj_rows = [
        [Paragraph("<b>🎯 Objetivo del Taller Práctico:</b>", p_prompt)],
        [Paragraph("Aprender a manejar Google NotebookLM de forma rigurosa y ordenada a partir de 3 fuentes complementarias sobre <b>La Llegada a la Luna (20 de julio de 1969)</b>: un texto personal (TXT), un informe histórico (PDF) y un vídeo conmemorativo (YouTube). Descubriremos el resumen automático central, cómo buscar datos con citas [1], cómo fijar notas, cómo pulsar los botones de Studio, cómo crear una presentación con estilo sobrio y revisarla en vivo, y cómo pulsar directamente el botón de resumen de vídeo.", p_box)]
    ]
    t_obj = Table(obj_rows, colWidths=[17.4*cm], splitByRow=1)
    t_obj.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), c_bg_box),
        ('BOX', (0,0), (-1,-1), 1.2, c_blue),
        ('PADDING', (0,0), (-1,-1), 6),
    ]))
    story.append(t_obj)
    story.append(Spacer(1, 6))

    # FASE 0
    story.append(Paragraph("FASE 0: Las 3 Fuentes de Partida (Panel Izquierdo)", p_h1))
    story.append(Paragraph("Carga en la columna izquierda («Fuentes») los 3 documentos preparados:", p_body))

    f_rows = [
        [Paragraph("Tipo de Fuente", p_th), Paragraph("Nombre / Enlace", p_th), Paragraph("Contenido Clave", p_th)],
        [Paragraph("<b>📄 FUENTE 1 (TXT)</b>", p_cell_b), Paragraph("Recuerdo_Luna_1969.txt", p_cell), Paragraph("Vivencia personal: calor en Madrid, tele Telefunken, voz de Jesús Hermida y lágrimas del abuelo.", p_cell)],
        [Paragraph("<b>📑 FUENTE 2 (PDF)</b>", p_cell_b), Paragraph("Historia_Apolo_11.pdf", p_cell), Paragraph("Datos técnicos: 400.000 ingenieros, Saturno V, ordenador AGC de 4 KB RAM y la estación de Fresnedillas (Madrid).", p_cell)],
        [Paragraph("<b>🎥 FUENTE 3 (YouTube)</b>", p_cell_b), Paragraph("Euronews 50 Aniversario", p_cell), Paragraph("Conmemoración oficial: 50 años del alunizaje, balance histórico y futuro de la exploración.", p_cell)]
    ]
    t_f = Table(f_rows, colWidths=[4.2*cm, 4.5*cm, 8.7*cm], splitByRow=1)
    t_f.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), c_primary),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E1')),
        ('PADDING', (0,0), (-1,-1), 4),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#F8FAFC')]),
    ]))
    story.append(t_f)
    story.append(Spacer(1, 6))

    # FASE 1
    story.append(Paragraph("FASE 1: La Zona Central (Observación, Resúmenes y Preguntas con Citas)", p_h1))
    story.append(Paragraph("<b>1.1. Observación guiada:</b> Con las 3 fuentes marcadas [✔], observa la tarjeta central «Guía del cuaderno». Comprueba cómo NotebookLM ha generado un resumen global automático sin haber tecleado nada.", p_body))
    story.append(Paragraph("<b>1.2. Resumen individual de cada fuente:</b> Deja marcada únicamente una casilla cada vez:", p_body))
    story.append(Paragraph("• <b>Solo TXT:</b> «Resume en 3 líneas los recuerdos familiares y la emoción de aquella noche.»", p_body))
    story.append(Paragraph("• <b>Solo PDF:</b> «Resume en 3 puntos técnicos los datos más asombrosos del programa Apolo 11.»", p_body))
    story.append(Paragraph("• <b>Solo YouTube:</b> «Resume en 3 frases qué conmemora el vídeo y qué destaca de la misión.»", p_body))
    story.append(Paragraph("<b>1.3. Preguntas con citas numeradas [1]:</b> Marca las 3 fuentes a la vez y pregunta:", p_body))
    story.append(Paragraph("• <i>«¿Qué marca era la tele y qué dijo el abuelo?»</i> -> Haz clic en [1]: subrayará <i>Telefunken</i> y el <i>carro de mulas</i>.", p_body))
    story.append(Paragraph("• <i>«¿Cuánta RAM tenía el Apolo 11 y qué estación española oyó a Armstrong?»</i> -> Haz clic en [2]: subrayará <i>4 KB</i> y <i>Fresnedillas</i>.", p_body))
    story.append(Paragraph("<b>1.4. Fijar como Nota (Pin):</b> Pide: <i>«Compara en una tabla la vivencia humana con la frialdad de los 4 KB»</i>. Pulsa el icono de la chincheta (Pin) sobre la respuesta para fijarla permanentemente en el corcho derecho (Studio).", p_body))
    story.append(Spacer(1, 6))

    # FASE 2
    story.append(Paragraph("FASE 2: El Panel de Studio — Los Botones Automáticos (Solo con YouTube)", p_h1))
    story.append(Paragraph("Desmarca el TXT y el PDF. Deja <b>marcado únicamente el vídeo de YouTube [✔]</b>. Ve pulsando uno a uno en Studio:", p_body))

    b_rows = [
        [Paragraph("Botón en Studio", p_th), Paragraph("Acción", p_th), Paragraph("Resultado Automático", p_th)],
        [Paragraph("1. Preguntas frecuentes (FAQ)", p_cell_b), Paragraph("Pulsa 'Preguntas frecuentes'", p_cell), Paragraph("Cuestionario con las dudas esenciales del tema.", p_cell)],
        [Paragraph("2. Guía de estudio 📚", p_cell_b), Paragraph("Pulsa 'Guía de estudio'", p_cell), Paragraph("Temario pedagógico con preguntas de autoevaluación.", p_cell)],
        [Paragraph("3. Cronología ⏱️", p_cell_b), Paragraph("Pulsa 'Línea de tiempo'", p_cell), Paragraph("Eje cronológico con los hitos temporales del vídeo.", p_cell)],
        [Paragraph("4. Documento informativo 📄", p_cell_b), Paragraph("Pulsa 'Documento informativo'", p_cell), Paragraph("Informe formal listo para imprimir o enviar.", p_cell)],
        [Paragraph("5. Tabla de contenidos 📑", p_cell_b), Paragraph("Pulsa 'Tabla de contenidos'", p_cell), Paragraph("Índice estructurado por capítulos y temas.", p_cell)],
        [Paragraph("6. Puntos clave 💡", p_cell_b), Paragraph("Pulsa 'Puntos clave'", p_cell), Paragraph("Decálogo con las ideas esenciales en viñetas.", p_cell)],
        [Paragraph("7. Ensayo / Borrador ✍️", p_cell_b), Paragraph("Pulsa 'Borrador / Ensayo'", p_cell), Paragraph("Texto continuo con redacción literaria fluida.", p_cell)],
        [Paragraph("8. Resumen breve", p_cell_b), Paragraph("Pulsa 'Resumen'", p_cell), Paragraph("Síntesis concisa en un párrafo del vídeo.", p_cell)],
        [Paragraph("9. + Añadir nota manual 📝", p_cell_b), Paragraph("Pulsa '+ Añadir nota'", p_cell), Paragraph("Corcho personal para tus propias notas.", p_cell)],
    ]
    t_b = Table(b_rows, colWidths=[4.2*cm, 4.2*cm, 9.0*cm], splitByRow=1)
    t_b.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), c_primary),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E1')),
        ('PADDING', (0,0), (-1,-1), 3.5),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#F8FAFC')]),
    ]))
    story.append(t_b)
    story.append(Spacer(1, 6))

    # FASE 3
    story.append(Paragraph("FASE 3: El Botón de Presentación y Revisión en Vivo (TXT + PDF)", p_h1))
    story.append(Paragraph("<b>3.1. Selección:</b> Marca el <b>TXT [✔] y el PDF [✔]</b> (desmarca el vídeo de YouTube).", p_body))
    story.append(Paragraph("<b>3.2. Pulsar Botón «Presentación» en Studio:</b> Haz clic en el botón «Presentación» y añade este prompt de estilo visual en la casilla de personalización:", p_body))
    story.append(Paragraph("👉 <i>«Crea una presentación con un estilo visual sobrio, elegante y de crónica documental histórica. Tonos espaciales limpios (azul marino oscuro, gris grafito y blanco de alto contraste), con tipografía grande y legible para personas mayores. Equilibra la técnica de la NASA con la emoción familiar en Madrid, con citas entrecomilladas y sin estilos infantiles ni colores estridentes.»</i>", p_prompt))
    story.append(Paragraph("<b>3.3. Revisar dentro de la presentación:</b> Entra en la presentación generada y pulsa el botón <b>«Revisar»</b> (o Editar). Realizaremos la revisión y los ajustes directamente sobre la marcha en clase según lo que queramos perfeccionar.", p_body))
    story.append(Spacer(1, 6))

    # FASE 4
    story.append(Paragraph("FASE 4: El Gran Resumen de Vídeo (Las 3 Fuentes Juntas)", p_h1))
    story.append(Paragraph("<b>4.1. Selección Total:</b> Marca las <b>3 fuentes a la vez [✔] TXT + [✔] PDF + [✔] YouTube</b>.", p_body))
    story.append(Paragraph("<b>4.2. Pulsar Botón «Resumen de vídeo» en Studio (Sin prompts en el chat):</b> Haz clic directamente en el botón <b>«Resumen de vídeo»</b> en Studio. Sin escribir nada en el chat ni añadir más instrucciones, observa cómo NotebookLM genera automáticamente el vídeo integrando la vivencia familiar, la ciencia de la NASA y el aniversario de Euronews.", p_body))
    story.append(Spacer(1, 6))

    # TABLA EVALUACIÓN
    story.append(Paragraph("📋 Checklist de Autoevaluación del Alumno", p_h1))
    chk_rows = [
        [Paragraph("Paso", p_th), Paragraph("Habilidad / Competencia Práctica", p_th), Paragraph("Autoevaluación", p_th)],
        [Paragraph("1", p_cell_b), Paragraph("Sé subir un archivo TXT, un PDF y un enlace de YouTube a Fuentes.", p_cell), Paragraph("[  ] SÍ  /  [  ] DUDAS", p_cell)],
        [Paragraph("2", p_cell_b), Paragraph("Sé localizar y observar el Resumen Automático Central en la Guía del Cuaderno.", p_cell), Paragraph("[  ] SÍ  /  [  ] DUDAS", p_cell)],
        [Paragraph("3", p_cell_b), Paragraph("Sé marcar y desmarcar casillas para analizar una sola fuente de forma aislada.", p_cell), Paragraph("[  ] SÍ  /  [  ] DUDAS", p_cell)],
        [Paragraph("4", p_cell_b), Paragraph("Sé hacer preguntas concretas y pulsar las citas [1] para comprobar el texto original.", p_cell), Paragraph("[  ] SÍ  /  [  ] DUDAS", p_cell)],
        [Paragraph("5", p_cell_b), Paragraph("Sé fijar cualquier respuesta útil como Nota permanente en Studio con la chincheta (Pin).", p_cell), Paragraph("[  ] SÍ  /  [  ] DUDAS", p_cell)],
        [Paragraph("6", p_cell_b), Paragraph("He probado los botones de Studio seleccionando únicamente la fuente de YouTube.", p_cell), Paragraph("[  ] SÍ  /  [  ] DUDAS", p_cell)],
        [Paragraph("7", p_cell_b), Paragraph("Sé pulsar el botón Presentación (TXT+PDF) con estilo sobrio y realizar la revisión en vivo.", p_cell), Paragraph("[  ] SÍ  /  [  ] DUDAS", p_cell)],
        [Paragraph("8", p_cell_b), Paragraph("Sé seleccionar las 3 fuentes a la vez y pulsar directamente el botón Resumen de vídeo.", p_cell), Paragraph("[  ] SÍ  /  [  ] DUDAS", p_cell)],
    ]
    t_chk = Table(chk_rows, colWidths=[1.8*cm, 12.0*cm, 3.6*cm], splitByRow=1)
    t_chk.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), c_primary),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E1')),
        ('PADDING', (0,0), (-1,-1), 3.5),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#F8FAFC')]),
    ]))
    story.append(t_chk)

    doc.build(story)
    print(f"✅ Documento PDF actualizado con éxito en:\n   {TARGET_PDF}")

def main():
    print("=" * 70)
    print("🚀 ACTUALIZANDO FICHA PRÁCTICA SISTEMATIZADA DE NOTEBOOKLM (DOCX Y PDF)")
    print("=" * 70)
    generate_docx()
    generate_pdf()
    print("\n🎉 ¡ACTUALIZACIÓN COMPLETADA AL 100%!")

if __name__ == "__main__":
    main()
