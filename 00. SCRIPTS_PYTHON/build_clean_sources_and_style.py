# -*- coding: utf-8 -*-
"""
Genera los 3 archivos limpios:
1. FUENTE_TEMARIO_GEMINI.docx (en 1. INTRODUCCION_GEMINI)
2. FUENTE_TEMARIO_NLM.docx (en 2. INTRODUCCION_NLM)
3. PROMPT_ESTILO_VISUAL_SOBRIO_NLM.docx (en ambas carpetas)
"""

import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

ROOT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
DIR_GEMINI = os.path.join(ROOT_DIR, "CLASES", "1. INTRODUCCION_GEMINI")
DIR_NLM = os.path.join(ROOT_DIR, "CLASES", "2. INTRODUCCION_NLM")

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=200, right=200):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def create_document(file_path, title_header, title_main, intro_note, body_text):
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    if os.path.exists(file_path):
        try: os.remove(file_path)
        except Exception: pass

    doc = docx.Document()
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)

    p_top = doc.add_paragraph()
    p_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_top = p_top.add_run(title_header)
    r_top.font.name = 'Calibri'
    r_top.font.size = Pt(11)
    r_top.font.bold = True
    r_top.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

    p_main = doc.add_paragraph()
    p_main.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_main = p_main.add_run(title_main)
    r_main.font.name = 'Calibri'
    r_main.font.size = Pt(15)
    r_main.font.bold = True
    r_main.font.color.rgb = RGBColor(0x1A, 0x0A, 0x2E)

    if intro_note:
        p_note = doc.add_paragraph()
        p_note.paragraph_format.space_before = Pt(6)
        p_note.paragraph_format.space_after = Pt(10)
        r_note = p_note.add_run(intro_note)
        r_note.font.name = 'Calibri'
        r_note.font.size = Pt(10)
        r_note.font.italic = True
        r_note.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    cell = table.cell(0, 0)
    cell.width = Inches(6.8)
    set_cell_background(cell, "F8FAFC")
    set_cell_margins(cell, top=180, bottom=180, left=220, right=220)

    for line in body_text.split('\n'):
        if not line.strip():
            p_line = cell.add_paragraph()
            p_line.paragraph_format.space_before = Pt(2)
            p_line.paragraph_format.space_after = Pt(2)
            continue
        p_line = cell.add_paragraph()
        p_line.paragraph_format.line_spacing = 1.15
        p_line.paragraph_format.space_after = Pt(2)
        r = p_line.add_run(line)
        r.font.name = 'Consolas'
        r.font.size = Pt(9.5)
        if line.startswith(('DIAPOSITIVA', 'SLIDE', 'ESTRUCTURA', 'DIRECCION', 'REGLAS')):
            r.font.bold = True
            r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        elif line.strip().startswith(('-', '•')):
            r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        else:
            r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    if len(cell.paragraphs) > 1 and not cell.paragraphs[0].text.strip():
        p_elem = cell.paragraphs[0]._p
        p_elem.getparent().remove(p_elem)

    doc.save(file_path)
    print(f"✅ Creado: {file_path}")

# ==============================================================================
# 1. FUENTE LIMPIA: GEMINI EN PC Y MÓVIL
# ==============================================================================
gemini_content = """DOCUMENTO FUENTE: CONTENIDOS DIDÁCTICOS DE GEMINI EN PC Y MÓVIL
Audiencia: Personas mayores de 60 a 75 años sin conocimientos informáticos previos.
Objetivo: Guía paso a paso para conocer la pantalla real de Gemini y practicar en casa sin miedo.

DESGLOSE CONTENIDO DIAPOSITIVA A DIAPOSITIVA:

DIAPOSITIVA 1: PORTADA
- Título: Gemini: Tu Asistente en el Ordenador y en el Móvil.
- Mensaje: Guía visual para conocer tu pantalla real y practicar en casa. En este curso no se puede romper nada: todo se aprende paso a paso.

DIAPOSITIVA 2: CÓMO ENTRAR DESDE EL ORDENADOR (EL ACCESO)
- Dirección web exacta: gemini.google.com en tu navegador habitual (Google Chrome o Microsoft Edge).
- Inicio de sesión: Con tu cuenta normal de Google (Gmail). Si ya tienes correo de Gmail, ya tienes la puerta abierta a Gemini sin registros nuevos.

DIAPOSITIVA 3: EL MAPA GENERAL DE LA PANTALLA REAL (3 ESPACIOS)
- Zona Izquierda: El Menú Lateral (tu archivador de conversaciones).
- Zona Central: La Pizarra Limpia (donde Gemini escribe sus respuestas).
- Zona Inferior: La Barra de Entrada (donde tú escribes o hablas).

DIAPOSITIVA 4: EL MENÚ LATERAL: «NUEVA CONVERSACIÓN»
- Botón superior: "+ Nueva conversación".
- Utilidad: Sirve para empezar una consulta desde cero sin mezclar temas diferentes (por ejemplo, no mezclar una duda médica con una receta de cocina).

DIAPOSITIVA 5: EL MENÚ LATERAL: «RECIENTES» (TU ARCHIVADOR AUTOMÁTICO)
- Elemento: La lista de conversaciones que aparece bajo el título "Recientes" en la barra lateral.
- Gran ventaja: Todo lo que hablas con Gemini se guarda solo automáticamente. No hace falta darle a ningún botón de guardar.
- Tranquilidad: Puedes cerrar el ordenador, volver la semana que viene, hacer clic en la conversación y continuar justo donde lo dejaste.

DIAPOSITIVA 6: LA BARRA INFERIOR: «PREGUNTA A GEMINI»
- Elemento: La barra alargada inferior con el texto interior "Pregunta a Gemini".
- Explicación: Es tu zona de escritura. Escribe como hablas normalmente, con palabras sencillas, sin tecnicismos y sin prisas.
- Envío: Pulsa la flechita a la derecha o pulsa la tecla Enter en tu teclado.

DIAPOSITIVA 7: EL MICRÓFONO 🎙️ (HABLAR EN LUGAR DE TECLEAR)
- Elemento: El icono del micrófono en el extremo derecho de la barra de mensaje.
- Cómo se usa: Haces un clic en el micro, hablas despacio y con claridad, y ves cómo tus palabras aparecen escritas solas en la pantalla.
- Ideal para: Quienes prefieren no cansarse tecleando o les cuesta escribir con teclado.

DIAPOSITIVA 8: EL BOTÓN MÁS (+) Y SUBIR FOTOS 📎
- Elemento: El botón (+) a la izquierda de la barra que despliega la opción "Subir archivos".
- Qué hace: Te permite adjuntar una foto o documento que tengas guardado en el ordenador.
- El truco fácil: También puedes simplemente arrastrar cualquier foto con el ratón directamente desde la carpeta de tu ordenador y soltarla encima de Gemini (o usar Copiar y Pegar con Ctrl+V).
- Utilidad: La IA tiene ojos: puede explicarte un monumento de una foto, leerte una carta antigua o analizar una etiqueta.

DIAPOSITIVA 9: LAS 3 MODALIDADES PRINCIPALES DE GEMINI
- 1. Modo Conversación (Texto): Para hacer preguntas cotidianas, pedir recetas, explicaciones de salud, historia o redactar textos.
- 2. Modo Imágenes: Para pedirle que pinte una fotografía o ilustración artística desde cero ("Dibuja un atardecer en el mar").
- 3. Modo Vídeos: Para crear pequeñas animaciones de vídeo de unos segundos con inteligencia artificial.

DIAPOSITIVA 10: QUÉ HACER CON LA RESPUESTA DE GEMINI
- Botón Copiar (dos folios superpuestos): Para llevarte el texto a un Word o imprimirlo.
- Botón Altavoz (Escuchar): Para que Gemini te lea la respuesta en voz alta si no quieres cansarte la vista leyendo en pantalla.
- Flecha Reintentar: Si quieres que te vuelva a explicar la misma respuesta con otras palabras más sencillas.

DIAPOSITIVA 11: GEMINI EN EL TELÉFONO MÓVIL (LA CÁMARA MÁGICA 📷)
- La aplicación oficial: Descarga gratuita desde Google Play (Android) o App Store (iPhone).
- La función estrella: El icono de la Cámara. Enfocas con el móvil una factura de la luz, el prospecto de una medicina o una planta, disparas la foto y le preguntas por voz.
- Todo sincronizado: Lo que preguntas en el móvil lo puedes consultar después en el ordenador.

DIAPOSITIVA 12: LAS 3 REGLAS DE ORO PARA PRACTICAR EN CASA
- 1ª Regla: No hay prisa. La IA no se cansa ni se impacienta nunca.
- 2ª Regla: No se puede romper nada. Si algo no sale como quieres, abres una "Nueva conversación" y vuelves a empezar.
- 3ª Regla: En este curso no hay preguntas tontas. Pregunta todo lo que tengas curiosidad por saber."""

# ==============================================================================
# 2. FUENTE LIMPIA: NOTEBOOKLM
# ==============================================================================
nlm_content = """DOCUMENTO FUENTE: CONTENIDOS DIDÁCTICOS DE NOTEBOOKLM (TU CUADERNO INTELIGENTE)
Audiencia: Personas mayores de 60 a 75 años sin conocimientos informáticos previos.
Objetivo: Guía paso a paso para entender NotebookLM, subir documentos propios y generar resúmenes y audios.

DESGLOSE CONTENIDO DIAPOSITIVA A DIAPOSITIVA:

DIAPOSITIVA 1: PORTADA
- Título: NotebookLM: Tu Cuaderno Inteligente y Archivo Personal.
- Mensaje: Pon a trabajar a la Inteligencia Artificial con tus propios documentos, recuerdos familiares y recetas. Una biblioteca inteligente que solo lee lo que tú le pides.

DIAPOSITIVA 2: ¿EN QUÉ SE DIFERENCIA DE GEMINI?
- Concepto clave: Gemini sabe de todo internet; NotebookLM solo lee los papeles que tú le subes a su mesa.
- Gran ventaja: NotebookLM no se inventa cosas. Solo responde con la verdad escrita en tus documentos.
- Demostración de rigor: Cada dato que te da incluye una etiqueta con número para que veas la página original exacta de donde lo sacó.

DIAPOSITIVA 3: CÓMO ENTRAR A NOTEBOOKLM
- Dirección web: notebooklm.google.com en Chrome o Edge.
- Identificación: Con tu misma cuenta de Google (Gmail).
- Sin registros nuevos: Entras directamente con un solo clic porque ya estás identificado en Google.

DIAPOSITIVA 4: LA PANTALLA PRINCIPAL: «TUS CUADERNOS»
- Elemento: El botón central "+ Nuevo cuaderno".
- Metáfora: Un cuaderno es como una carpeta de cartón donde guardas papeles de un mismo tema (ej: "Mis recetas familiares", "Papeles de la casa", "La historia de mi pueblo").

DIAPOSITIVA 5: EL PANEL DE FUENTES: «AÑADIR TUS PAPELES»
- Elemento: El panel lateral con el botón "+ Añadir fuentes".
- Qué puedes meter: Archivos PDF, fotos de cartas antiguas, documentos de texto o enlaces web.
- El gesto fácil: Arrastrar el documento con el ratón directamente a la pantalla o pulsar "Subir desde el ordenador".

DIAPOSITIVA 6: EL CHAT INTELIGENTE: «PREGUNTA A TUS PAPELES»
- Elemento: La barra inferior de consulta.
- Cómo funciona: Le preguntas en español normal, como si hablaras con un archivador que se ha leído todos tus papeles en tres segundos.
- Ejemplos: "Dime los ingredientes exactos de la receta de mi abuela" o "¿En qué fecha nos mudamos de casa según las cartas?".

DIAPOSITIVA 7: LAS CITAS NUMERADAS: «DEMOSTRACIÓN DE VERDAD»
- Elemento: Un párrafo de respuesta con pequeños números [1], [2] al final de las frases.
- Utilidad: Al hacer clic sobre el número [1], la pantalla salta automáticamente al párrafo exacto de tu documento donde está escrito.
- Tranquilidad: Sabes al 100% que no es una invención de la IA.

DIAPOSITIVA 8: LA GUÍA DEL CUADERNO: ESQUEMAS Y RESÚMENES
- Elemento: Los botones superiores de "Guía del cuaderno", "Preguntas frecuentes" y "Línea de tiempo".
- Utilidad: Te resume automáticamente un documento largo de 20 páginas en 5 puntos clave para que no tengas que leerlo todo de golpe.

DIAPOSITIVA 9: LA FUNCIÓN ESTRELLA: EL PROGRAMA DE RADIO (AUDIO OVERVIEW) 🎙️
- Elemento: El botón "Generar" bajo el apartado de audio.
- La magia: La IA crea una tertulia de radio en español entre dos locutores simpáticos que comentan y explican tus documentos como en una charla de sobremesa.
- Cómo disfrutarlo: Le das al "Play" y lo escuchas tranquilamente mientras tomas un café o descansas.

DIAPOSITIVA 10: CASO PRÁCTICO 1: TUS MEMORIAS Y RECETAS FAMILIARES
- Qué subes: Fotos de cuadernos viejos de cocina, recuerdos de juventud o anécdotas escritas en un folio.
- Qué te da: Un recetario ordenado por platos, o un relato continuo de tu vida para regalar a tus hijos y nietos.

DIAPOSITIVA 11: CASO PRÁCTICO 2: INFORMES MÉDICOS Y PAPELES DE CASA
- Qué subes: El PDF de un análisis de sangre, el contrato de la luz o la carta del banco.
- Qué le pides: "Explícame en palabras sencillas qué significa este informe para que lo entienda sin tecnicismos".
- Privacidad total: Tus documentos son solo tuyos y privados; nadie más puede verlos.

DIAPOSITIVA 12: LAS 3 REGLAS DE ORO DE NOTEBOOKLM
- 1ª Regla: Cuanto mejor y más claro sea el documento que subas, mejor te ayudará.
- 2ª Regla: Comprueba siempre las citas [1] para ver el papel original.
- 3ª Regla: Disfruta escuchando tus audios como si tuvieras tu propia emisora de radio en casa."""

# ==============================================================================
# 3. PROMPT DE ESTILO VISUAL (DIRECCIÓN DE ARTE)
# ==============================================================================
style_prompt = """PROMPT DE ESTILO VISUAL Y DIRECCIÓN DE ARTE PARA PRESENTACIONES
Instrucción: Copia este texto y pégalo en la ventana de personalización o prompt de generación de presentación en NotebookLM.

ORDEN DE DIRECCIÓN DE ARTE (OBLIGATORIA Y ESTRICTA):

Transforma el contenido de la fuente adjunta en una presentación visual de 12 diapositivas con el siguiente estándar estético profesional:

1. ATMÓSFERA GENERAL:
   - Estilo: Sobrio, elegante, fotorrealista y de tecnología moderna (estilo Keynote de presentación oficial de Google / Apple).
   - Tono: Profesional, pedagógico y muy limpio.

2. PALETA DE COLOR:
   - Fondos: Modo oscuro mate elegante (Gris grafito profundo #131314 o azul noche sobrio #0B0F19).
   - Textos: Blanco puro nítido (#F8FAFC) para títulos principales y gris perla (#CBD5E1) para explicaciones secundarias.
   - Acentos: Azul tecnología Google (#4285F4) en botones y detalles destacados.

3. TRATAMIENTO VISUAL DE CADA DIAPOSITIVA:
   - Representación realista de interfaces digitales limpias (pantallas de ordenador y móvil reales).
   - Si aparecen personas, deben ser fotografías fotorrealistas de adultos normales en entornos luminosos y modernos de hogar o estudio.
   - Destacar los botones y zonas con recuadros nítidos y flechas de diseño sutil.

4. PROHIBICIONES TERMINANTES (CRUCIAL):
   - PROHIBIDO el estilo cuento de hadas, acuarela, dibujo animado, cómic o ilustraciones infantiles.
   - PROHIBIDO fondos de colores cálidos estridentes (naranjas chillones, amarillos saturados o tonos pastel de fantasía).
   - PROHIBIDO interfaces de WhatsApp o elementos que no pertenezcan a la pantalla real de Google.
   - Cero robots de juguete, circuitos futuristas o elementos mágicos."""

if __name__ == '__main__':
    # 1. Fuente Gemini
    f_gem = os.path.join(DIR_GEMINI, "FUENTE_TEMARIO_GEMINI.docx")
    create_document(
        f_gem,
        "CURSO DE INTELIGENCIA ARTIFICIAL PARA PERSONAS MAYORES",
        "DOCUMENTO FUENTE: TEMARIO DE GEMINI EN PC Y MÓVIL",
        "📋 Sube este archivo directamente a la sección 'Fuentes' de tu cuaderno de NotebookLM.",
        gemini_content
    )

    # 2. Fuente NLM
    f_nlm = os.path.join(DIR_NLM, "FUENTE_TEMARIO_NLM.docx")
    create_document(
        f_nlm,
        "CURSO DE INTELIGENCIA ARTIFICIAL PARA PERSONAS MAYORES",
        "DOCUMENTO FUENTE: TEMARIO DE NOTEBOOKLM",
        "📋 Sube este archivo directamente a la sección 'Fuentes' de tu cuaderno de NotebookLM.",
        nlm_content
    )

    # 3. Prompt de Estilo Visual (en ambas carpetas para tenerlo a mano)
    p_style_gem = os.path.join(DIR_GEMINI, "PROMPT_ESTILO_VISUAL_SOBRIO_NLM.docx")
    p_style_nlm = os.path.join(DIR_NLM, "PROMPT_ESTILO_VISUAL_SOBRIO_NLM.docx")
    
    create_document(
        p_style_gem,
        "DIRECCIÓN DE ARTE PARA PRESENTACIONES NOTEBOOKLM",
        "PROMPT DE ESTILO VISUAL: SOBRIO, ELEGANTE Y REALISTA",
        "📋 Copia este texto y pégalo en la orden de generación de la presentación de NotebookLM.",
        style_prompt
    )
    create_document(
        p_style_nlm,
        "DIRECCIÓN DE ARTE PARA PRESENTACIONES NOTEBOOKLM",
        "PROMPT DE ESTILO VISUAL: SOBRIO, ELEGANTE Y REALISTA",
        "📋 Copia este texto y pégalo en la orden de generación de la presentación de NotebookLM.",
        style_prompt
    )
