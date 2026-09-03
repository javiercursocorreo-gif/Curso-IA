# -*- coding: utf-8 -*-
"""
Genera el nuevo archivo Word con el prompt corregido y blindado para Gemini en PC y Móvil
"""

import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

ROOT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TARGET_DIR = os.path.join(ROOT_DIR, "CLASES", "1. INTRODUCCION_GEMINI")

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=200, right=200):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def build_doc(file_path):
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    if os.path.exists(file_path):
        try: os.remove(file_path)
        except Exception: pass

    doc = docx.Document()

    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("CURSO DE INTELIGENCIA ARTIFICIAL PARA PERSONAS MAYORES (60+)")
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(11)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("PROMPT MAESTRO ACTUALIZADO PARA NOTEBOOKLM (NLM)\nGUÍA VISUAL PASO A PASO: TU PANTALLA DE GEMINI EN PC Y MÓVIL")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(15)
    r_sub.font.bold = True
    r_sub.font.color.rgb = RGBColor(0x1A, 0x0A, 0x2E)

    p_inst = doc.add_paragraph()
    p_inst.paragraph_format.space_before = Pt(8)
    p_inst.paragraph_format.space_after = Pt(12)
    r_inst = p_inst.add_run("📋 INSTRUCCIÓN: Copia todo el texto del recuadro inferior y pégalo directamente en el chat de tu cuaderno de NotebookLM. Este prompt ya tiene incorporadas todas las correcciones reales (menú lateral oficial con 'Recientes', cero gráficos de WhatsApp, botón '+' para 'Subir archivos' y cero alucinaciones).")
    r_inst.font.name = 'Calibri'
    r_inst.font.size = Pt(10)
    r_inst.font.italic = True
    r_inst.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    prompt_text = """Actúa como un Diseñador Curricular y Pedagógico Senior especializado en la alfabetización digital e Inteligencia Artificial para adultos mayores (65 a 75 años sin experiencia tecnológica).

Tu misión es generar el guion y contenido EXACTO diapositiva por diapositiva para crear una PRESENTACIÓN VISUAL titulada:
«GUÍA VISUAL PASO A PASO: TU PANTALLA DE GEMINI EN EL ORDENADOR Y EN EL MÓVIL»

REGLAS ESTRICTAS DE DISEÑO:
1. PROHIBIDO mencionar o dibujar interfaces de WhatsApp (confunde a los alumnos).
2. PROHIBIDO hablar de iconos antiguos de "paisaje" o atajos de teclado raros como "Ctrl+M".
3. La interfaz descrita debe coincidir exactamente con el Gemini actual: menú lateral con "Nueva conversación", "Buscar", "Imágenes", "Vídeos", "Biblioteca", "Cuadernos" y "Recientes".
4. El botón para adjuntar es el signo más (+), que despliega "Subir archivos".
5. Lenguaje cálido, claro, con letra grande, sin tecnicismos y una sola idea por diapositiva.

ESTRUCTURA EXACTA DE LAS 12 DIAPOSITIVAS:

SLIDE 1: PORTADA
- Título: Gemini: Tu Asistente en el Ordenador y en el Bolsillo.
- Subtítulo: Guía visual para conocer tu pantalla y practicar en casa sin miedo.
- Mensaje: "Aquí no se puede romper nada. Todo se aprende paso a paso."

SLIDE 2: CÓMO ENTRAR DESDE EL ORDENADOR
- Dirección web: gemini.google.com (en Chrome o Edge).
- Inicio de sesión: Con tu cuenta habitual de Google (Gmail).
- Idea clave: Si tienes correo de Gmail, ya tienes la puerta abierta a Gemini.

SLIDE 3: EL MAPA GENERAL DE LA PANTALLA (LOS 3 ESPACIOS)
- Zona 1 (Izquierda): El Menú Lateral (tu archivador de conversaciones).
- Zona 2 (Centro): La Pizarra Limpia (donde Gemini escribe sus respuestas).
- Zona 3 (Abajo): La Barra de Entrada (donde tú escribes o hablas).

SLIDE 4: EL MENÚ LATERAL: «NUEVA CONVERSACIÓN»
- El botón superior: "+ Nueva conversación" (icono de lápiz o cruz).
- Para qué sirve: Para empezar una consulta nueva desde cero sin mezclar temas (por ejemplo, no mezclar una receta de cocina con una duda médica).

SLIDE 5: EL MENÚ LATERAL: «RECIENTES» (TU ARCHIVADOR AUTOMÁTICO)
- Elemento visual: La lista de conversaciones bajo el título "Recientes" en la barra lateral.
- Tranquilidad total: Todo lo que hablas con Gemini se guarda solo. No tienes que darle a "Guardar".
- Utilidad: Puedes cerrar el ordenador, volver la semana que viene, hacer clic en la conversación y continuar justo donde lo dejaste.

SLIDE 6: LA BARRA INFERIOR: «PREGUNTA A GEMINI»
- Elemento visual: El recuadro alargado inferior con el texto "Pregunta a Gemini".
- Explicación: Es tu zona de escritura. Escribe como hablas normalmente, con tus propias palabras y sin prisas.
- El botón de envío: La flecha a la derecha o pulsar la tecla Enter en tu teclado.

SLIDE 7: EL MICRÓFONO 🎙️ (HABLAR EN LUGAR DE TECLEAR)
- Elemento visual: El icono del micrófono en el extremo derecho de la barra de mensaje.
- Cómo se usa: Haces un clic en el micro, hablas despacio y con claridad, y ves cómo tus palabras aparecen escritas solas en la pantalla.
- Ideal para: Quienes prefieren no cansarse tecleando o les cuesta escribir rápido.

SLIDE 8: EL BOTÓN MÁS (+) Y SUBIR FOTOS 📎
- Elemento visual: El botón (+) a la izquierda de la barra que despliega el menú con "Subir archivos".
- Qué hace: Te permite adjuntar una foto que tengas guardada en el ordenador.
- El truco fácil: También puedes simplemente arrastrar cualquier foto desde la carpeta de tu ordenador y soltarla encima de Gemini, o usar Copiar y Pegar (Ctrl + V).
- "La IA tiene ojos": Al subirle una foto, puede explicarte qué monumento es, leerte una carta antigua o analizar una etiqueta.

SLIDE 9: LAS 3 MODALIDADES PRINCIPALES DE GEMINI
- 1. Modo Conversación (Texto): Para hacer preguntas, pedir consejos, recetas, resumir textos o redactar felicitaciones.
- 2. Modo Imágenes: Para pedirle que pinte una lámina o dibujo artístico ("Dibuja un atardecer en el mar").
- 3. Modo Vídeos: Para crear pequeñas animaciones de vídeo de unos segundos con inteligencia artificial.

SLIDE 10: QUÉ HACER CON LA RESPUESTA DE GEMINI
- El botón Copiar (dos hojas superpuestas): Para llevarte el texto a un documento de Word o imprimirlo.
- El botón del Altavoz (Escuchar): Para que Gemini te lea la respuesta en voz alta si no te apetece leer en pantalla.
- La flecha de Reintentar: Si quieres que te lo vuelva a explicar de otra manera más sencilla.

SLIDE 11: GEMINI EN EL TELÉFONO MÓVIL (LA CÁMARA MÁGICA 📷)
- La App oficial: Disponible en Google Play (Android) y App Store (iPhone).
- La función estrella: El icono de la Cámara. Enfocas con el móvil una factura de la luz, el prospecto de una medicina o una planta, disparas la foto y le preguntas por voz.
- Todo sincronizado: Lo que haces en el móvil lo puedes ver luego en el ordenador.

SLIDE 12: LAS 3 REGLAS DE ORO PARA CASA
- 1ª Regla: No hay prisa. La IA no se impacienta ni se cansa de esperar.
- 2ª Regla: No se puede romper nada. Si algo no te gusta, abres una "Nueva conversación" y vuelves a empezar.
- 3ª Regla: No hay preguntas tontas. Pregunta todo lo que tengas curiosidad por saber."""

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    cell = table.cell(0, 0)
    cell.width = Inches(6.8)
    set_cell_background(cell, "F8FAFC")
    set_cell_margins(cell, top=200, bottom=200, left=240, right=240)

    p_box = cell.paragraphs[0]
    p_box.paragraph_format.line_spacing = 1.15
    p_box.paragraph_format.space_after = Pt(4)

    for line in prompt_text.split('\n'):
        if not line.strip():
            p_line = cell.add_paragraph()
            p_line.paragraph_format.space_before = Pt(2)
            p_line.paragraph_format.space_after = Pt(2)
            continue
        p_line = cell.add_paragraph()
        p_line.paragraph_format.line_spacing = 1.15
        p_line.paragraph_format.space_after = Pt(2)
        r = p_line.add_run(line)
        r.font.name = 'Consolas'
        r.font.size = Pt(9)
        if line.startswith(('SLIDE', 'REGLAS', 'ESTRUCTURA')):
            r.font.bold = True
            r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        elif line.strip().startswith(('-', '•')):
            r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        else:
            r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    if len(cell.paragraphs) > 1 and not cell.paragraphs[0].text.strip():
        p_elem = cell.paragraphs[0]._p
        p_elem.getparent().remove(p_elem)

    doc.save(file_path)
    print(f"✅ Archivo creado: {file_path}")

if __name__ == '__main__':
    # Creamos tanto PROMPT_PRESENTACION_GEMINI_NLM.docx como actualizamos PROMPT_INTRO_GEMINI_NLM.docx
    f1 = os.path.join(TARGET_DIR, "PROMPT_PRESENTACION_GEMINI_NLM.docx")
    f2 = os.path.join(TARGET_DIR, "PROMPT_INTRO_GEMINI_NLM.docx")
    build_doc(f1)
    build_doc(f2)
