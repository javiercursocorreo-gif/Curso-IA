# -*- coding: utf-8 -*-
"""
Genera el archivo Word: CLASE_0_PROMPT_INTRO_GEMINI_NLM.docx
en la carpeta CLASES/0. INTRODUCCION A LA IA
con el prompt maestro para diseñar la presentación de Gemini en PC y Móvil.
"""

import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

ROOT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TARGET_DIR = os.path.join(ROOT_DIR, "CLASES", "0. INTRODUCCION A LA IA")
TARGET_FILE = os.path.join(TARGET_DIR, "CLASE_0_PROMPT_INTRO_GEMINI_NLM.docx")

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=200, right=200):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def create_docx():
    os.makedirs(TARGET_DIR, exist_ok=True)
    if os.path.exists(TARGET_FILE):
        try: os.remove(TARGET_FILE)
        except Exception: pass

    doc = docx.Document()

    # Configurar márgenes a 2 cm
    for section in doc.sections:
        section.top_margin = Inches(0.8)
        section.bottom_margin = Inches(0.8)
        section.left_margin = Inches(0.8)
        section.right_margin = Inches(0.8)

    # Título Principal
    p_title = doc.add_paragraph()
    p_title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_title = p_title.add_run("CURSO DE INTELIGENCIA ARTIFICIAL PARA PERSONAS MAYORES (60+)")
    r_title.font.name = 'Calibri'
    r_title.font.size = Pt(11)
    r_title.font.bold = True
    r_title.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5) # Índigo

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_sub = p_sub.add_run("PROMPT MAESTRO PARA NOTEBOOKLM (NLM)\nGUÍA VISUAL PASO A PASO: TU PANTALLA DE GEMINI EN PC Y MÓVIL")
    r_sub.font.name = 'Calibri'
    r_sub.font.size = Pt(15)
    r_sub.font.bold = True
    r_sub.font.color.rgb = RGBColor(0x1A, 0x0A, 0x2E) # Morado oscuro

    # Instrucción para el profesor
    p_inst = doc.add_paragraph()
    p_inst.paragraph_format.space_before = Pt(8)
    p_inst.paragraph_format.space_after = Pt(12)
    r_inst = p_inst.add_run("📋 INSTRUCCIÓN: Copia todo el texto del recuadro inferior y pégalo directamente en el chat de tu cuaderno de NotebookLM (o Gemini) para generar la presentación completa diapositiva por diapositiva que se llevarán los alumnos a casa.")
    r_inst.font.name = 'Calibri'
    r_inst.font.size = Pt(10)
    r_inst.font.italic = True
    r_inst.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    # Tabla contenedor con el Prompt
    prompt_text = """Actúa como un Diseñador Curricular y Pedagógico Senior especializado en la enseñanza de tecnología e Inteligencia Artificial para adultos mayores (65 a 75 años sin conocimientos informáticos previos).

Tu misión es diseñar el contenido completo, redactado diapositiva a diapositiva, para una PRESENTACIÓN VISUAL DE RESPALDO titulada:
«GUÍA VISUAL PASO A PASO: TU PANTALLA DE GEMINI EN PC Y EN EL MÓVIL»

El objetivo de esta presentación es que el alumno, al llegar a su casa después de la clase práctica, pueda abrirla y recordar exactamente qué significa cada botón, dónde hacer clic y cómo usar Gemini sin miedo a equivocarse.

ESTRUCTURA DE LA PRESENTACIÓN (DIAPOSITIVA POR DIAPOSITIVA):

Genera para cada diapositiva:
- TÍTULO DE LA DIAPOSITIVA (Claro, grande y amigable).
- ELEMENTO VISUAL CENTRAL (Qué captura de pantalla, icono o botón debe mostrarse en grande con flechas indicativas).
- EXPLICACIÓN EN LENGUAJE COTIDIANO (Máximo 3 viñetas cortas, letra grande, sin tecnicismos, con metáforas sencillas).
- CONSEJO PRÁCTICO PARA CASA (Una frase tranquilizadora).

DESGLOSE OBLIGATORIO DE CONTENIDOS QUE DEBE CUBRIR:

1. DIAPOSITIVA 1: PORTADA
   - Título: Gemini: Tu Asistente Personal en el Ordenador y en el Bolsillo.
   - Mensaje de bienvenida y confianza.

2. DIAPOSITIVA 2: CÓMO ENTRAR DESDE EL ORDENADOR (EL ACCESO)
   - La dirección web exacta: gemini.google.com en el navegador (Chrome o Edge).
   - Iniciar sesión con la cuenta de Google (Gmail): «Si ya tienes correo de Gmail, ¡ya tienes Gemini abierto!».

3. DIAPOSITIVA 3: EL MAPA GENERAL DE LA PANTALLA (LOS 3 GRANDES ESPACIOS)
   - Visión panorámica dividida en 3 zonas marcadas con colores:
     A) El Menú de la izquierda (Tu cajón de conversaciones).
     B) La Zona central limpia (La pizarra de respuestas).
     C) La Barra inferior (Donde tú escribes o hablas).

4. DIAPOSITIVA 4: EL MENÚ LATERAL IZQUIERDO (HISTORIAL Y ORDEN)
   - El botón «+ Nueva conversación»: para empezar de cero un tema nuevo sin mezclar cosas.
   - El historial de conversaciones: cómo volver a ver una receta o un texto que te hizo la semana pasada.
   - Los tres puntitos para renombrar o borrar.

5. DIAPOSITIVA 5: LA BARRA INFERIOR (EL CORAZÓN DE GEMINI)
   - El cuadro de texto: «Escribe aquí lo que necesitas».
   - Metáfora: es exactamente igual que el recuadro para escribir en WhatsApp.
   - La flecha de envío (Enter o clic en la flechita).

6. DIAPOSITIVA 6: EL BOTÓN DEL MICRÓFONO 🎙️ (HABLAR EN VEZ DE TECLEAR)
   - Dónde está el icono del micrófono en el ordenador.
   - Cómo pulsar, hablar despacio y con naturalidad, y ver cómo tus palabras se escriben solas.
   - Ideal para quienes se cansan de escribir con el teclado.

7. DIAPOSITIVA 7: EL BOTÓN DEL CLIP (+) O IMÁGENES (ENSEÑARLE FOTOS)
   - El icono de la cruz (+) o del paisaje para adjuntar imágenes.
   - El truco del «Copiar y Pegar»: clic derecho en cualquier foto de Google y Ctrl+V en Gemini.
   - «La IA tiene ojos: si le subes una foto, la analiza al segundo».

8. DIAPOSITIVA 8: LAS MODALIDADES DE GEMINI (TEXTO, IMÁGENES Y VÍDEO)
   - Modo Conversación / Texto: para redactar, resolver dudas, menús, salud o recuerdos.
   - Modo Imágenes: cuando le pedimos que pinte algo («Fotografía de...», «Dibujo de...»).
   - Modo Vídeo: cómo reproduce o interpreta contenido en movimiento.

9. DIAPOSITIVA 9: QUÉ HACER CON LA RESPUESTA DE GEMINI
   - El botón de «Copiar» (los dos folios superpuestos) para pegar el texto en un Word o WhatsApp.
   - El botón del altavoz (Escuchar): para que el ordenador te lea la respuesta en voz alta.
   - La flecha curva de «Reintentar / Volver a generar»: si la primera respuesta no te convence del todo.

10. DIAPOSITIVA 10: GEMINI EN TU TELÉFONO MÓVIL (LA APP OFICIAL)
    - Descarga en Google Play Store (Android) o App Store (iPhone).
    - Mismo usuario y misma contraseña: lo que haces en el PC aparece en el móvil.

11. DIAPOSITIVA 11: EL SALVAVIDAS DE LA CÁMARA EN EL MÓVIL 📷
    - Los 3 toques mágicos: Abrir app ➔ Tocar la Cámara ➔ Hacer foto a una etiqueta, factura o cuadro.
    - Dictarle la duda por voz con el micrófono del móvil.

12. DIAPOSITIVA 12: LAS 3 REGLAS DE ORO PARA PRACTICAR EN CASA
    - 1ª: No se puede romper nada (si te equivocas, abres una conversación nueva).
    - 2ª: La IA nunca se enfada ni tiene prisa.
    - 3ª: No hay preguntas tontas: cuanto más curiosidad tengas, mejor te lo pasarás.

REGLAS DE TONO:
- Emplea un lenguaje sumamente respetuoso, ameno, cálido y estimulante.
- Evita anglicismos innecesarios (traduce "prompt" por "tu orden o mensaje").
- Redacta el contenido de forma directa para que pueda pasarse a diapositivas limpias de PowerPoint."""

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    cell = table.cell(0, 0)
    cell.width = Inches(6.8)
    set_cell_background(cell, "F8FAFC") # Gris azulado clarito
    set_cell_margins(cell, top=200, bottom=200, left=240, right=240)

    p_box = cell.paragraphs[0]
    p_box.paragraph_format.line_spacing = 1.15
    p_box.paragraph_format.space_after = Pt(4)

    for line in prompt_text.split('\n'):
        if not line.strip():
            p_line = cell.add_paragraph()
            p_line.paragraph_format.space_before = Pt(2)
            p_line.paragraph_format.space_after = Pt(2)
            continue
        p_line = cell.add_paragraph()
        p_line.paragraph_format.line_spacing = 1.15
        p_line.paragraph_format.space_after = Pt(2)
        r = p_line.add_run(line)
        r.font.name = 'Consolas'
        r.font.size = Pt(9)
        if line.startswith(('1.', '2.', '3.', '4.', '5.', '6.', '7.', '8.', '9.', '10.', '11.', '12.', 'ESTRUCTURA', 'DESGLOSE', 'REGLAS')):
            r.font.bold = True
            r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        elif line.strip().startswith(('-', '•')):
            r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        else:
            r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    # Eliminar primer párrafo vacío en la celda
    if len(cell.paragraphs) > 1 and not cell.paragraphs[0].text.strip():
        p_elem = cell.paragraphs[0]._p
        p_elem.getparent().remove(p_elem)

    doc.save(TARGET_FILE)
    print(f"✅ Archivo creado con éxito: {TARGET_FILE}")

if __name__ == '__main__':
    create_docx()

