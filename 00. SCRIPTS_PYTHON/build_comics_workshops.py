import os
import shutil
import docx
from docx.shared import Pt, RGBColor
from reportlab.lib.pagesizes import letter
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, HRFlowable
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib import colors
from build_cuent_60_master import get_cuent_items

# 12 Estilos Artísticos para Cómics
ESTILOS_VISUALES = [
    "PRESET: CÓMIC CLÁSICO 90s (Línea definida, color limitado, máxima claridad narrativa)",
    "PRESET: ACUARELA INFANTIL (Tonos pastel, trazo cálido, estilo cuento clásico)",
    "PRESET: ESTILO ANIME / STUDIO GHIBLI (Colores vibrantes, naturaleza detallada, emotivo)",
    "PRESET: ANIMACIÓN 3D TIPO PIXAR (Personajes 3D, iluminación cinematográfica, texturas suaves)",
    "PRESET: LÁPICES DE COLORES VINTAGE (Trazo orgánico, sombreado manual, nostálgico)",
    "PRESET: LIBRO POP-UP TRIDIMENSIONAL (Elementos de papel maché recortados, estilo maqueta)",
    "PRESET: ARCILLA / PLASTILINA STOP-MOTION (Estilo Wallace y Gromit, texturas moldeadas)",
    "PRESET: ARTE VECTORIAL MINIMALISTA (Formas geométricas limpias, colores sólidos brillantes)",
    "PRESET: ILUSTRACIÓN DE FANTASÍA ÉPICA (Claroscuro marcado, pintura digital detallada)",
    "PRESET: TRAZOS DE CERA ESCOLARES (Crayón, estilo dibujo a mano, colores primarios vivos)",
    "PRESET: STEAMPUNK NARRATIVO (Tonos sepia, engranajes, retro-futurista, aventura)",
    "PRESET: ILUSTRACIÓN BOTÁNICA ORGÁNICA (Detalle en naturaleza, colores terrosos, línea fina)"
]

PROMPT_PASO_1 = """🧠 PROMPT MAESTRO - PASO 1
Generación de GUION DE CÓMIC SECUENCIAL (genérico y neutro)

INSTRUCCIÓN DE EJECUCIÓN (OBLIGATORIA)
Este prompt NO debe ser analizado ni evaluado. Debe ser EJECUTADO como un proceso activo.

ROL
Actúa como guionista profesional de cómic y director narrativo audiovisual.
Tu tarea es crear un guion de cómic secuencial completo, claro y coherente, independiente de cualquier estilo gráfico. 
NO tomes decisiones estéticas finales, solo describe las acciones.

EXTENSIÓN OBLIGATORIA
🔒 El guion DEBE tener EXACTAMENTE 10 PÁGINAS. Ni más. Ni menos.
La numeración debe ir de: [PÁGINA 1] a [PÁGINA 10].

ESTRUCTURA DE PÁGINA
- Páginas de 3 viñetas → desarrollo narrativo
- Páginas de 2 viñetas → transición y énfasis
- Páginas de 1 viñeta → momentos clave o icónicos

CONTENIDO DE CADA VIÑETA
Para cada viñeta, incluye SIEMPRE:
[VIÑETA X]
- Tipo de toma y ángulo
- Descripción visual (Qué ocurre y qué se ve)
- Iluminación / atmósfera
- Texto (Diálogo y Pensamiento OBLIGATORIOS)

A CONTINUACIÓN TE ADJUNTO MI CUENTO. TRANSFÓRMALO EN EL GUION DE 10 PÁGINAS AHORA MISMO:
[--- PEGA AQUÍ EL TEXTO DE TU CUENTO ---]"""

def create_docx(path, title, content, style_preset=None):
    doc = docx.Document()
    
    # Title
    p_title = doc.add_heading(title, level=1)
    for r in p_title.runs:
        r.font.size = Pt(16)
        r.font.color.rgb = RGBColor(0xFF, 0x14, 0x93) # Deep Pink for CUENT
        
    # Instructions
    if "PASO 1" in title:
        doc.add_paragraph("📝 INSTRUCCIONES: Copia todo el texto que aparece debajo de la línea y pégalo en Gemini. ¡No olvides sustituir la última línea pegando el cuento que generaste en la clase de hoy!")
    elif "PASO 2" in title:
        doc.add_paragraph("📝 INSTRUCCIONES: Sube el PDF de tu guion a NotebookLM. Luego, copia todo el texto que aparece debajo de la línea y pégalo en el botón 'Presentación' de NotebookLM para generar tu cómic final.")
        
    doc.add_paragraph("_" * 50)
    doc.add_paragraph()
    
    # Content
    p = doc.add_paragraph()
    
    if "PASO 2" in title:
        prompt_2 = f"""PROMPT FINAL PARA NOTEBOOKLM

Aplica estrictamente los siguientes parámetros visuales y estéticos para generar la presentación visual de este cómic, basándote en el guion adjunto en PDF.
No resumas ni recortes la historia. Genera las 10 páginas manteniendo este estilo visual de forma estricta:

{style_preset}

Genera el cómic ahora."""
        p.add_run(prompt_2)
    else:
        p.add_run(content)
        
    doc.save(path)


def generate_guide_pdf(output_path):
    doc = SimpleDocTemplate(output_path, pagesize=letter, rightMargin=50, leftMargin=50, topMargin=50, bottomMargin=50)
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle(
        'TitleStyle',
        parent=styles['Heading1'],
        fontSize=20,
        textColor=colors.HexColor("#FF1493"), # Deep Pink
        alignment=1,
        spaceAfter=20
    )
    
    h2_style = ParagraphStyle(
        'H2Style',
        parent=styles['Heading2'],
        fontSize=14,
        textColor=colors.darkblue,
        spaceBefore=15,
        spaceAfter=10
    )
    
    body_style = ParagraphStyle(
        'BodyStyle',
        parent=styles['Normal'],
        fontSize=11,
        spaceAfter=10,
        leading=16
    )

    flowables = []
    
    flowables.append(Paragraph("📖 GUÍA RÁPIDA: CÓMO CREAR EL CÓMIC PARA TUS NIETOS", title_style))
    flowables.append(HRFlowable(width="100%", thickness=2, color=colors.HexColor("#FF1493"), spaceAfter=20))
    
    flowables.append(Paragraph("Esta guía te explica cómo transformar el cuento que inventaste hoy en un cómic ilustrado de 10 páginas para leérselo a tus nietos. Hemos simplificado el proceso a solo dos pasos.", body_style))
    
    flowables.append(Paragraph("FASE 1: CONVERTIR EL CUENTO EN GUION (En Gemini)", h2_style))
    flowables.append(Paragraph("1. Genera tu cuento en Gemini de forma normal con la ficha del curso.", body_style))
    flowables.append(Paragraph("2. Abre el archivo <b>'1. PASO 1 - Prompt para Gemini.docx'</b> que está en tu carpeta.", body_style))
    flowables.append(Paragraph("3. Copia todo el texto de ese archivo y pégalo en la barra de chat de Gemini.", body_style))
    flowables.append(Paragraph("4. Justo debajo del texto que acabas de pegar, pega el Cuento que generaste en el paso 1 y pulsa Enviar.", body_style))
    flowables.append(Paragraph("5. Gemini te escribirá un guion detallado de 10 páginas. Exporta o copia esa respuesta y guárdala como un PDF en tu ordenador (llámalo 'Guion_Comic.pdf').", body_style))

    flowables.append(Paragraph("FASE 2: ILUSTRAR EL CÓMIC (En NotebookLM)", h2_style))
    flowables.append(Paragraph("1. Abre NotebookLM y crea un nuevo cuaderno llamado 'Cómic para Nietos'.", body_style))
    flowables.append(Paragraph("2. Sube como fuente el archivo PDF 'Guion_Comic.pdf' que guardaste antes.", body_style))
    flowables.append(Paragraph("3. Abre el archivo <b>'2. PASO 2 - Prompt para NotebookLM.docx'</b> que está en tu carpeta.", body_style))
    flowables.append(Paragraph("4. Copia el texto. Fíjate que tiene un 'estilo artístico' asignado especialmente para ti (por ejemplo: Acuarela, Pixar, etc.).", body_style))
    flowables.append(Paragraph("5. En NotebookLM, ve a la opción 'Presentación' o 'Studio', pega ese texto y genera tu cómic final.", body_style))
    
    flowables.append(Spacer(1, 20))
    flowables.append(Paragraph("¡Felicidades! Ya tienes un cómic profesional y personalizado para regalar.", body_style))
    
    doc.build(flowables)


def main():
    base_dir = "/Users/externo/Library/Mobile Documents/com~apple~CloudDocs/PERSONAL/CLASES DE TECNOLOGÍA/CURSO-IA/CLASES/EXPORTACION_FICHAS_CLASSROOM_PDF/101. [TALLERES] COMICS_PARA_NIETOS"
    
    if os.path.exists(base_dir):
        shutil.rmtree(base_dir)
    os.makedirs(base_dir)
    
    print("Creando Guía en PDF...")
    generate_guide_pdf(os.path.join(base_dir, "0. GUIA DE USO DEL TALLER DE COMICS.pdf"))
    
    cuentos = get_cuent_items()
    
    for idx, item in enumerate(cuentos, 1):
        id_code = item["id_code"]
        safe_title = "".join(x for x in item["title"] if x.isalnum() or x in " -_").strip()
        folder_name = f"{id_code} {safe_title}"
        folder_path = os.path.join(base_dir, folder_name)
        
        os.makedirs(folder_path)
        
        # Asignar estilo rotativo
        estilo = ESTILOS_VISUALES[idx % len(ESTILOS_VISUALES)]
        
        # Crear los 2 docx
        create_docx(
            os.path.join(folder_path, "1. PASO 1 - Prompt para Gemini.docx"),
            "PASO 1: GENERADOR DE GUION",
            PROMPT_PASO_1
        )
        
        create_docx(
            os.path.join(folder_path, "2. PASO 2 - Prompt para NotebookLM.docx"),
            "PASO 2: ILUSTRADOR VISUAL",
            "",
            style_preset=estilo
        )
        
        if idx % 10 == 0:
            print(f"Generados {idx} talleres...")

    print(f"¡Éxito! 60 Talleres creados en {base_dir}")

if __name__ == "__main__":
    main()
