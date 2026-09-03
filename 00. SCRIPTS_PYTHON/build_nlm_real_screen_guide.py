# -*- coding: utf-8 -*-
"""
Reescribe FUENTE_TEMARIO_NLM.docx y PROMPT_ESTILO_VISUAL_SOBRIO_NLM.docx
para basar la presentación al 100% en la PANTALLA REAL de NotebookLM:
- Los 3 paneles: Fuentes (Izquierda) | Chat y Campo Prompt (Centro) | Panel Studio (Derecha)
- Explicación detallada de STUDIO (Audio Overview, Guías, FAQ, Notas)
- Explicación detallada del CAMPO PROMPT y las citas [1]
- Cero estilo cuento de hadas, 100% interfaz técnica y profesional de Google.
"""

import os
import docx
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import parse_xml
from docx.oxml.ns import nsdecls

ROOT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TARGET_DIR = os.path.join(ROOT_DIR, "CLASES", "2. INTRODUCCION_NLM")

def set_cell_background(cell, fill_hex):
    tcPr = cell._tc.get_or_add_tcPr()
    shd = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{fill_hex}"/>')
    tcPr.append(shd)

def set_cell_margins(cell, top=140, bottom=140, left=200, right=200):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = parse_xml(f'''
        <w:tcMar {nsdecls("w")}>
            <w:top w:w="{top}" w:type="dxa"/>
            <w:bottom w:w="{bottom}" w:type="dxa"/>
            <w:left w:w="{left}" w:type="dxa"/>
            <w:right w:w="{right}" w:type="dxa"/>
        </w:tcMar>
    ''')
    tcPr.append(tcMar)

def create_document(file_path, title_header, title_main, intro_note, body_text):
    os.makedirs(os.path.dirname(file_path), exist_ok=True)
    if os.path.exists(file_path):
        try: os.remove(file_path)
        except Exception: pass

    doc = docx.Document()
    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.8)
        s.right_margin = Inches(0.8)

    p_top = doc.add_paragraph()
    p_top.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_top = p_top.add_run(title_header)
    r_top.font.name = 'Calibri'
    r_top.font.size = Pt(11)
    r_top.font.bold = True
    r_top.font.color.rgb = RGBColor(0x4F, 0x46, 0xE5)

    p_main = doc.add_paragraph()
    p_main.alignment = WD_ALIGN_PARAGRAPH.CENTER
    r_main = p_main.add_run(title_main)
    r_main.font.name = 'Calibri'
    r_main.font.size = Pt(15)
    r_main.font.bold = True
    r_main.font.color.rgb = RGBColor(0x1A, 0x0A, 0x2E)

    if intro_note:
        p_note = doc.add_paragraph()
        p_note.paragraph_format.space_before = Pt(6)
        p_note.paragraph_format.space_after = Pt(10)
        r_note = p_note.add_run(intro_note)
        r_note.font.name = 'Calibri'
        r_note.font.size = Pt(10)
        r_note.font.italic = True
        r_note.font.color.rgb = RGBColor(0x47, 0x55, 0x69)

    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False

    cell = table.cell(0, 0)
    cell.width = Inches(6.8)
    set_cell_background(cell, "F8FAFC")
    set_cell_margins(cell, top=180, bottom=180, left=220, right=220)

    for line in body_text.split('\n'):
        if not line.strip():
            p_line = cell.add_paragraph()
            p_line.paragraph_format.space_before = Pt(2)
            p_line.paragraph_format.space_after = Pt(2)
            continue
        p_line = cell.add_paragraph()
        p_line.paragraph_format.line_spacing = 1.15
        p_line.paragraph_format.space_after = Pt(2)
        r = p_line.add_run(line)
        r.font.name = 'Consolas'
        r.font.size = Pt(9.5)
        if line.startswith(('SLIDE', 'DIRECCION', 'REGLAS', 'ESTRUCTURA')):
            r.font.bold = True
            r.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)
        elif line.strip().startswith(('-', '•')):
            r.font.color.rgb = RGBColor(0x33, 0x41, 0x55)
        else:
            r.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

    if len(cell.paragraphs) > 1 and not cell.paragraphs[0].text.strip():
        p_elem = cell.paragraphs[0]._p
        p_elem.getparent().remove(p_elem)

    doc.save(file_path)
    print(f"✅ Documento creado: {file_path}")

# ==============================================================================
# CONTENIDO REAL SOBRE LA PANTALLA DE NOTEBOOKLM
# ==============================================================================
nlm_real_content = """DOCUMENTO FUENTE: GUÍA VISUAL SOBRE LA PANTALLA REAL DE NOTEBOOKLM
Audiencia: Adultos mayores de 60 a 75 años sin experiencia tecnológica.
Enfoque: Explicación directa sobre la interfaz real de NotebookLM dividida en sus 3 columnas (Fuentes, Chat/Campo Prompt y Panel Studio).

DESGLOSE DIAPOSITIVA A DIAPOSITIVA:

SLIDE 1: PORTADA
- Título: NotebookLM: Tu Cuaderno Inteligente y Archivo Personal.
- Subtítulo: Conoce tu pantalla real: cómo subir tus papeles, preguntar y escuchar resúmenes en audio.
- Visual: Pantalla de ordenador mostrando la interfaz web real y moderna de Google NotebookLM con sus tres columnas.

SLIDE 2: ¿EN QUÉ SE DIFERENCIA DE GEMINI?
- Concepto real: Gemini busca en todo internet y puede opinar o equivocarse; NotebookLM es un archivador privado que SOLO lee lo que tú le subes.
- Seguridad absoluta: Si no está en tus papeles, NotebookLM te dice "no aparece en los documentos". Cero invenciones.
- Visual: Dos pantallas comparadas: el icono de Gemini (web abierta) frente a la carpeta cerrada de NotebookLM (tus propios archivos).

SLIDE 3: CÓMO ENTRAR A NOTEBOOKLM
- Dirección web: notebooklm.google.com en Chrome o Edge.
- Acceso: Clic en "Iniciar sesión" con tu cuenta habitual de Google (Gmail).
- Visual: Captura de la pantalla de bienvenida con el botón de inicio de sesión de Google.

SLIDE 4: EL MAPA GENERAL DE LA PANTALLA REAL (LAS 3 COLUMNAS)
- Columna 1 (Izquierda): Panel de FUENTES (los papeles que le das a leer).
- Columna 2 (Centro): ZONA DE CONVERSACIÓN Y CAMPO PROMPT (donde preguntas y lees las respuestas con citas).
- Columna 3 (Derecha): PANEL STUDIO (donde se generan los audios, resúmenes automáticos y guías).
- Visual: Captura completa de la interfaz real de NotebookLM con las tres zonas delimitadas y rotuladas.

SLIDE 5: COLUMNA IZQUIERDA: EL PANEL DE FUENTES («+ AÑADIR FUENTES»)
- El botón superior: "+ Añadir fuentes".
- Qué formatos admite:
  1. Subir archivos PDF o Word desde tu ordenador.
  2. Enlaces de vídeos de YouTube (lee su transcripción al segundo).
  3. Texto copiado (pegar cartas, recetas o notas manuscritas).
  4. Google Drive.
- Las casillas de activación: Puedes marcar o desmarcar qué documentos quieres que tenga en cuenta en cada momento.

SLIDE 6: COLUMNA CENTRAL: EL CAMPO PROMPT (DÓNDE Y CÓMO ESCRIBIR)
- Elemento visual: El recuadro inferior alargado de texto ("Empieza a escribir...").
- Cómo usarlo: Escribe tu pregunta con tus propias palabras cotidianas.
  * Ejemplos: "¿Qué día ocurrió el alunizaje según los informes?", "¿Qué receta contiene almendras?", "¿Cuánto dinero costaba el alquiler?".
- Envío: Pulsa la flecha o la tecla Enter.

SLIDE 7: LAS CITAS NUMERADAS [1] (LA PRUEBA DEL ALGODÓN)
- Elemento visual: El texto de respuesta con pequeños botones grises o azules numerados: [1], [2].
- La magia de la cita: Al pulsar sobre el número [1], la pantalla se desplaza y subraya el párrafo exacto en tu documento original donde está escrito el dato.
- Tranquilidad: El alumno comprueba por sí mismo que la respuesta es 100% verídica y no un invento.

SLIDE 8: COLUMNA DERECHA: EL PANEL STUDIO (LA FÁBRICA DE RESÚMENES)
- Elemento visual: El panel lateral derecho titulado "Studio".
- Para qué sirve: Es tu mesa de trabajo para crear contenidos automáticos sin tener que pedirlos uno a uno.
- Las herramientas de Studio:
  * Guía del cuaderno (Notebook Guide).
  * Preguntas frecuentes (FAQ automáticas).
  * Cronología / Línea de tiempo de los hechos.
  * Notas fijadas.

SLIDE 9: STUDIO: EL RESUMEN DE AUDIO (AUDIO OVERVIEW / PODCAST) 🎙️
- Elemento visual: El botón destacado "Generar" bajo el apartado de Audio en Studio.
- Cómo funciona: En un solo clic, crea una conversación de radio en español entre dos locutores que debaten y explican tus documentos con tono cercano y entretenido.
- La experiencia: Pulsas "Play", te pones cómodo y escuchas tus papeles convertidos en tertulia de sobremesa.

SLIDE 10: GUARDAR NOTAS EN STUDIO («+ AÑADIR NOTA»)
- Elemento visual: El botón "+ Añadir nota" en Studio.
- Utilidad: Si una respuesta del chat te gusta mucho, pulsas "Guardar en nota" y se queda archivada en tu panel derecho como un post-it permanente que nunca se pierde.

SLIDE 11: PRIVACIDAD Y CONTROL TOTAL DE TUS DATOS
- Tus papeles no se comparten: Google no utiliza tus documentos personales para entrenar modelos públicos de inteligencia artificial.
- Borrado sencillo: Puedes eliminar cualquier fuente o cuaderno entero con los tres puntitos cuando termines tu consulta.

SLIDE 12: LAS 3 REGLAS DE ORO DE NOTEBOOKLM EN CLASE
- 1ª Regla: La calidad de la respuesta depende de la calidad de tus fuentes.
- 2ª Regla: Pulsa siempre las citas [1] para ver el papel de origen.
- 3ª Regla: Explora Studio sin miedo: genera resúmenes y escucha tus audios cuantas veces quieras."""

# ==============================================================================
# PROMPT DE ESTILO VISUAL BLINDADO PARA NOTEBOOKLM
# ==============================================================================
nlm_style_prompt = """PROMPT DE DIRECCIÓN DE ARTE PARA PRESENTACIÓN DE NOTEBOOKLM
Instrucción: Pega este texto en el campo de instrucciones o prompt de generación de la presentación en NotebookLM.

ORDEN DE DIRECCIÓN DE ARTE (OBLIGATORIA Y ESTRICTA):

Transforma la fuente adjunta en una presentación visual de 12 diapositivas técnicas y pedagógicas basadas DIRECTAMENTE EN LA PANTALLA REAL DE NOTEBOOKLM:

1. ATMÓSFERA Y ENFOQUE REALISTA (CRUCIAL):
   - Estilo: Guía visual oficial de software (estilo manual técnico interactivo de Google Workspace).
   - Fondo: Pantalla real de la aplicación (gris grafito moderno #1E1F20 o pizarra tecnológica oscura #131314).
   - Interfaz gráfica: Muestra la pantalla real con sus 3 columnas: FUENTES (izquierda), CHAT Y CAMPO PROMPT (centro) y PANEL STUDIO (derecha).

2. ELEMENTOS VISUALES PRIORITARIOS:
   - Debe verse claramente la caja del CAMPO PROMPT ("Empieza a escribir...") con llamadas indicativas.
   - Debe verse el PANEL STUDIO a la derecha con el reproductor de Audio Overview (ondas de sonido, botón Play y duración) y las tarjetas de Guía del cuaderno.
   - Debe verse el botón "+ Añadir fuentes" y los iconos reales de PDF, YouTube y Texto.
   - Resalta los botones con recuadros limpios de color azul Google (#4285F4) y flechas sutiles.

3. PROHIBICIONES TOTALES:
   - PROHIBIDO el estilo cuento de hadas, fantasía, acuarela, cómics o dibujos infantiles.
   - PROHIBIDO fondos naranjas, pasteles cálidos o colores chillones.
   - PROHIBIDO metáforas abstractas irreales. Todo debe representar la interfaz de ordenador real que el alumno ve en su navegador."""

if __name__ == '__main__':
    f_content = os.path.join(TARGET_DIR, "FUENTE_TEMARIO_NLM.docx")
    f_style = os.path.join(TARGET_DIR, "PROMPT_ESTILO_VISUAL_SOBRIO_NLM.docx")

    create_document(
        f_content,
        "CURSO DE INTELIGENCIA ARTIFICIAL PARA PERSONAS MAYORES",
        "DOCUMENTO FUENTE: GUÍA VISUAL DE LA PANTALLA REAL DE NOTEBOOKLM",
        "📋 Sube este archivo a 'Fuentes' en NotebookLM. Contiene los 12 puntos explicados sobre la pantalla real: Fuentes, Campo Prompt y Panel Studio.",
        nlm_real_content
    )

    create_document(
        f_style,
        "DIRECCIÓN DE ARTE PARA PRESENTACIÓN DE NOTEBOOKLM",
        "PROMPT DE ESTILO VISUAL: PANTALLA REAL Y PANEL STUDIO",
        "📋 Copia este texto y pégalo en la orden de generación de la presentación de NotebookLM para que use la pantalla real y elimine el estilo de cuento de hadas.",
        nlm_style_prompt
    )
