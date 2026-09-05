# -*- coding: utf-8 -*-
"""
Generador de la Ficha Didáctica del Monográfico 5:
«EDUCANIETOS IA: Tutor de Matemáticas Razonadas y Sesiones en Pantalla con NotebookLM»
Genera:
- FICHA_MONOGRAFICO_EDUCANIETOS_IA.docx y .pdf
- FICHA_MONOGRAFICO_EL_ABUELO_TUTOR.docx y .pdf (para compatibilidad de enlaces)
en CLASES/5. EL_ABUELO_TUTOR_MATEMATICAS_HISTORIA/
"""

import os
import shutil
import docx
from docx.shared import Inches, Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.enum.table import WD_TABLE_ALIGNMENT
from docx.oxml import OxmlElement, parse_xml
from docx.oxml.ns import nsdecls, qn

from reportlab.lib import colors
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle, HRFlowable
from reportlab.lib.units import cm

ROOT_DIR = os.path.dirname(os.path.dirname(os.path.abspath(__file__)))
TARGET_DIR = os.path.join(ROOT_DIR, "CLASES", "5. EL_ABUELO_TUTOR_MATEMATICAS_HISTORIA")
TARGET_DOCX = os.path.join(TARGET_DIR, "FICHA_MONOGRAFICO_EDUCANIETOS_IA.docx")
TARGET_PDF = os.path.join(TARGET_DIR, "FICHA_MONOGRAFICO_EDUCANIETOS_IA.pdf")

LEGACY_DOCX = os.path.join(TARGET_DIR, "FICHA_MONOGRAFICO_EL_ABUELO_TUTOR.docx")
LEGACY_PDF = os.path.join(TARGET_DIR, "FICHA_MONOGRAFICO_EL_ABUELO_TUTOR.pdf")

os.makedirs(TARGET_DIR, exist_ok=True)

# ---------------------------------------------------------------------------
# UTILIDADES DOCX
# ---------------------------------------------------------------------------
def set_cell_background(cell, hex_color):
    shading_elm = parse_xml(f'<w:shd {nsdecls("w")} w:fill="{hex_color}"/>')
    cell._tc.get_or_add_tcPr().append(shading_elm)

def set_cell_margins(cell, top=140, bottom=140, left=180, right=180):
    tcPr = cell._tc.get_or_add_tcPr()
    tcMar = OxmlElement('w:tcMar')
    for m, val in [('top', top), ('bottom', bottom), ('left', left), ('right', right)]:
        node = OxmlElement(f'w:{m}')
        node.set(qn('w:w'), str(val))
        node.set(qn('w:type'), 'dxa')
        tcMar.append(node)
    tcPr.append(tcMar)

def add_callout_box(doc, title_text, body_paragraphs, border_color="0B4F6C", bg_color="F0F7FA"):
    table = doc.add_table(rows=1, cols=1)
    table.alignment = WD_TABLE_ALIGNMENT.CENTER
    table.autofit = False
    table.columns[0].width = Inches(6.5)

    cell = table.cell(0, 0)
    set_cell_background(cell, bg_color)
    set_cell_margins(cell, top=160, bottom=160, left=220, right=220)

    tcPr = cell._tc.get_or_add_tcPr()
    tcBorders = parse_xml(
        f'<w:tcBorders {nsdecls("w")}>'
        f'  <w:top w:val="none"/>'
        f'  <w:left w:val="single" w:sz="30" w:space="0" w:color="{border_color}"/>'
        f'  <w:bottom w:val="none"/>'
        f'  <w:right w:val="none"/>'
        f'</w:tcBorders>'
    )
    tcPr.append(tcBorders)

    p_title = cell.paragraphs[0]
    p_title.paragraph_format.space_before = Pt(0)
    p_title.paragraph_format.space_after = Pt(4)
    run_t = p_title.add_run(title_text)
    run_t.bold = True
    run_t.font.name = "Calibri"
    run_t.font.size = Pt(11)
    run_t.font.color.rgb = RGBColor.from_string(border_color)

    for body_text in body_paragraphs:
        p_body = cell.add_paragraph()
        p_body.paragraph_format.space_before = Pt(0)
        p_body.paragraph_format.space_after = Pt(3)
        run_b = p_body.add_run(body_text)
        run_b.font.name = "Calibri"
        run_b.font.size = Pt(10)
        run_b.font.color.rgb = RGBColor(40, 50, 60)

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

# ---------------------------------------------------------------------------
# GENERADOR DOCX
# ---------------------------------------------------------------------------
def generate_docx():
    print("📝 Generando Ficha Didáctica en DOCX...")
    doc = docx.Document()

    for s in doc.sections:
        s.top_margin = Inches(0.8)
        s.bottom_margin = Inches(0.8)
        s.left_margin = Inches(0.85)
        s.right_margin = Inches(0.85)

    C_NAVY = RGBColor(11, 37, 69)      # #0B2545
    C_BLUE = RGBColor(0, 102, 161)     # #0066A1
    C_DARK = RGBColor(33, 37, 41)
    C_MUTED = RGBColor(100, 116, 139)

    # Cabecera
    p_inst = doc.add_paragraph()
    p_inst.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_inst.paragraph_format.space_after = Pt(2)
    r_inst = p_inst.add_run("CURSO DE INTELIGENCIA ARTIFICIAL Y TECNOLOGÍA PARA ADULTOS MAYORES (60+)")
    r_inst.font.name = "Calibri"
    r_inst.font.size = Pt(9.5)
    r_inst.font.bold = True
    r_inst.font.color.rgb = C_BLUE

    p_main = doc.add_paragraph()
    p_main.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_main.paragraph_format.space_after = Pt(4)
    r_main = p_main.add_run("MONOGRÁFICO 5: EDUCANIETOS IA")
    r_main.font.name = "Calibri"
    r_main.font.size = Pt(22)
    r_main.font.bold = True
    r_main.font.color.rgb = C_NAVY

    p_sub = doc.add_paragraph()
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    p_sub.paragraph_format.space_after = Pt(14)
    r_sub = p_sub.add_run("Tutor de Matemáticas Razonadas con IA y Sesiones en Pantalla de Historia y Ciencias con NotebookLM")
    r_sub.font.name = "Calibri"
    r_sub.font.size = Pt(11.5)
    r_sub.font.italic = True
    r_sub.font.color.rgb = C_MUTED

    # Cuadro de enfoque
    add_callout_box(
        doc,
        "🎯 El Enfoque Pedagógico: Matemáticas en la App + Historia en Directo en Pantalla",
        [
            "En Matemáticas: Usamos la aplicación «Educanietos IA». El abuelo introduce cualquier problema en lenguaje libre y la IA lo desglosa en 4 capas pedagógicas (analogía intuitiva, paso a paso lógico, rigor de examen y reto gemelo imprimible con cuadrícula).",
            "En Historia y Ciencias: Trabajamos directamente en pantalla dentro de Google NotebookLM. Al ser el mapa mental interactivo, las tarjetas y los cuestionarios multipantalla herramientas vivas, abuelo y nieto se sientan juntos frente a la pantalla para explorar, jugar y contrastar citas sin necesidad de imprimir ni hacer capturas engorrosas.",
            "El valor insustituible del abuelo tutor: Tranquilidad, paciencia sin reproches y el afecto necesario para devolverle al nieto la seguridad en sí mismo."
        ],
        border_color="0066A1",
        bg_color="F0F7FC"
    )

    # Sección 1
    p_h1 = doc.add_paragraph()
    p_h1.paragraph_format.space_before = Pt(12)
    p_h1.paragraph_format.space_after = Pt(6)
    r = p_h1.add_run("1. Módulo de Matemáticas: La Técnica de las 4 Capas Pedagógicas")
    r.font.name = "Calibri"
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = C_NAVY

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run("Cuando un niño se atasca en matemáticas, casi siempre es porque le han enseñado la fórmula antes de que entienda la idea. En la aplicación ")
    p.add_run("«Educanietos IA»").bold = True
    p.add_run(", cualquier problema o duda planteada en lenguaje cotidiano se resuelve en 4 niveles:")

    pasos_mates = [
        ("🌟 Nivel 1: Para comprenderlo sin miedo (La analogía cotidiana)", "Una metáfora sencilla sin números que conecta con la vida real (balanzas de cocina, cintas métricas, repartos de merienda o cajas sorpresa)."),
        ("🧠 Nivel 2: El paso a paso razonado (El puente lógico)", "La explicación detallada de cada movimiento sin dar nada por sentado, eliminando los 'saltos mágicos' que confunden al estudiante."),
        ("📝 Nivel 3: Para el examen del colegio (Rigor formal)", "El desarrollo algebraico estricto con las fórmulas del currículo escolar para que el profesor le ponga la máxima calificación."),
        ("🎯 Nivel 4: El Reto Gemelo (Problema espejo imprimible)", "Un ejercicio idéntico con números cambiados. La aplicación permite imprimir una ficha en PDF con una cuadrícula milimetrada para que el nieto lo resuelva a lápiz y demuestre su dominio.")
    ]

    for tit, desc in pasos_mates:
        p_item = doc.add_paragraph()
        p_item.paragraph_format.left_indent = Inches(0.2)
        p_item.paragraph_format.space_after = Pt(3)
        r_t = p_item.add_run(tit + ": ")
        r_t.bold = True
        r_t.font.color.rgb = C_BLUE
        p_item.add_run(desc).font.color.rgb = C_DARK

    doc.add_paragraph().paragraph_format.space_after = Pt(4)

    # Sección 2
    p_h1 = doc.add_paragraph()
    p_h1.paragraph_format.space_before = Pt(12)
    p_h1.paragraph_format.space_after = Pt(6)
    r = p_h1.add_run("2. Historia y Ciencias: La Sesión en Directo sobre Pantalla con NotebookLM")
    r.font.name = "Calibri"
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = C_NAVY

    p = doc.add_paragraph()
    p.paragraph_format.space_after = Pt(4)
    p.add_run("Las funciones de estudio de NotebookLM (el Mapa Mental interactivo, las Tarjetas de memoria y los Cuestionarios dinámicos) no están pensadas para imprimirse en un papel estático ni guardarse en capturas. ").font.color.rgb = C_DARK
    p.add_run("Están diseñadas para vivirse en directo en la pantalla entre el abuelo y el nieto:").bold = True

    add_callout_box(
        doc,
        "💻 LA DINÁMICA ABUELO-NIETO FRENTE A NOTEBOOKLM",
        [
            "1. Subir los apuntes: El nieto hace 2 fotos de las páginas del libro con el móvil y se suben a Fuentes de NotebookLM. NLM extrae el texto exacto sin inventar nada.",
            "2. El Mapa Mental en pantalla: En Studio, pulsáis «Mapa mental». Vais desplegando juntos los nodos en pantalla para entender las causas y consecuencias visualmente.",
            "3. El Concurso de Tarjetas: Pulsáis «Tarjetas». El abuelo lee la pregunta en voz alta, el nieto piensa la respuesta y hacen clic en la tarjeta para comprobar si acierta.",
            "4. Cuestionario interactivo con Citas: Contestáis juntos el test pantalla a pantalla. Si el nieto duda, pulsa el número de cita y NLM resalta el renglón exacto del libro donde está la respuesta."
        ],
        border_color="0B7285",
        bg_color="E6FCF5"
    )

    # Sección 3
    p_h1 = doc.add_paragraph()
    p_h1.paragraph_format.space_before = Pt(12)
    p_h1.paragraph_format.space_after = Pt(6)
    r = p_h1.add_run("3. Guía Paso a Paso para la Sesión Práctica")
    r.font.name = "Calibri"
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = C_NAVY

    p_s1 = doc.add_paragraph()
    p_s1.paragraph_format.left_indent = Inches(0.2)
    p_s1.paragraph_format.space_after = Pt(3)
    p_s1.add_run("[  ] Paso 1: Abrir Educanietos IA en el Navegador: ").bold = True
    p_s1.add_run("Abre ").font.color.rgb = C_DARK
    p_s1.add_run("index.html").bold = True
    p_s1.add_run(" en Google Chrome o Edge. En el Módulo 1 escribe la duda matemática que tenga tu nieto.")

    p_s2 = doc.add_paragraph()
    p_s2.paragraph_format.left_indent = Inches(0.2)
    p_s2.paragraph_format.space_after = Pt(3)
    p_s2.add_run("[  ] Paso 2: Explicarle la metáfora y darle la ficha: ").bold = True
    p_s2.add_run("Léele la analogía cotidiana mientras merendáis. Pulsa «Imprimir Ficha con Cuadrícula» y dale la hoja para que intente el Reto Gemelo a solas con lápiz.")

    p_s3 = doc.add_paragraph()
    p_s3.paragraph_format.left_indent = Inches(0.2)
    p_s3.paragraph_format.space_after = Pt(3)
    p_s3.add_run("[  ] Paso 3: Sesión de Historia/Ciencias en NotebookLM: ").bold = True
    p_s3.add_run("Abre ").font.color.rgb = C_DARK
    p_s3.add_run("notebooklm.google.com").bold = True
    p_s3.add_run(". Sube las fotos de los apuntes y jugad juntos a las Tarjetas y al Cuestionario interactivo en la pantalla.")

    # Tabla de Autoevaluación
    p_h1 = doc.add_paragraph()
    p_h1.paragraph_format.space_before = Pt(12)
    p_h1.paragraph_format.space_after = Pt(6)
    r = p_h1.add_run("📋 Tabla de Autoevaluación del Abuelo Tutor")
    r.font.name = "Calibri"
    r.font.size = Pt(14)
    r.font.bold = True
    r.font.color.rgb = C_NAVY

    t_check = doc.add_table(rows=6, cols=3)
    t_check.alignment = WD_TABLE_ALIGNMENT.CENTER
    t_check.autofit = False
    c_col_widths = [Inches(1.0), Inches(3.8), Inches(1.7)]

    c_headers = ["Módulo", "Competencia del Abuelo Tutor", "¿Superado?"]
    for j, h in enumerate(c_headers):
        c = t_check.cell(0, j)
        c.width = c_col_widths[j]
        set_cell_background(c, "0B2545")
        set_cell_margins(c, top=100, bottom=100, left=120, right=120)
        p_h = c.paragraphs[0]
        r = p_h.add_run(h)
        r.bold = True
        r.font.name = "Calibri"
        r.font.size = Pt(9.5)
        r.font.color.rgb = RGBColor(255, 255, 255)

    checklist_data = [
        ("Matemáticas", "Sé escribir cualquier problema libremente en Educanietos IA y obtener la analogía intuitiva.", "[  ] SÍ  /  [  ] DUDAS"),
        ("Matemáticas", "Sé imprimir la ficha con cuadrícula para que mi nieto resuelva el Reto Gemelo a lápiz.", "[  ] SÍ  /  [  ] DUDAS"),
        ("Historia/CC", "Sé subir fotos de los apuntes o libros del nieto a Google NotebookLM como fuentes.", "[  ] SÍ  /  [  ] DUDAS"),
        ("Historia/CC", "Sé explorar el Mapa Mental y jugar a las Tarjetas interactivas en pantalla junto a mi nieto.", "[  ] SÍ  /  [  ] DUDAS"),
        ("Historia/CC", "Sé resolver el Cuestionario dinámico de NLM usando las citas directas para investigar.", "[  ] SÍ  /  [  ] DUDAS")
    ]

    for i, (p_num, p_desc, p_eval) in enumerate(checklist_data, start=1):
        bg = "FFFFFF" if i % 2 != 0 else "F8FAFC"
        row_vals = [p_num, p_desc, p_eval]
        for j, val in enumerate(row_vals):
            c = t_check.cell(i, j)
            c.width = c_col_widths[j]
            set_cell_background(c, bg)
            set_cell_margins(c, top=70, bottom=70, left=100, right=100)
            p_cell = c.paragraphs[0]
            p_cell.paragraph_format.space_after = Pt(0)
            r = p_cell.add_run(val)
            r.font.name = "Calibri"
            r.font.size = Pt(9)
            if j == 0:
                r.bold = True
            r.font.color.rgb = C_DARK

    doc.save(TARGET_DOCX)
    shutil.copyfile(TARGET_DOCX, LEGACY_DOCX)
    print(f"✅ Documento Word generado en:\n   {TARGET_DOCX}\n   {LEGACY_DOCX}")

# ---------------------------------------------------------------------------
# GENERADOR PDF
# ---------------------------------------------------------------------------
def generate_pdf():
    print("\n📄 Generando Ficha Didáctica en PDF (ReportLab)...")
    for p in [TARGET_PDF, LEGACY_PDF]:
        if os.path.exists(p):
            try: os.remove(p)
            except Exception: pass

    doc = SimpleDocTemplate(
        TARGET_PDF,
        pagesize=A4,
        leftMargin=1.8*cm,
        rightMargin=1.8*cm,
        topMargin=1.8*cm,
        bottomMargin=1.8*cm
    )

    styles = getSampleStyleSheet()

    c_primary = colors.HexColor('#0B2545')
    c_blue = colors.HexColor('#0066A1')
    c_dark = colors.HexColor('#1E293B')
    c_bg_box = colors.HexColor('#F0F7FC')

    p_header = ParagraphStyle('Head', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=8.5, textColor=c_blue, alignment=1, spaceAfter=2)
    p_title = ParagraphStyle('Title', parent=styles['Heading1'], fontName='Helvetica-Bold', fontSize=17, leading=20, textColor=c_primary, alignment=1, spaceAfter=4)
    p_sub = ParagraphStyle('Sub', parent=styles['Normal'], fontName='Helvetica-Oblique', fontSize=10, leading=13, textColor=colors.HexColor('#64748B'), alignment=1, spaceAfter=10)
    
    p_h1 = ParagraphStyle('H1', parent=styles['Heading2'], fontName='Helvetica-Bold', fontSize=12, leading=15, textColor=c_primary, spaceBefore=8, spaceAfter=4)
    p_body = ParagraphStyle('Body', parent=styles['Normal'], fontName='Helvetica', fontSize=8.5, leading=11.5, textColor=c_dark, spaceAfter=3)
    p_prompt = ParagraphStyle('Prompt', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=8.5, leading=12, textColor=c_primary)
    p_box = ParagraphStyle('Box', parent=styles['Normal'], fontName='Helvetica', fontSize=8.2, leading=11.5, textColor=c_dark)
    p_cell = ParagraphStyle('Cell', parent=styles['Normal'], fontName='Helvetica', fontSize=8, leading=10.5, textColor=c_dark)
    p_cell_b = ParagraphStyle('CellB', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=8, leading=10.5, textColor=c_primary)
    p_th = ParagraphStyle('TH', parent=styles['Normal'], fontName='Helvetica-Bold', fontSize=8.2, leading=11, textColor=colors.white)

    story = []

    story.append(Paragraph("CURSO DE INTELIGENCIA ARTIFICIAL Y TECNOLOGÍA PARA ADULTOS MAYORES (60+)", p_header))
    story.append(Paragraph("MONOGRÁFICO 5: EDUCANIETOS IA", p_title))
    story.append(Paragraph("Tutor de Matemáticas Razonadas y Sesiones en Pantalla con NotebookLM", p_sub))
    story.append(HRFlowable(width="100%", thickness=1.5, color=c_blue, spaceAfter=8))

    # Caja Enfoque
    obj_rows = [
        [Paragraph("<b>🎯 Metodología Diferenciada: Matemáticas en la App + Historia en Pantalla</b>", p_prompt)],
        [Paragraph("• <b>En Matemáticas:</b> Usamos la aplicación <b>«Educanietos IA»</b> para introducir dudas en lenguaje libre y obtener una explicación desglosada en 4 niveles (analogía cotidiana, paso a paso lógico, rigor de examen y reto gemelo imprimible en cuadrícula).<br/>• <b>En Historia y Ciencias:</b> Trabajamos en directo frente a la pantalla de <b>NotebookLM</b>. Las tarjetas interactivas, el mapa mental y el cuestionario dinámico multipantalla están pensados para explorarse cara a cara entre abuelo y nieto, investigando juntos sin necesidad de imprimir ni hacer capturas engorrosas.", p_box)]
    ]
    t_obj = Table(obj_rows, colWidths=[17.4*cm], splitByRow=1)
    t_obj.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,-1), c_bg_box),
        ('BOX', (0,0), (-1,-1), 1.2, c_blue),
        ('PADDING', (0,0), (-1,-1), 6),
    ]))
    story.append(t_obj)
    story.append(Spacer(1, 6))

    # MÓDULO MATES
    story.append(Paragraph("1. Módulo de Matemáticas: La Técnica de las 4 Capas Pedagógicas", p_h1))
    story.append(Paragraph("Cualquier problema introducido libremente en <b>Educanietos IA</b> se desglosa en 4 niveles:", p_body))
    story.append(Paragraph("• <b>1. Para comprenderlo sin miedo (Analogía cotidiana):</b> Metáfora sin números que conecta con la intuición.", p_body))
    story.append(Paragraph("• <b>2. El paso a paso razonado (El puente lógico):</b> Explicación ordenada renglón por renglón sin saltos.", p_body))
    story.append(Paragraph("• <b>3. Para el examen del colegio (Rigor formal):</b> Fórmulas y desarrollo formal del libro de texto.", p_body))
    story.append(Paragraph("• <b>4. El Reto Gemelo (Problema espejo):</b> Ejercicio idéntico con números cambiados e impresión con cuadrícula.", p_body))
    story.append(Spacer(1, 6))

    # MÓDULO HISTORIA
    story.append(Paragraph("2. Historia y Ciencias: Sesión en Vivo con NotebookLM (En Pantalla)", p_h1))
    story.append(Paragraph("Al ser herramientas interactivas y dinámicas, no tiene sentido intentar imprimirlas; se disfrutan juntos frente al ordenador:", p_body))
    story.append(Paragraph("• <b>Fotos a Fuentes:</b> Sube fotos de los apuntes o libros del nieto a NotebookLM sin transcribir nada a mano.", p_body))
    story.append(Paragraph("• <b>Mapa Mental en Vivo:</b> Exploráis juntos el árbol visual de causas y consecuencias nodo a nodo.", p_body))
    story.append(Paragraph("• <b>Concurso de Tarjetas (Flashcards):</b> El abuelo formula la pregunta y el nieto intenta adivinar antes de voltear.", p_body))
    story.append(Paragraph("• <b>Cuestionario con Citas:</b> Resuelven el test interactivo en pantalla y usan las citas para contrastar las fuentes.", p_body))
    story.append(Spacer(1, 6))

    # CÓMO USAR LA APP
    story.append(Paragraph("3. Guía Paso a Paso para la Sesión en Casa", p_h1))
    story.append(Paragraph("<b>[  ] Paso 1:</b> Abre la aplicación <code>app/index.html</code> en tu navegador web.", p_body))
    story.append(Paragraph("<b>[  ] Paso 2:</b> En Matemáticas, introduce el problema, revisa la explicación y pulsa 'Imprimir Ficha'.", p_body))
    story.append(Paragraph("<b>[  ] Paso 3:</b> En Historia, abre <code>notebooklm.google.com</code>, sube los apuntes y comparte la pantalla con tu nieto.", p_body))
    story.append(Spacer(1, 6))

    # TABLA EVALUACIÓN
    story.append(Paragraph("📋 Checklist de Autoevaluación del Abuelo Tutor", p_h1))
    chk_rows = [
        [Paragraph("Módulo", p_th), Paragraph("Habilidad / Competencia Práctica", p_th), Paragraph("Autoevaluación", p_th)],
        [Paragraph("Matemáticas", p_cell_b), Paragraph("Sé escribir cualquier problema libremente en Educanietos IA y obtener la analogía intuitiva.", p_cell), Paragraph("[  ] SÍ  /  [  ] DUDAS", p_cell)],
        [Paragraph("Matemáticas", p_cell_b), Paragraph("Sé imprimir la ficha con cuadrícula para que mi nieto resuelva el Reto Gemelo a lápiz.", p_cell), Paragraph("[  ] SÍ  /  [  ] DUDAS", p_cell)],
        [Paragraph("Historia/CC", p_cell_b), Paragraph("Sé subir fotos de los apuntes o libros del nieto a Google NotebookLM como fuentes.", p_cell), Paragraph("[  ] SÍ  /  [  ] DUDAS", p_cell)],
        [Paragraph("Historia/CC", p_cell_b), Paragraph("Sé explorar el Mapa Mental y jugar a las Tarjetas interactivas en pantalla junto a mi nieto.", p_cell), Paragraph("[  ] SÍ  /  [  ] DUDAS", p_cell)],
        [Paragraph("Historia/CC", p_cell_b), Paragraph("Sé resolver el Cuestionario dinámico de NLM usando las citas directas para investigar.", p_cell), Paragraph("[  ] SÍ  /  [  ] DUDAS", p_cell)],
    ]
    t_chk = Table(chk_rows, colWidths=[2.2*cm, 11.6*cm, 3.6*cm], splitByRow=1)
    t_chk.setStyle(TableStyle([
        ('BACKGROUND', (0,0), (-1,0), c_primary),
        ('GRID', (0,0), (-1,-1), 0.5, colors.HexColor('#CBD5E1')),
        ('PADDING', (0,0), (-1,-1), 3.5),
        ('ROWBACKGROUNDS', (0,1), (-1,-1), [colors.white, colors.HexColor('#F8FAFC')]),
    ]))
    story.append(t_chk)

    doc.build(story)
    shutil.copyfile(TARGET_PDF, LEGACY_PDF)
    print(f"✅ Documento PDF generado en:\n   {TARGET_PDF}\n   {LEGACY_PDF}")

def main():
    print("=" * 70)
    print("🚀 GENERANDO FICHA DIDÁCTICA DE EDUCANIETOS IA EN DOCX Y PDF")
    print("=" * 70)
    generate_docx()
    generate_pdf()
    print("\n🎉 ¡PROCESO COMPLETADO AL 100%!")

if __name__ == "__main__":
    main()
